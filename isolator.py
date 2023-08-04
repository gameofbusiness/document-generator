# isolator.py
# for every time you need to sort something

import generator, determiner, reader, writer
import numpy as np
import re

# needed to isolate data from existing word doc, which could be moved to the reader
from docx import Document

def isolate_detail_field(all_details, field_title):

	#print("\n=== Isolate Detail Field: " + field_title + " ===")

	detail_field_values = []

	item_name_idx = 1 # zoho
	
	coll_name_idx = 1
	
	field_idx = 0
	if field_title == "title":
		field_idx = item_name_idx
	elif field_title == "coll name":
		field_idx = coll_name_idx

	for item_idx in range(len(all_details)):
		item_details = all_details[item_idx]
		#print("Item Details: " + str(item_details))
		
		field_value = ''
		if field_title == "handle":
			field_value = generate_handle(item_details)
		elif field_title == "title": # zoho import where title is part of item name title/opt_value
			field_value = item_details[field_idx]
			field_data = field_value.split("/")
			field_value = field_data[0]
		else:
			field_value = item_details[field_idx]
		#print("Init Field Value: " + field_value)

		#print("Final Field Value: " + field_value)
		detail_field_values.append(field_value)

	#print("=== Isolated Detail Field: " + field_title + " ===\n")

	return detail_field_values

# returns list of all unique variants no longer grouped into products
def isolate_unique_variants(all_sorted_products, import_type):
	unique_variants = []
	for sorted_product in all_sorted_products:
		for variant in sorted_product:
			if generator.determine_unique_variant(variant, sorted_product, import_type):
				unique_variants.append(variant)

	return unique_variants
	
def isolate_incorrect_skus(skus):
	incorrect_skus = []
	
	init_sku_idx = 0
	part_sku_idx = 1

	init_skus = []
	part_skus = []
	for sku in skus:
		init_skus.append(sku[init_sku_idx])
		part_skus.append(sku[part_sku_idx])

	#print("init skus: " + str(skus))
	#print("part skus: " + str(skus))

	num_skus = len(init_skus)
	print("num skus: " + str(num_skus))

	for sku_idx in range(num_skus):
	
		incorrect_sku = { "sku":"", "idx":"" }
	
		init_sku = init_skus[sku_idx]
		part_sku = part_skus[sku_idx]
	
		#print("init sku: " + init_sku)
		#print("part sku: " + part_sku)

		if init_sku != part_sku:
			print("Determine correct sku for " + init_sku)
			incorrect_sku["sku"] = init_sku
			incorrect_sku["idx"] = sku_idx
			
			incorrect_skus.append(incorrect_sku)
			
	return incorrect_skus
	
def isolate_unknown_skus(inventory_skus, vrnt_skus):
	unknown_skus = []
	known_sku = False

	#print("inventory_skus: " + str(inventory_skus))
	#print("vrnt_skus: " + str(vrnt_skus))

	for inv_sku in inventory_skus:
	
		for vrnt_sku in vrnt_skus:

			if inv_sku == vrnt_sku:
				#print("SKU is known: " + inv_sku)
				known_sku = True
				break
				
		if not known_sku:
			unknown_skus.append(inv_sku)
			
	return unknown_skus
	
#def isolate_unknown_item_descriptions():
	
def isolate_data_field(all_data, table_title, field_title):

	print("\n=== Isolate Data Field: " + field_title + " ===\n")

	data_field_values = []
	
	field_value = ''
	
	field_idx = 0
	
	# aw and gb - standard comparison of aw and gb
	aw_idx = 0
	gb_idx = 1
	
	# principle data
	princ_num_idx = 0
	princ_content_idx = 1
	
	if field_title == "aw":
		field_idx = aw_idx
	elif field_title == "gb":
		#print('field_title is gb')
		field_idx = gb_idx
	
	if table_title == 'all principle data':
		
		if field_title == 'principle number':
			field_idx = princ_num_idx
		elif field_title == 'principle content':
			field_idx = princ_content_idx
		
		for item_idx in range(len(all_data)):
			item_data = all_data[item_idx]
			field_value = item_data[field_idx]
	
			data_field_values.append(field_value)
	elif table_title == 'tables':
		
		if field_title == 'aw':
			field_idx = aw_idx
		elif field_title == 'gb':
			field_idx = gb_idx
		
		for item_idx in range(len(all_data)):
			item_data = all_data[item_idx]
			field_value = item_data[field_idx]
	
			data_field_values.append(field_value)
	elif table_title == "details":
	
		# order of detail fields
		sku_idx = 0
		handle_idx = 1 # old
		collection_idx = 1
		title_idx = 2
		intro_idx = 3
		color_idx = 4
		mat_idx = 5
		finish_idx = 6
		width_idx = 7
		depth_idx = 8
		height_idx = 9
		weight_idx = 10
		features_idx = 11
		cost_idx = 12
		img_src_idx = 13
		barcode_idx = 14
		gen_vendor_idx = 15
		gen_handle_idx = 16

		item_name_idx = 1 # zoho
		
		if field_title == "init sku":
			field_idx = item_name_idx
		elif field_title == "gen handle":
			field_idx = gen_handle_idx
		elif field_title == "title":
			field_idx = title_idx

		for item_idx in range(len(all_data)):
			item_data = all_data[item_idx]
			#print("Item Details: " + str(item_details))
		
			field_value = ''
			if field_title == "handle":
				field_value = generator.generate_handle(item_data)
			else:
				field_value = item_data[field_idx]
			#print("Init Field Value: " + field_value)

			#print("Final Field Value: " + field_value)
			data_field_values.append(field_value)

		#print("=== Isolated Detail Field: " + field_title + " ===\n")
		
	elif table_title == "zoho item":
	
		for item_idx in range(len(all_data)):
			item_data = all_data[item_idx]
			#print("Item Details: " + str(item_details))
		
			field_value = ''
			if field_title == "handle":
				field_value = generator.generate_handle(item_data)
			elif field_title == "title": # zoho import where title is part of item name title/opt_value
				field_data = field_value.split("/")
				field_value = field_data[0]
			else:
				field_value = item_data[field_idx]
			#print("Init Field Value: " + field_value)

			#print("Final Field Value: " + field_value)
			data_field_values.append(field_value)
		
	elif table_title == "correct sku":
		init_sku_idx = 0
		final_sku_idx = 1
		
		if field_title == "init sku":
			field_idx = init_sku_idx
		elif field_title == "final sku":
			field_idx = final_sku_idx
			
		for item_idx in range(len(all_data)):
			item_data = all_data[item_idx]
			field_value = item_data[field_idx]
	
			data_field_values.append(field_value)
	elif table_title == "current sku" or table_title == "inventory sku":
		sku_idx = 0
		
		if field_title == "sku":
			field_idx = sku_idx
			
		for item_idx in range(len(all_data)):
			item_data = all_data[item_idx]
			field_value = item_data[field_idx]
	
			data_field_values.append(field_value)
			
	elif table_title == "name":
		#print("Isolate " + field_title + " in " + table_title + " table.")
		if field_title == "coll name" or field_title == "collection name":
			field_idx = 2
		if field_title == "handle":
			field_idx = 4
			
		for item_idx in range(len(all_data)):
			item_data = all_data[item_idx]
			if len(item_data) > field_idx:
				field_value = item_data[field_idx]
			else:
				field_value = 'n/a'
	
			data_field_values.append(field_value)
		
	elif field_title == "sku":
		for item_idx in range(len(all_data)):
			item_data = all_data[item_idx]
			field_value = item_data[field_idx]
	
			data_field_values.append(field_value)
	elif field_title == "aw" or field_title == "gb":
		for item_idx in range(len(all_data)):
			item_data = all_data[item_idx] #item_data formatted ['14. Now...win or lose.']
			item_num = item_idx + 1
			
			if len(item_data) > field_idx:
				field_value = item_data[field_idx]
				
			#print("Principle " + str(item_num) + ": \"" + field_value + "\"")
	
			data_field_values.append(field_value)
			
	else:
		print("Warning: unknown table title")

	#print("data_field_values: " + str(data_field_values))
	return data_field_values
	
def isolate_product_from_details(all_details, start_idx, stop_idx):
	product_rows = []

	for variant_idx in range(start_idx, stop_idx):
		product_rows.append(all_details[variant_idx])

	return product_rows
	
def isolate_products(all_details):
	products = []

	field_title = "handle" # we know that all variants of the same product have the same handle

	# ensure same handles grouped together, sort by handle 
	#unsorted_handles = generate_all_handles(all_details)
	#sorted_details = sorter.sort_items_by_handle(unsorted_handles, all_details)

	#if handle given in details (old version) handles = np.array(isolate_detail_field(all_details, field_title))
	#sorted_handles = np.array(generate_all_handles(sorted_details)) # details already sorted by handle
	handles = np.array(isolate_detail_field(all_details, field_title))

	_, idx, cnt = np.unique(handles, return_index=True, return_counts=True)

	unique_handles = handles[np.sort(idx)]
	counts = cnt[np.argsort(idx)]
	indices = np.sort(idx)

	num_products = len(unique_handles)

	# isolate products and append to products array
	# handles must already be adjacent b/c uses start idx + count to determine stop idx
	for product_idx in range(num_products):
		product_start_idx = indices[product_idx]
		product_stop_idx = product_start_idx + counts[product_idx]

		product_rows = isolate_product_from_details(all_details, product_start_idx, product_stop_idx)
		products.append(product_rows)

		product_start_idx = product_stop_idx
		if product_start_idx > len(all_details) - 1:
			break;

	#print("Products: " + str(products) + "\n")
	return products
	
def isolate_group_from_data(all_data, start_idx, stop_idx):
	group_rows = []

	for item_idx in range(start_idx, stop_idx):
		item_data = all_data[item_idx]
		if len(item_data) > 0:
			group_rows.append(item_data)

	return group_rows
	
def isolate_groups(all_data, field_title):

	groups = []
	
	field_values = np.array(isolate_detail_field(all_data, field_title))
	
	_, idx, cnt = np.unique(field_values, return_index=True, return_counts=True)

	unique_field_values = field_values[np.sort(idx)]
	counts = cnt[np.argsort(idx)]
	#print("counts: " + str(counts))
	indices = np.sort(idx)

	num_groups = len(unique_field_values)

	# isolate products and append to products array
	# handles must already be adjacent b/c uses start idx + count to determine stop idx
	for group_idx in range(num_groups):
		group_start_idx = indices[group_idx]
		group_stop_idx = group_start_idx + counts[group_idx]

		group_rows = isolate_group_from_data(all_data, group_start_idx, group_stop_idx)
		groups.append(group_rows)

		group_start_idx = group_stop_idx
		if group_start_idx > len(all_data) - 1:
			break;

	#print("groups: " + str(groups) + "\n")
	return groups

# each product row is of length 20 uniquely	
#def isolate_products_in_collection():

def isolate_all_item_skus(collections):

	all_item_skus = []
	
	vendor = "Global"
	keyword = "item/sku"
	data_type = "raw data"
	item_sku_idx = determiner.determine_field_idx(vendor, keyword, data_type)

	prod_row_len = 20
	for coll in collections:
		for row_idx in range(len(coll)):
			row = coll[row_idx]
			# check if next row is less than 20 to determine if pkg row
			# if pkg then do not include sku
			if not determiner.determine_pkg_row(coll, row_idx):
				item_sku = row[item_sku_idx]
				if item_sku == '': # error bc must have sku
					item_sku = 'n/a'
				all_item_skus.append(item_sku)
				
	return all_item_skus
	
def isolate_all_item_barcodes(collections):

	all_item_barcodes = []
	
	vendor = "Global"
	keyword = "upc"
	data_type = "raw data"
	item_barcode_idx = determiner.determine_field_idx(vendor, keyword, data_type)

	prod_row_len = 20
	for coll in collections:
		for row_idx in range(len(coll)):
			
			row = coll[row_idx]
			# check if next row is less than 20 to determine if pkg row
			# if pkg then do not include barcode
			if not determiner.determine_pkg_row(coll, row_idx):
				item_barcode = row[item_barcode_idx]
				if item_barcode == '':
					item_barcode = 'n/a'
				all_item_barcodes.append(item_barcode)
				
	return all_item_barcodes
	
def isolate_all_item_dims(collections):

	all_item_dims = []
	
	vendor = "Global"
	keyword = "assembled dimensions (l x w x h)"
	data_type = "raw data"
	item_dims_idx = determiner.determine_field_idx(vendor, keyword, data_type)

	prod_row_len = 20
	for coll in collections:
	
		for row_idx in range(len(coll)):
			
			row = coll[row_idx]
			
			coll_sku_idx = 0
			coll_sku = row[coll_sku_idx]
			
			# check if next row is less than 20 to determine if pkg row
			# if pkg then do not include barcode
			if not determiner.determine_pkg_row(coll, row_idx):
				item_dims = 'n/a'
				if len(row) > item_dims_idx:
					item_dims = row[item_dims_idx]
					if item_dims == '':
						item_dims = 'n/a'
				else:
					print("Warning at row index " + str(row_idx) + ": " + coll_sku + " pkg row does not have dims!")
				
				all_item_dims.append(item_dims)
				
	return all_item_dims
	
def isolate_raw_item_data(vendor, field_title, collections):

	all_raw_item_data = []
	
	data_type = "raw data"
	item_data_idx = determiner.determine_field_idx(vendor, field_title, data_type)

	for coll in collections:
		for row_idx in range(len(coll)):
			
			row = coll[row_idx]
			
			coll_sku_idx = 0
			coll_sku = row[coll_sku_idx]
			
			# check if next row is less than 20 to determine if pkg row
			# if pkg then do not include barcode
			if not determiner.determine_pkg_row(coll, row_idx):
				item_data = 'n/a'
				if len(row) > item_data_idx:
					item_data = row[item_data_idx]
					if item_data == '':
						item_data = 'n/a'
				else:
					print("Warning at row index " + str(row_idx) + ": " + coll_sku + " pkg row does not have data at idx " + str(item_data_idx) + "!")
				
				all_raw_item_data.append(item_data)
				
	return all_raw_item_data

	
def isolate_prods_in_coll(coll):
	prods = []
	prod = []
	
	prod_row_len = 20
	
	# this is 1 method
	# another method is to list row nums with new prod and then gather items bt those row nums
	for row_idx in range(len(coll)):
		row = coll[row_idx]
		if row_idx == 0:
			prod.append(row)
		elif len(row) != prod_row_len:
			prod.append(row)
		else:
			prods.append(prod)
			prod = []
			prod.append(row)
			
		if row_idx == len(coll)-1:
			prods.append(prod)
	
	return prods
	
def isolate_coll_sku(vendor, item_sku):

	coll_sku = item_sku
	if vendor == 'Global':
		sku_data = item_sku.split('-')
		
		coll_sku = sku_data[0].strip()
		
	return coll_sku
	
# in post process, rows deleted from shopify import
# now we ned zoho import to match so isolate valid ref nums
# if ref num in shopify import after post process, consider valid
def isolate_valid_ref_nums(shopify_import, zoho_import):
	print("isolate valid ref nums")
	
	valid_zoho_import = []
	
	return valid_zoho_import
	
def isolate_all_titles(vendor,all_data):

	all_titles = []
	
	data_type = "raw data"
	keyword = "item description" # description bc raw data has both titles and details in same column
	title_idx = determiner.determine_field_idx(vendor, keyword, data_type)
	
	for data in all_data:
		title = data[title_idx]
		
		all_titles.append(title)
	
	return all_titles
	
def isolate_all_dims(vendor,all_data):

	all_dims = []
	
	data_type = "raw data"
	keyword = "item dimensions" 
	dims_idx = determiner.determine_field_idx(vendor, keyword, data_type)
	
	for data in all_data:
		dims = data[dims_idx]
		
		all_dims.append(dims)
	
	return all_dims
	
def isolate_all_weights(vendor,all_data):

	all_weights = []
	
	data_type = "raw data"
	keyword = "(lbs)" 
	weight_idx = determiner.determine_field_idx(vendor, keyword, data_type)
	
	for data in all_data:
		weight = data[weight_idx]
		
		all_weights.append(weight)
	
	return all_weights
	
def isolate_all_costs(vendor,all_data):

	all_costs = []
	
	data_type = "raw data"
	keyword = "cost" 
	cost_idx = determiner.determine_field_idx(vendor, keyword, data_type)
	
	for data in all_data:
		cost = data[cost_idx]
		
		all_costs.append(cost)
	
	return all_costs
	
def isolate_unique_raw_data(all_item_data,all_item_skus):
	# isolate unique items based on sku
	print("\n=== Isolate Unique Items ===\n")
	
	final_unique_data = []
	
	all_unique_data = []
	added_skus = []
	for item_idx in range(len(all_item_data)):
		item = all_item_data[item_idx]
		item_sku = all_item_skus[item_idx]
		#print("item_sku: " + item_sku)
	
		unique_item = True
		for sku in added_skus:
			if sku == item_sku:
				unique_item = False
				break
	
		if unique_item:
			all_unique_data.append(item)
			added_skus.append(item_sku)
			
	final_unique_data = [all_unique_data,added_skus]		
	
	return final_unique_data
	
def isolate_product_in_collection(product_handle,collection):
	product_vrnts = []
	for vrnt in collection:
		vrnt_handle = generator.generate_handle(vrnt)
		if vrnt_handle == product_handle:
			product_vrnts.append(vrnt)
			
	return product_vrnts
	
#=== GB ===
# ====== Principle Generator ======

principle_num_idx = 0
principle_content_idx = 1

# for multiple translations, iso principles by lines starting with \d+\. and not ending with \w
def isolate_all_principle_data(raw_principle_data):

	#print("\n=== Isolate All Principle Data ===\n")
	
	all_principle_data = []
	
	for raw_data_translations in raw_principle_data:
		#print("raw_data_translations: " + str(raw_data_translations))
	
		trans_princ_data = []
	
		# go for all possible transformations of aw, currently only aw and gb
		for translation in raw_data_translations:
		
			#print("translation: " + translation)
		
			if determiner.determine_principle(translation):
				trans_princ_data.append(translation)
				
		#print("trans_princ_data: " + str(trans_princ_data))
		if len(trans_princ_data) > 0:
			all_principle_data.append(trans_princ_data)
	
	return all_principle_data
	
# separate/isolate princ num and princ content
def isolate_principle_num_and_content(principle):

	#print("\n=== Isolate Principle Number and Content ===\n")

	principle_data = []
	
	if len(principle) > 0:
		principle_data = re.split("(?<=\d\.)\s",principle)
		
	# ensure princ num ends with dot
	init_princ_num = principle_data[0]
	#print("init_princ_num: " + init_princ_num)
	final_princ_num = reader.add_rear_dot(init_princ_num)
	#print("final_princ_num: " + final_princ_num)
	principle_data[0] = final_princ_num
	
	return principle_data
	
def isolate_chapters(desired_principles):

	#print("\n=== Isolate Chapters ===\n")

	chapters = []
	
	current_chapter = []
	
	for principle_idx in range(len(desired_principles)):
		principle = desired_principles[principle_idx]
		#print("principle: " + principle)
		
		# if the first principle simply add to init current chap list
		if principle_idx != 0:
		
			# how to determine new chapter? number restarts at 1. content starts with sun zi said. either works.
		
			principle_data = isolate_principle_num_and_content(principle)
		
			principle_num = principle_data[principle_num_idx]
		
			# if not the first principle overall but the first principle of a chapter
			if principle_num == '1.': # if principle num = 1. (number 1 followed by dot)
				chapters.append(current_chapter)
				current_chapter = [] # reset for new chapt
	
		current_chapter.append(principle)
		
	chapters.append(current_chapter)
	
	return chapters
	
# input raw data in format: ['1. Planning','O1','1. Sun Tzu said: The art of war is of vital importance to the State.','2. Waging War']
def isolate_whole_chapters(raw_data):

	#print("\n=== Isolate Whole Chapters ===\n")

	whole_chapters = []
	
	current_chapter = []
	
	for line_idx in range(len(raw_data)):
		line = raw_data[line_idx]
		#print("line: " + line)
		
		# if the first row simply add to init current chap list
		if line_idx != 0:
		
			# how to determine new chapter? heading format like '1. Planning'. 
		
			# if not the first line overall but the first line of a chapter
			if re.search('^\d+\.\s+\w+(\s\w+)*$', line): # if line in heading format like '1. Planning'. 
				#print("New chapter found!")
				whole_chapters.append(current_chapter)
				current_chapter = [] # reset for new chapt
	
		current_chapter.append(line)
		
	whole_chapters.append(current_chapter)
	#print("whole_chapters: " + str(whole_chapters))
	
	return whole_chapters
	
# input aw_chapters is output of isolate_whole_chapters, raw data separated by chapter
def isolate_all_ch_titles(book_chapters):

	print("\n=== Isolate All Chapter Titles ===\n")

	ch_titles = []
	
	for raw_ch in book_chapters:
		# chapter title is always first line of chapter since the input is the chapters already separated
		ch_title = raw_ch[0]
		if ch_title != '':
			ch_titles.append(ch_title)
		
	print("ch_titles: " + str(ch_titles))
	
	return ch_titles
	
# input aw_chapters is output of isolate_whole_chapters, raw data separated by chapter
def isolate_all_ch_overviews(aw_chapters):

	#print("\n=== Isolate All Chapter Overviews ===\n")

	ch_overviews = [] # all overviews, which are lists of overview points themselves
	
	for raw_ch in aw_chapters:
		ch_overview = [] # list of overview points
		# chapter title is always first line of chapter since the input is the chapters already separated
		for line in raw_ch:
			if re.search("^[A-Z]",line):
				ch_overview.append(line)
		ch_overviews.append(ch_overview)
		
	#print("ch_overviews: " + str(ch_overviews))
	
	return ch_overviews
	
# input aw_chapters is output of isolate_whole_chapters, raw data separated by chapter
def isolate_all_ch_principles(aw_chapters):

	#print("\n=== Isolate All Chapter Principles ===\n")

	all_ch_principles = [] # all chapter principles, which are lists of principles themselves
	
	for raw_ch in aw_chapters:
		ch_principles = [] # list of principles
		# chapter title is always first line of chapter since the input is the chapters already separated
		for line in raw_ch:
			if determiner.determine_principle(line):
				ch_principles.append(line)
		all_ch_principles.append(ch_principles)
	
	return all_ch_principles
	
def isolate_principle_content(desired_princ_num, ch_princ_data):

	#print("\n=== Isolate Principle Content ===\n")
	
	desired_princ_content = ''
	
	for row in ch_princ_data:
	
		cur_princ_num = row[0]
		
		if cur_princ_num == desired_princ_num:
			desired_princ_content = row[1]
			break
			
	return desired_princ_content
	
def isolate_principle_paragraph(desired_princ_num, ch_princ_data, ch_paragraphs):

	print("\n=== Isolate Principle Paragraph ===\n")
	
	#print("ch_princ_data: " + str(ch_princ_data))
	#print("ch_paragraphs: " + str(ch_paragraphs))
	
	princ_paragraph = ch_paragraphs[0]
	
	for row_idx in range(len(ch_princ_data)):
		row = ch_princ_data[row_idx]
	
		cur_princ_num = row[0]
		
		if cur_princ_num == desired_princ_num:
			princ_paragraph = ch_paragraphs[row_idx]
			break
			
	return princ_paragraph
	
# isolate sub-principles in a group of principles 
# bc each sub-principle gets its own row in the table

def isolate_sub_principles(group, ch_princ_data):

	#print("\n=== Isolate Sub-Principles ===\n")

	all_sub_principles = []
	sub_princ_set = []
	
	for princ_idx in range(len(group)):
	
		#print("sub_princ_set: " + str(sub_princ_set))
		
		princ = group[princ_idx]
		#print("princ: " + princ)
		
		if princ_idx == 0:
			sub_princ_set = [princ]
			#print("sub_princ_set: " + str(sub_princ_set))
			continue
	
		if determiner.determine_sub_sub_principle(princ):
			sub_princ_set.append(princ)
			
		# if first principle in set contains the phrase "the following", 
		# then keep the following sub-principles in the same sub-princ-set (even though not sub-sub-principle)
		elif determiner.determine_strict_princ_set(princ, ch_princ_data): # strict, meaning keep on 1 page strictly
			sub_princ_set.append(princ)
			
		else:
			#print("Not sub-sub-principle, nor strict set, so start new sub-principle set!")
			all_sub_principles.append(sub_princ_set)
			#print("all_sub_principles: " + str(all_sub_principles))
			sub_princ_set = [princ] # reset the sub princ set before looping again
			
	all_sub_principles.append(sub_princ_set)
	
	return all_sub_principles
	
# must separate by chapter when you see format like "1. Planning" bc raw data also includes overview points which are not in ind
def isolate_src_chapters(all_paragraphs, book_title='aw'):

	print("\n=== Isolate Chapters ===\n")

	chapters = []
	chapter = []
	
	after_intro = False
	
	for p_idx in range(len(all_paragraphs)):
		p = all_paragraphs[p_idx]
		
		if book_title == 'aw':
			if p.text == 'Appendix': # distinct to aw bc gb ends with conclusion
				#print("Reached Appendix so we have all main content")
				chapters.append(chapter)
				break # we only need main content
		elif re.search("gb",book_title):
			if p.text == 'Conclusion': # distinct to gb bc aw ends with appendix
				#print("Reached Appendix so we have all main content")
				chapters.append(chapter)
				break # we only need main content
	
		#print("paragraph text: " + p.text)
		if re.search("^\d+\.\s+[A-Z]",p.text): # new chapter 
			#print("Found new chapter!")
			
			if len(chapter) > 0: # first time chapter title encountered, no previous chapter to add to chapters list
				#print("Length of chapter > 0 so append it to list of chapters.")
				chapters.append(chapter)
				#print("chapters: " + str(chapters))
				chapter = []
			
			if len(chapters) == 0:
				#print("Chapter 1")
				after_intro = True # only needs to be set first time we encounter ch title
			
			#print("append \"" + p.text + "\" to chapter")
			chapter.append(p)
				
		elif after_intro:
			#print("after_intro: " + str(after_intro))
			if p.text != '':
				#print("append \"" + p.text + "\" to chapter")
				chapter.append(p)
				
		#print("chapters: " + str(chapters))
	
	return chapters
	
def isolate_chapter_paragraphs(chapter):

	print("\n=== Isolate Chapter Paragraphs ===\n")

	chapter_paragraphs = []
	
	after_overview = False
	
	for p in chapter:
	
		if re.search("sun zi said",p.text.lower()):
		
			#print("Found new chapter!")
				
			if len(chapter_paragraphs) == 0:
				#print("AFTER OVERVIEW")
				after_overview = True # only needs to be set first time we encounter ch main content
				
			#print("append \"" + p.text + "\" to chapter")
			# remove line number element
			p_children = p._element.getchildren()
			#print("p_children: " + str(p_children))
			p_line_num_element = p_children[0]
			#print("p_line_num_element: " + str(p_line_num_element))
			p_line_num_element.getparent().remove(p_line_num_element)
			#print("p_children: " + str(p_children))
			chapter_paragraphs.append(p)
			
		elif after_overview:
			if p.text != '':
				#print("append \"" + p.text + "\" to chapter")
				# remove line number element
				p_children = p._element.getchildren()
				#print("p_children: " + str(p_children))
				p_line_num_element = p_children[0]
				#print("p_line_num_element: " + str(p_line_num_element))
				p_line_num_element.getparent().remove(p_line_num_element)
				#print("p_children: " + str(p_children))
				chapter_paragraphs.append(p)
				
	
	#print("chapter_paragraphs: " + str(chapter_paragraphs))
	#writer.display_ch_paragraphs(chapter_paragraphs)
	return chapter_paragraphs
	
def isolate_all_ch_paragraphs(book_title, book_idx=0):

	print("\n=== Isolate All Chapter Paragraphs for " + book_title.upper() + " ===\n")

	all_chapter_paragraphs = []
	
	if re.search(",",book_title):
		book_titles = book_title.split(",")
		cur_book_title = book_titles[book_idx]
		book_title = cur_book_title

	src_filename = "../data/" + book_title.upper() + "-input.docx"
	src_doc = Document(src_filename)
	
	src_chapters = isolate_src_chapters(src_doc.paragraphs, book_title)
	
	for ch in src_chapters:
	
		ch_paragraphs = isolate_chapter_paragraphs(ch)
		#print("ch_paragraphs: " + str(ch_paragraphs))
		
		all_chapter_paragraphs.append(ch_paragraphs)
	
	#print("all_chapter_paragraphs: " + str(all_chapter_paragraphs))
	return all_chapter_paragraphs
	
def isolate_all_ch_title_paragraphs(book_title, book_idx=0):

	print("\n=== Isolate All Chapter Title Paragraphs for " + book_title.upper() + " ===\n")

	all_chapter_title_paragraphs = []
	
# 	aw_idx=0
# 	gb_idx=1
	if re.search(",",book_title):
		book_titles = book_title.split(",")
		cur_book_title = book_titles[book_idx]
		book_title = cur_book_title

	src_filename = "../data/" + book_title.upper() + "-input.docx"
	src_doc = Document(src_filename)
	
	src_chapters = isolate_src_chapters(src_doc.paragraphs, book_title)
	
	for ch in src_chapters:
		
		for p in ch:
	
			if re.search("^\d+\.\s+[A-Z]",p.text): # new chapter 
		
				#print("Found chapter title!")
				ch_title_p = p
		
				all_chapter_title_paragraphs.append(ch_title_p)
	
	#print("all_chapter_title_paragraphs: " + str(all_chapter_title_paragraphs))
	return all_chapter_title_paragraphs
	
def isolate_ch_num(ch_title):

	ch_num = ''
	
	if ch_title != '':
		ch_num = ch_title.split('.')[0]
	else:
		print("Warning: ch_title is blank, so we cannot determine ch_num!")
		
	#print("ch_num: " + ch_num)
	return ch_num
	
def isolate_book_appendix_data(valid_ap_tables, book_title):

	print("\n=== Isolate Book Appendix Data ===\n")

	aw_idx = 0
	gb_idx = 1
	
	book_idx = aw_idx
	if book_title == 'gb':
		book_idx = gb_idx
		
	num_total_cols = len(valid_ap_tables[0])
	num_books = 2 # 2 bc 2 books included with input
	num_book_cols = num_total_cols / num_books
	
	#all_valid_ap_tables = [] # format [[c11,c12],[c21,c22]],[[c11,c12],[c21,c22]]
	
	book_ap_data = []
	
	all_book_rows = []
	for row in valid_ap_tables: # format [c11,c12,c13,c14],[c21,c22,c23,c24]
	
		cur_book_idx = 0
		book_row = [] # book row bc row of single book
		for col_idx in range(len(row)):
			cell = row[col_idx]
			#print("col_idx: " + str(col_idx))
			# if col_idx != 0 and col_idx % num_book_cols == 0:
# 				print("cur_book_idx: " + str(cur_book_idx))
# 				print("book_idx: " + str(book_idx))
# 				if cur_book_idx == book_idx:
# 					book_ap_data.append(book_row)
# 			
# 				all_book_rows.append(book_row)
# 				book_row = []
# 				
# 				cur_book_idx += 1
				
			col_num = col_idx + 1
			book_row.append(cell)
			
			if col_num % num_book_cols == 0:
			
				#print("cur_book_idx: " + str(cur_book_idx))
				#print("book_idx: " + str(book_idx))
				if cur_book_idx == book_idx:
					book_ap_data.append(book_row)
					
				all_book_rows.append(book_row)
				book_row = []
				
				cur_book_idx += 1
				
		
	print("all_book_rows: " + str(all_book_rows))	
	print("book_ap_data: " + str(book_ap_data))
	return book_ap_data	
		
		
		
