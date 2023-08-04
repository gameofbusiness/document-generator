# generator.py
# functions for a generator

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.text.tabstops import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import reader, sorter, determiner, isolator, writer, converter
import re, datetime, math
import numpy as np

# order of detail fields
sku_idx = 0
collection_idx = 1
handle_idx = 1 # OLD
title_idx = 2
intro_idx = 3
color_idx = 4
mat_idx = 5
finish_idx = 6
width_idx = 7
depth_idx = 8
height_idx = 9
weight_idx = 10
features_idx = 11
cost_idx = 12
img_src_idx = 13
barcode_idx = 14
gen_vendor_idx = 15
gen_handle_idx = 16

# order of shopify import fields
import_sku_idx = 0
import_handle_idx = 1
import_opt1_name_idx = 6
import_opt1_val_idx = 7
import_opt2_name_idx = 8
import_opt2_val_idx = 9
import_opt3_name_idx = 10
import_opt3_val_idx = 11

# order of zoho import fields
item_name_idx = 1

# get data from a file and format into a list
def extract_data(vendor, input, extension):
	vendor = vendor.title()
	catalog_filename = ''
	if input == "name":
		catalog_filename = "../Data/product-names - " + vendor + "." + extension
	elif input == "handle":
		catalog_filename = "../Data/" + vendor + "-product-import - " + input.capitalize() + "s." + extension
	else:
		vendor = re.sub(' ','-',vendor)
		#catalog_filename = catalog_filename = "../Data/" + vendor + "-catalog - " + input.capitalize() + "." + extension
		catalog_filename = catalog_filename = "../Data/" + vendor + "-catalog - New." + extension

	lines = []
	data = []
	all_data = []

	with open(catalog_filename, encoding="UTF8") as catalog_file:

		current_line = ""
		for catalog_info in catalog_file:
			current_line = catalog_info.strip()
			lines.append(current_line)

		catalog_file.close()

	# skip header line
	for line in lines[1:]:
		if len(line) > 0:
			if extension == "csv":
				data = line.split(",")
			else:
				data = line.split("\t")
		all_data.append(data)

	return all_data

# write data from a list into a file
def write_data(arranged_data, vendor, output, extension):
	output = re.sub(' ','-',output)
	catalog_filename = "../Data/" + vendor + "-" + output + "-import." + extension
	catalog_file = open(catalog_filename, "w", encoding="utf8") # overwrite existing content

	for row_idx in range(len(arranged_data)):
		catalog_file.write(arranged_data[row_idx])
		catalog_file.write("\n")
		#print(catalog[row_idx])

	catalog_file.close()
	
# handle is based on descrip
def generate_handle(item_details):

	#print("\n=== Generate Handle ===\n")

	descrip = final_title_suffix = final_handle_suffix = final_handle = ''

	descrip_data = []

	output = "title"
	all_keywords = reader.read_keywords(output)

	# look at item handle to determine type
	if len(item_details) > 0:
		# need to know field number in given table
		sku = item_details[sku_idx].strip().lower()
		collection_name = item_details[collection_idx].strip().lower()
		descrip = item_details[title_idx].strip().lower()
		vendor = item_details[gen_vendor_idx]
		
		#print("Collection Name: " + collection_name)
		# keywords in form without dashes so remove excess from descrip to compare to keywords
		
		compare_key = ''
		
		if vendor == "Klaussner" or vendor == "Comfort Design":
			vendor = "Klaussner"
			output = vendor + " title"
			all_keywords = reader.read_keywords(output)
			compare_key = sku
		else:
			compare_key = descrip

		for title_suffix, title_keywords in all_keywords.items():
			#print("Title Suffix: " + title_suffix)
			#print("Title Keywords: " + str(title_keywords))
			for keyword in title_keywords:
				#print("Keyword: " + keyword + "\n")
				if re.search(keyword,compare_key):
					final_title_suffix = title_suffix
					break

			if final_title_suffix != '':
				break

		# go from title format to handle format by adding dashes b/t words, b/c already lowercase
		final_handle_suffix = re.sub(' ','-',final_title_suffix)
		#print("Final Handle Suffix: " + final_handle_suffix)
		collection_name = re.sub(' ','-',collection_name)
		#print("collection_name: " + collection_name)

		final_handle = collection_name + "-" + final_handle_suffix
		final_handle = re.sub("\)|\(","",final_handle) # remove parentheses
		#print("final_handle: " + final_handle)
	else:
		print("Warning: No details for this item!")

	#print(final_handle)
	#print("===Final Handle: " + final_handle + "===\n")
	return final_handle

# generate all handles for a set of products
def generate_all_handles(all_details):
	all_handles = []

	for item_details in all_details:
		handle = generate_handle(item_details)
		all_handles.append(handle)

	return all_handles
	
# gen title directly from handle
def generate_title(handle):
	title = ''
	if handle != '':
		handle_words = re.split('-',handle)
		#print("Handle Words: " + str(handle_words))

		# capitalize words known as common all caps
		all_caps = ["TV","II","USB"]

		for word_idx in range(len(handle_words)):
			word = handle_words[word_idx]
			word = word.capitalize()
			
			for caps in all_caps:
				if word.lower() == caps.lower():
					word = caps
			
			title += word + ' '
			
		title = title.rstrip() # get the title with no hyphens
		
		# add hyphens to title
		all_hyphenated = ["Power-Swivel-Gliding-Reclining","Manual-Swivel-Gliding-Reclining","Power-Swivel-Rocking-Reclining","Manual-Swivel-Rocking-Reclining","Power-Gliding-Reclining","Manual-Gliding-Reclining","Power-Rocking-Reclining","Manual-Rocking-Reclining","Power-Reclining","Manual-Reclining","Counter-Height","Left-Arm-Facing","Right-Arm-Facing","Sofa-Bed","Day-Bed","Chair-Bed","Low-Profile","High-Profile","Cup-Holder","Extra-Large","Lift-Top","Sofa-Chaise"]
		for hyphenated in all_hyphenated:
			unhyphenated = re.sub("-"," ",hyphenated)
			title = re.sub(unhyphenated,hyphenated,title)
			
		# add parentheses for known important points
		all_parentheses = ["(1 Stool Only)","(1 Chair Only)","(2 Pack)","(5 Piece)","(7 Piece)"]
		for parenthesized in all_parentheses:
			unparenthesized = re.sub("\(|\)","",parenthesized)
			title = re.sub(unparenthesized,parenthesized,title)
		#print(title)
	else:
		print("Warning: Blank handle while generating title, so the title was set to an empty string (title = '')!")
	
	return title

# title is based on handle which is based on item details
def generate_title_from_details(item_details):
	title = ''

	if len(item_details) > 0:
		handle = generate_handle(item_details) #item_details[1].strip().lower()
		#print("Handle: " + handle)
		
		title = generate_title(handle) # gen title
		
	else:
		print("Warning: No Details found for this item while generating title!")

	return title

# generate all titles for a set of products
# def generate_all_titles(all_details):
# 	all_titles = []
# 
# 	for item_details in all_details:
# 		title = generate_title_from_details(item_details)
# 		all_titles.append(title)
# 
# 	return all_titles
	
def generate_all_titles(all_inputs,input_type="details"):
	all_titles = []

	for input in all_inputs:
		if input_type == 'handles':
			title = generate_title(input)
		else:
			title = generate_title_from_details(input)
			
		all_titles.append(title)

	return all_titles
	
# generate all titles for a set of products based on handles
# def generate_all_titles_from_handles(all_handles):
# 	all_titles = []
# 
# 	for handle in all_handles:
# 		title = generate_title_from_handle(handle)
# 		all_titles.append(title)
# 
# 	return all_titles

# tags based on vendor, publication year, color, material, and finish
def generate_tags(item_details, vendor):
	now = datetime.datetime.now()
	publication_year = str(now.year) # get current year

	sku = handle = colors = materials = finishes = ''

	color_data = []
	material_data = []
	finish_data = []

	tags = color_tags = material_tags = finish_tags = ''

	if len(item_details) > 0:
		sku = item_details[sku_idx].strip().lower()
		handle = generate_handle(item_details) #item_details[1].strip().lower()
		colors = item_details[4].strip().lower()
		#print("Colors: " + colors)
		materials = item_details[5].strip().lower()
		materials = re.sub('full ','',materials)
		materials = re.sub(' front','',materials)
		materials = re.sub(' back','',materials)
		#print("Materials: " + materials)
		finishes = item_details[6].strip().lower()
		#print("Finishes: " + finishes)

	if colors != "n/a":
		color_data = re.split(',|/|&|\\band|\\bwith',colors)
	#print("Color Data: " + str(color_data))
	if materials != "n/a":
		material_data = re.split(',|/|&|\\band|\\bwith',materials)
	#print("Material Data: " + str(material_data))
	if finishes != "n/a":
		finish_data = re.split(',|/|&|\\band|\\bwith',finishes)
	#print("Finish Data: " + str(finish_data))

	for color in color_data:
		color = color.strip()
		color = color.rstrip(' -') # for Global but maybe also for others
		color = color.rstrip(' w') # b/c splits on slash so abbrev w/ needs special handling
		if color != '':
			color_tags += "color-" + color + ", "
	color_tags = color_tags.rstrip(', ')
	#print("Color Tags: " + color_tags)
	for material in material_data:
		material = material.strip()
		material = material.rstrip(' -') # for Global but maybe also for others
		material = material.rstrip(' w') # b/c splits on slash so abbrev w/ needs special handling
		if material != '':
			material_tags += "material-" + material + ", "
	material_tags = material_tags.rstrip(', ')
	#print("Material Tags: " + material_tags)
	for finish in finish_data:
		finish = finish.strip()
		finish = finish.rstrip(' -') # for Global but maybe also for others
		finish = finish.rstrip(' w') # b/c splits on slash so abbrev w/ needs special handling
		if finish != '':
			finish_tags += "finish-" + finish + ", "
	finish_tags = finish_tags.rstrip(', ')
	#print("Finish Tags: " + finish_tags)

	vendor = re.sub("\s","",vendor)
	tags = vendor + publication_year

	if colors != "n/a":
		tags += ", " + color_tags

	if materials != "n/a":
		tags += ", " + material_tags

	if finishes != "n/a":
		tags += ", " + finish_tags
		
	product_type = generate_product_type(item_details)
	if product_type == "sofabeds" or product_type == "sofa chaise beds":
		tags += ", sleeper"
		
	elif product_type == "futons":
		tags += ", futon"

	return tags

def generate_all_tags(all_details, vendor):
	all_tags = []

	for item_details in all_details:
		tags = generate_tags(item_details, vendor)
		all_tags.append(tags)

	return all_tags

def generate_product_type(item_details):
	handle = final_type = ''

	handle_data = []

	output = "type"
	all_keywords = reader.read_keywords(output)

	# look at item handle to determine type
	if len(item_details) > 0:
		handle = generate_handle(item_details) #item_details[1].strip().lower() # need to know field number in given table

		if handle != '':
			# keywords in form without dashes so remove dashes from handle to compare to keywords
			dashless_handle = re.sub('-', ' ', handle)

			for type, type_keywords in all_keywords.items():
				#print("Type: " + type)
				#print("Type Keywords: " + str(type_keywords))
				for keyword in type_keywords:
					#print("Keyword: " + keyword)
					if re.search(keyword,dashless_handle):
						final_type = type
						break

				if final_type != '':
					break
		else:
			print("Warning: Blank handle while generating type, so the type was set to an empty string (type = '')!")
	else:
		print("Warning: No details for this item!")

	return final_type

def generate_all_product_types(all_details):
	all_product_types = []

	for item_details in all_details:
		types = generate_product_type(item_details)
		all_product_types.append(types)

	return all_product_types

def generate_options(item_details, init_item_details, handle=''):

	#print("\n=== Generate Options ===\n")

	init_width = init_item_details[width_idx].strip().lower()
	#print("Init Width: " + init_width)
	if handle == '':
		handle = generate_handle(item_details) #item_details[handle_idx]
	meas_type = reader.determine_measurement_type(init_width, handle)

	sku = color = title = ''

	output = "option"
	all_keywords = reader.read_keywords(output)

	final_opt_names = []
	final_opt_values = []

	# look at item sku to determine options
	# if nothing apparent from sku, then check other fields like color and material
	# do not rely entirely on sku b/c could be ambiguous codes that may appear as part of other words not meant to indicate options
	# example: W is code for wenge brown for vendor=Global, but W is likely to mean something else for other vendors
	if len(item_details) > 0:
		sku = item_details[sku_idx].strip().lower()
		#print("===Generate Options for SKU: " + sku)
		title = item_details[title_idx].strip().lower()
		color = item_details[color_idx].strip().lower()
		material = item_details[mat_idx].strip().lower()
		#print("Color: " + color)

		# option codes must only be considered valid when they are the entire word in the sku, so must remove dashes to separate words and isolate codes
		dashless_sku = re.sub('-',' ',sku)

		final_opt_string = ''

		type = generate_product_type(item_details)
		if type == 'rugs':
			option_name = 'Size' # width-depth combos are options for rugs
			# see if dims given
			width = item_details[width_idx].strip()
			depth = item_details[depth_idx].strip()
			if width != 'n/a' and depth != 'n/a':
				dim_string = width + "\" x " + depth + "\""
				if meas_type == 'round':
					dim_string = width + "\" Diameter"
				final_opt_values.append(dim_string)
				final_opt_names.append(option_name)

		# loop for each type of option, b/c need to fill in value for each possible option (eg loop for size and then loop for color in case item has both size and color options)
		for option_name, option_dict in all_keywords.items():
			#print("======Check for Option Name: " + option_name)
			#print("Option Dict: " + str(option_dict))

			final_opt_value = ''

			for option_value, option_keywords in option_dict.items():
				#print("Option Value: " + option_value)
				#print("Option Keywords: " + str(option_keywords))

				for keyword in option_keywords:
					#print("Keyword: " + keyword)
					
					#print("Plain SKU: " + dashless_sku)
					if re.search(keyword,dashless_sku):
						final_opt_value = option_value
						final_opt_values.append(final_opt_value)

						final_opt_names.append(option_name)

						final_opt_string += option_name + "," + final_opt_value + ","
						break
					
					#put precedent on color code over sku or title bc specifically for opt gen
					# if no codes found in sku, check other fields for this item such as color field
					if re.search(keyword,color):
						final_opt_value = option_value
						final_opt_values.append(final_opt_value)

						final_opt_names.append(option_name)

						final_opt_string += option_name + "," + final_opt_value + ","
						break
						
					if re.search(keyword,material):
						final_opt_value = option_value
						final_opt_values.append(final_opt_value)

						final_opt_names.append(option_name)

						final_opt_string += option_name + "," + final_opt_value + ","
						break

					# if no codes found in sku, check other fields for this item such as title field
					if re.search(keyword,title):
						final_opt_value = option_value
						final_opt_values.append(final_opt_value)

						final_opt_names.append(option_name)

						final_opt_string += option_name + "," + final_opt_value + ","
						break

				if final_opt_value != '':
					#print("Final Option Name: " + option_name)
					#print("Final Option Value: " + final_opt_value)
					#print("Option String: " + final_opt_string + "\n")
					break

			#print("======Checked for Option Name: " + option_name + "\n")\

		#print("===Generated Options for SKU: " + sku + "\n")
	else:
		print("Warning: No details for this item!")

	final_opt_data = [ final_opt_names, final_opt_values ]		

	#print("=== Generated Options ===\n")

	return final_opt_data

def generate_all_options(all_details, init_all_details, all_handles=[]):
	all_options = []

	for item_idx in range(len(all_details)):
		item_details = all_details[item_idx]
		init_item_details = init_all_details[item_idx]
		item_handle = ''
		if len(all_handles) > 0:
			item_handle = all_handles[item_idx]

		options = generate_options(item_details, init_item_details, item_handle) # we need init details to detect measurement type
		
		# generate options string
		option_names = options[0]
		option_values = options[1]
		#print("Options: " + str(options))
		option_string = ''
		for opt_idx in range(len(option_names)):
			option_name = option_names[opt_idx]
			option_value = option_values[opt_idx]
			if opt_idx == 0:
				option_string += option_name + "," + option_value
			else:
				option_string += "," + option_name + "," + option_value
				
		all_options.append(option_string) # return as list of opt string

	return all_options

def generate_glossary_fmla(product):
	#print("\n=== Generate Glossary Formula ===\n")
	
	glossary_fmla = ""
	
	# gather terms in options to be defined like mattress opts like innerspring
	include_terms = []
	for variant in product:
		opt_data = generate_options(variant,variant) # [['Size','Mattress'],['Full','Innerspring']]
		#print("opt_data: " + str(opt_data))
		opt_vals = opt_data[1]
		#print("opt_vals: " + str(opt_vals))
		
		all_terms = ['Inner Spring','Dreamquest','Enso Memory Foam','Air Coil']
		
		for term in all_terms:
			if term in opt_vals and term not in include_terms:
				include_terms.append(term)
				
	#print("include_terms: " + str(include_terms) + "\n")
	
	glossary = reader.read_glossary("JF")
	#print("glossary: " + str(glossary) + "\n")
	
	# get def for each desired term
	glossary_defs = ''
	for term_idx in range(len(include_terms)):
		include_term = include_terms[term_idx]
		term_def = ''
		for term, definition in glossary.items():
			#print("term: " + term)
			#print("definition: " + definition)
			if include_term == term:
				term_def = definition
				break
		#print("term_def: " + term_def)
		if term_idx == 0:		
			glossary_defs += term_def
		else:		
			glossary_defs += "\",CHAR(10),\"" + term_def
		#print("glossary_defs: " + glossary_defs)
	
	glossary_fmla = "\"" + glossary_defs + "\""
	#print("glossary_fmla: " + glossary_fmla)

	
	#print("\n=== Generated Glossary Formula ===\n")

	return glossary_fmla

# what form is the product param? when is it set? Catalog Details [sku,...,barcode]
def generate_description(product, init_product,collection=[]):
	descrip_instances = []

	intro_fmla = generate_intro_fmla(product)

	colors_fmla = generate_colors_fmla(product,init_product)

	materials_fmla = generate_materials_fmla(product,init_product)

	finishes_fmla = generate_finishes_fmla(product)

	dimensions_fmla = generate_dimensions_fmla(product,init_product,collection)

	features_fmla = generate_features_fmla(product)

	#arrival_fmla = generate_arrival_fmla(product) # arrival time, such as Arrives: 3-4 weeks from Date of Purchase (eventually update dynamically based on date of purchase)

	# glossary to define terms like innerspring for mattress options
	glossary_fmla = generate_glossary_fmla(product)

	descrip_fmla = "==CONCATENATE("
	# add info if not blank
	if intro_fmla != '\"\"':
	 	descrip_fmla += intro_fmla
	if colors_fmla != '\"\"':
	 	descrip_fmla += ",CHAR(10),CHAR(10)," + colors_fmla
	if materials_fmla != '\"\"':
	 	descrip_fmla += ",CHAR(10),CHAR(10)," + materials_fmla
	if finishes_fmla != '\"\"':
	 	descrip_fmla += ",CHAR(10),CHAR(10)," + finishes_fmla
	if dimensions_fmla != '\"\"':
	 	descrip_fmla += ",CHAR(10),CHAR(10)," + dimensions_fmla
	if features_fmla != '\"\"':
	 	descrip_fmla += ",CHAR(10),CHAR(10)," + features_fmla
	# add glossary if not blank
	if glossary_fmla != '\"\"':
	 	descrip_fmla += ",CHAR(10),CHAR(10)," + glossary_fmla + ")"
	else:
	 	descrip_fmla += ")"
	 	
	if descrip_fmla == "==CONCATENATE()":
		descrip_fmla = ''

	# all variants of the product get the same description
	# the variants must be ordered by options, based on knowledge of desired option order and available options
	for variant in product:
		descrip_instances.append(descrip_fmla)

	return descrip_instances

def generate_all_descriptions(all_details, init_all_details):
	all_descriptions = []

	init_products = isolate_products(init_all_details)
	products = isolate_products(all_details)

	for product_idx in range(len(products)):
		product = products[product_idx]
		init_product = init_products[product_idx]
		descrip_instances = generate_description(product, init_product, all_details) # if it's a set you will need to know about other prods in coll. todo: isolate colls instead of going thru all given item details
		for descrip_instance in descrip_instances:
			all_descriptions.append(descrip_instance)

	return all_descriptions
	
def generate_item_name_from_export(variant):
	final_name = ''
	
	title_idx = 1
	opt1_idx = 3
	opt2_idx = 5
	opt3_idx = 7

	# look at item handle to determine title, and other details to determine options
	if len(variant) > 0:

		product_title = variant[title_idx]
		#print("Product Title: " + product_title)
		final_name += product_title

		opt1 = variant[opt1_idx].strip()
		if opt1 != "":
			final_name += "/" + opt1 
			opt2 = variant[opt2_idx].strip()
			if opt2 != "":
				final_name += "/" + opt2 
				
				opt3 = variant[opt3_idx].strip()
				if opt3 != "":
					final_name += "/" + opt3
	else:
		print("Warning: No info for this item!")

	#print("Final Name: " + final_name + "\n")
	return final_name

# needs init_item_details to gen opts
def generate_item_name(item_details, init_item_details, product_title='', product_options=''):

	#print("\n=== Generate Item Name ===\n")
	#print("product_title: \"" + product_title + "\"")
	#print("product_options: \"" + product_options + "\"")
	
	final_name = ''

	# look at item handle to determine title, and other details to determine options
	if len(item_details) > 0:

		if product_title == '':
			product_title = generate_title_from_details(item_details)
		#print("Product Title: " + product_title)
		final_name += product_title

		option_values = []
		option_data = []
		
		# eg product_options = 'optname,optval'
		if len(product_options) > 0:
			option_data = converter.convert_prod_opt_string_to_data(product_options)
		else:
			option_data = generate_options(item_details, init_item_details) # eg [[opt1name,opt2name],[opt1val,opt2val]]
			
		if len(option_data) > 0:
			option_values = option_data[1]
			#print("Option Values: " + str(option_values))
			
		total_num_opts = len(option_values)
		num_opts = total_num_opts
		if total_num_opts > 3:
			num_opts = 3 

		for value_idx in range(num_opts):
			value = option_values[value_idx]
			if value_idx == 0:
				final_name += "/" + value
			else:
				final_name += " / " + value

	else:
		print("Warning: No details for this item!")

	#print("Final Name: " + final_name + "\n")
	return final_name

def generate_all_item_names(all_details, init_all_details, product_titles=[], all_product_options=[]):
	all_item_names = []

	for item_idx in range(len(all_details)):
		item_details = all_details[item_idx]
		init_item_details = init_all_details[item_idx]
		
		product_title = ''
		if len(product_titles) > 0:
			product_title = product_titles[item_idx]
			
		product_options = ''
		if len(all_product_options) > 0:
			product_options = all_product_options[item_idx]

		item_name = generate_item_name(item_details, init_item_details, product_title, product_options)
		all_item_names.append(item_name)

	return all_item_names

def generate_collection_type(item_details):
	final_collection_type = ''

	output = "collection type"
	all_keywords = reader.read_keywords(output)

	# look at item handle to determine type
	if len(item_details) > 0:
		product_type = generate_product_type(item_details)

		for type, type_keywords in all_keywords.items():
			#print("Type: " + type)
			#print("Type Keywords: " + str(type_keywords))
			for keyword in type_keywords:
				#print("Keyword: " + keyword)
				if re.search(keyword,product_type):
					final_collection_type = type
					break

			if final_collection_type != '':
				break
	else:
		print("Warning: No details for this item!")

	#print(product_type + ", " + final_collection_type)
	return final_collection_type

def generate_all_collection_types(all_details):
	all_collection_types = []

	for item_details in all_details:
		collection_type = generate_collection_type(item_details)
		all_collection_types.append(collection_type)

	return all_collection_types

# extract options from product string and format as string
def generate_option_string(raw_product_string):

	option_string = ''

	max_options =  3

	raw_product_data = raw_product_string.split(';')

	init_opt_idx = import_opt1_name_idx
	for opt_idx in range(max_options):
		opt_name_idx = init_opt_idx + opt_idx * 2
		opt_val_idx = init_opt_idx + 1 + opt_idx * 2
		opt_name = raw_product_data[opt_name_idx]
		opt_val = raw_product_data[opt_val_idx]

		if opt_idx == 0:
			option_string += opt_name + ";" + opt_val
		else:
			option_string += ";" + opt_name + ";" + opt_val

	return option_string

# isolate unique variants by comparing option lists such as "["Size","10'x10'"]"
# def isolate_unique_variants(raw_product_vrnt_strings):
# 	print("\n=== Isolate Unique Variants ===\n")
#
# 	option_string = generate_option_string(raw_product_string)
#
#
#
# 	print("\n=== Isolated Unique Variants ===\n")

def get_unique_vrnt_idx(question_variant, sorted_product):
	unique_vrnt_idx = 0

	for vrnt_idx in range(len(sorted_product)):
		variant = sorted_product[vrnt_idx]
		if variant == question_variant:
			unique_vrnt_idx = vrnt_idx
			break

	#print("Unique Variant Index: " + str(unique_vrnt_idx))

	return unique_vrnt_idx

# def count_max_product_options(sorted_product, import_type):
# 	max_product_options = num_product_options = 0
#
# 	all_num_product_opts = []
#
# 	for vrnt_idx in range(len(sorted_product)):
# 		num_product_options = 0
#
# 		variant = sorted_product[vrnt_idx]
# 		vrnt_data = variant.split(';')
#
# 		relevant_vrnt_data = vrnt_data[6:12]
# 		if import_type == 'zoho':
# 			item_name = vrnt_data[item_name_idx]
#
# 			relevant_vrnt_data = item_name.split("/")
# 			relevant_vrnt_data = relevant_vrnt_data[1:]
#
# 		print("Variant Data: " + str(relevant_vrnt_data))
#
# 		for opt_idx in range(len(relevant_vrnt_data)):
# 			opt_name_or_value = relevant_vrnt_data[opt_idx]
# 			if import_type == 'zoho':
# 				if opt_name_or_value != '':
# 					num_product_options += 1
# 			else:
# 				# if even number idx then check if empty
# 				if opt_idx % 2 == 0:
# 					if opt_name_or_value != '':
# 						num_product_options += 1
#
# 		all_num_product_opts.append(num_product_options)
#
# 	max_product_options = max(all_num_product_opts)
#
# 	print("Max Product Options: " + str(max_product_options))
# 	return max_product_options
#
#
# def count_variant_options(question_variant, import_type):
# 	num_vrnt_opts = 0
#
# 	vrnt_data = question_variant.split(';')
#
# 	#print("Variant Data: " + str(vrnt_data[6:12]))
# 	#print("Question Variant Data: " + str(q_vrnt_data[6:12]))
# 	#print()
#
# 	relevant_vrnt_data = vrnt_data[6:12]
# 	if import_type == 'zoho':
# 		item_name = vrnt_data[item_name_idx]
#
# 		relevant_vrnt_data = item_name.split("/")
# 		relevant_vrnt_data = relevant_vrnt_data[1:]
#
# 	for opt_idx in range(len(relevant_vrnt_data)):
# 		opt_name_or_value = relevant_vrnt_data[opt_idx]
# 		if import_type == 'zoho':
# 			if opt_name_or_value != '':
# 				num_vrnt_opts += 1
# 		else:
# 			# if even number idx then check if empty
# 			if opt_idx % 2 == 0:
# 				if opt_name_or_value != '':
# 					num_vrnt_opts += 1
#
# 	print("Num Variant Options: " + str(num_vrnt_opts))
# 	return num_vrnt_opts

def determine_unique_variant(question_variant, sorted_product, import_type):

	unique_vrnt = True

	# max_product_options = count_max_product_options(sorted_product, import_type)
	# num_variant_options = count_variant_options(question_variant, import_type)
	# print()

	#print("Question Variant: " + question_variant)

	q_vrnt_data = question_variant.split(';')
	#print("Question Variant Data: " + str(q_vrnt_data))
	#print()

	unique_vrnt_idx = get_unique_vrnt_idx(question_variant, sorted_product)

	# we know that sku might actually be different (idx=0) but if the rest of the line is the same then it is a duplicate variant
	# really we just need the option data but the whole string should be the same
	for vrnt_idx in range(len(sorted_product)):
		variant = sorted_product[vrnt_idx]
		vrnt_data = variant.split(';')

		#print("Variant Data: " + str(vrnt_data[6:12]))
		#print("Question Variant Data: " + str(q_vrnt_data[6:12]))
		#print()

		relevant_vrnt_data = vrnt_data[6:12] # option fields in product import table
		relevant_q_vrnt_data = q_vrnt_data[6:12] # option fields in product import table
		if import_type == 'zoho':
			item_name = vrnt_data[item_name_idx] # item name has option fields in product import table, so get opt fields from itm nme
			q_item_name = q_vrnt_data[item_name_idx]

			relevant_vrnt_data = item_name.split("/")
			relevant_q_vrnt_data = q_item_name.split("/")

		if relevant_vrnt_data == relevant_q_vrnt_data: # remember last item is not included in range
			#print("Variant Data = Question Variant Data so could be duplicate variant.")
			if vrnt_idx != unique_vrnt_idx and vrnt_idx < unique_vrnt_idx:
				#print("New Variant matches an Earlier Variant (" + str(vrnt_idx) + " != " + str(unique_vrnt_idx) + "), so duplicate variant!")
				unique_vrnt = False
				break

	#print("Unique Variant? " + str(unique_vrnt) + "\n")

	return unique_vrnt

def isolate_detail_field(all_details, field_title):

	#print("\n=== Isolate Detail Field: " + field_title + " ===")

	detail_field_values = []

	item_name_idx = 1 # zoho
	
	field_idx = 0
	if field_title == "title":
		field_idx = item_name_idx

	for item_idx in range(len(all_details)):
		item_details = all_details[item_idx]
		#print("Item Details: " + str(item_details))
		
		field_value = ''
		if field_title == "handle":
			field_value = generate_handle(item_details)
		elif field_title == "title": # zoho import where title is part of item name title/opt_value
			field_value = item_details[field_idx]
			field_data = field_value.split("/")
			field_value = field_data[0]
		else:
			field_value = item_details[field_idx]
		#print("Init Field Value: " + field_value)

		#print("Final Field Value: " + field_value)
		detail_field_values.append(field_value)

	#print("=== Isolated Detail Field: " + field_title + " ===\n")

	return detail_field_values
	
def isolate_import_field(all_import_data, field_title):

	#print("\n=== Isolate Import Field: " + field_title + " ===")

	detail_field_values = []

	handle_idx = 1 # shopify
	item_name_idx = 1 # zoho
	field_idx = 0
	if field_title == "handle":
		field_idx = gen_handle_idx
	elif field_title == "title":
		field_idx = item_name_idx

	for item_idx in range(len(all_import_data)):
		item_details = all_import_data[item_idx]
		#print("Item Details: " + str(item_details))
		
		field_value = item_details[field_idx]
		
		if field_title == "title": # zoho import where title is part of item name title/opt_value
			field_data = field_value.split("/")
			field_value = field_data[0]
		#print("Init Field Value: " + field_value)

		#print("Final Field Value: " + field_value)
		detail_field_values.append(field_value)

	#print("=== Isolated Detail Field: " + field_title + " ===\n")

	return detail_field_values

def isolate_product_from_details(all_details, start_idx, stop_idx):
	product_rows = []

	for variant_idx in range(start_idx, stop_idx):
		product_rows.append(all_details[variant_idx])

	return product_rows

def isolate_products(all_details):
	products = []

	field_title = "handle" # we know that all variants of the same product have the same handle

	# ensure same handles grouped together, sort by handle 
	#unsorted_handles = generate_all_handles(all_details)
	#sorted_details = sorter.sort_items_by_handle(unsorted_handles, all_details)

	#if handle given in details (old version) handles = np.array(isolate_detail_field(all_details, field_title))
	#sorted_handles = np.array(generate_all_handles(sorted_details)) # details already sorted by handle
	handles = np.array(isolate_detail_field(all_details, field_title))

	_, idx, cnt = np.unique(handles, return_index=True, return_counts=True)

	unique_handles = handles[np.sort(idx)]
	counts = cnt[np.argsort(idx)]
	indices = np.sort(idx)

	num_products = len(unique_handles)

	# isolate products and append to products array
	# handles must already be adjacent b/c uses start idx + count to determine stop idx
	for product_idx in range(num_products):
		product_start_idx = indices[product_idx]
		product_stop_idx = product_start_idx + counts[product_idx]

		product_rows = isolate_product_from_details(all_details, product_start_idx, product_stop_idx)
		products.append(product_rows)

		product_start_idx = product_stop_idx
		if product_start_idx > len(all_details) - 1:
			break;

	#print("Products: " + str(products) + "\n")
	return products

def isolate_product_strings(all_imports, import_type):
	products = []

	#print("All Imports: " + str(all_imports) + "\n")
	#print("Isolate Product Strings for Import Type: \"" + import_type + "\"")

	field_title = "handle" # we know that all variants of the same product have the same handle

	if import_type == "zoho":
		field_title = "title"

	all_import_data = []
	for variant_import in all_imports:
		import_data = []
		import_data = variant_import.split(";")
		all_import_data.append(import_data)

	handles = np.array(isolate_import_field(all_import_data, field_title))

	_, idx, cnt = np.unique(handles, return_index=True, return_counts=True)

	unique_handles = handles[np.sort(idx)]
	counts = cnt[np.argsort(idx)]
	indices = np.sort(idx)

	num_products = len(unique_handles)

	# isolate products and append to products array
	for product_idx in range(num_products):
		product_start_idx = indices[product_idx]
		product_stop_idx = product_start_idx + counts[product_idx]

		product_rows = isolate_product_from_details(all_imports, product_start_idx, product_stop_idx)
		products.append(product_rows)

		product_start_idx = product_stop_idx
		if product_start_idx > len(all_imports) - 1:
			break;

	#print("Products: " + str(products) + "\n")
	return products

def capitalize_sentences(intro):
	final_intro = ''

	intro_sentences = intro.split('.')
	#print("intro_sentences: " + str(intro_sentences))
	for sentence in intro_sentences:
		if sentence != '':
			#print("sentence: " + sentence)
			sentence = sentence.strip().capitalize() + '. '
			final_intro += sentence

	return final_intro

def generate_intro_fmla(product):
	intro_fmla = "\"\""

	for variant in product:

		intro = ''

		if len(variant) > 3:
			handle = generate_handle(variant) #variant[1].strip().lower()
			#print("Handle: " + handle)
			intro = variant[3].strip().lower()
			intro = re.sub(";","\.",intro)
			intro = re.sub("u\.s\.","US",intro)
			#print("Intro: " + intro)

		part_warning = '' # if image not perfectly clear what it is like with slipcover same image for sofa and slipcover
		type = generate_product_type(variant)
		if type == 'slipcovers':
			part_warning = 'Slipcover only! Make sure the photo and dimensions match your furniture.'

		if intro != '' and intro != 'n/a':
			intro = capitalize_sentences(intro)
			intro = re.sub('\"','\",CHAR(34),\"', intro) # if quotes found in intro, such as for dimensions, then fmla will incorrectly interpret that as closing string
			if part_warning != '':
				intro_fmla = "\"" + part_warning + "\",CHAR(10),CHAR(10),\"" + intro + "\""
			else:
				intro_fmla = "\"" + intro + "\""
		#print("Intro Formula: " + intro_fmla)

	return intro_fmla

def determine_given_colors(product_details):
	given_colors = True

	for variant_details in product_details:
		colors = ''

		if len(variant_details) > 4:
			colors = variant_details[4].strip()

			if colors == '' or colors.lower() == 'n/a':
				# no colors given
				given_colors = False

		else:
			given_colors = False

	return given_colors

def generate_colors_fmla(product, init_product):

	final_colors_fmla = "\"\"" # init fmla for this part of the description
	if determine_given_colors(product): # if we are NOT given colors we do not include colors in the description
		final_colors_fmla = "\"Colors: \"" # only if at least 1 of the variants has colors

		opt_name = 'Color' # choose what option we want to show
		standard_type = opt_name.lower() + "s" # standards are defined in the data/standards folder
		options = reader.read_standards(standard_type) # get dict of options
		color_options = options[opt_name]
		#print("Color Options: " + str(color_options))

		valid_opt = False

		for vrnt_idx in range(len(product)):
			variant = product[vrnt_idx]
			init_variant = init_product[vrnt_idx]

			valid_opt = False

			colors = ''

			if len(variant) > 4:
				handle = generate_handle(variant) #variant[1].strip().lower()
				#print("Handle: " + handle)
				colors = variant[4].strip().lower()
				#print("Colors: " + colors)

			colors_fmla = '\"\"' # init option fmla so if no value given it is empty quotes
			if colors != '' and colors != 'n/a':
				colors = re.sub('\"','\",CHAR(34),\"', colors) # if something like "red" brown is given as colors, then fmla will incorrectly interpret that as closing string
				colors_fmla = "\"" + colors + "\""
			#print("Colors Formula: " + colors_fmla)

			option_data = generate_options(variant, init_variant)

			if len(option_data) > 0:
				option_names = option_data[0]
				#print("Option Names: " + str(option_names))
				option_values = option_data[1]
				#print("Option Values: " + str(option_values))

			# get the value of the size option, if there is one
			opt_idx = 0
			for current_opt_name in option_names:
				if current_opt_name == opt_name:
					#print("Valid Opt: " + opt_name)
					valid_opt = True
					break
				opt_idx += 1


			#print("Opt Idx: " + str(opt_idx))
			if valid_opt:
				opt_value = option_values[opt_idx]
				#print("Option Value: " + opt_value)

				opt_fmla = colors_fmla

				color_options[opt_value] = opt_fmla

		#print("Populated Option Values: " + opt_name + ": " + str(color_options))

		# now we have populated all color values for this product
		# so create color fmla by looping through colors and printing those with valid values
		if valid_opt:
			#print("Colors: ")
			opt_idx = 0
			for color_name, color_value in color_options.items():
				if color_value != '':
					variant_color_fmla = color_value
					#print(variant_color_fmla)
					if opt_idx == 0:
						final_colors_fmla += "," + variant_color_fmla
					else:
						final_colors_fmla += ",\", or \"," + variant_color_fmla
					opt_idx += 1
			#print()
			final_colors_fmla += ",\". \""
		else:
			final_colors_fmla += "," + colors_fmla + ",\". \""

	#print("Colors Formula: " + final_colors_fmla + "\n")

	return final_colors_fmla

def determine_given_materials(product_details):
	given_materials = True

	for variant_details in product_details:
		materials = ''

		if len(variant_details) > 5:
			materials = variant_details[5].strip()

			if materials == '' or materials.lower() == 'n/a':
				# no colors given
				given_materials = False

		else:
			given_materials = False

	return given_materials

def generate_materials_fmla(product, init_product):

	final_materials_fmla = "\"\"" # init fmla for this part of the description
	if determine_given_materials(product): # if we are NOT given materials we do not include materials in the description
		final_materials_fmla = "\"Materials: \"" # only if at least 1 of the variants has materials

		# for now, only handle cases where all variants have same material
		variant1 = product[0]
		
		opt_name = 'Material' # choose what option we want to show
		standard_type = opt_name.lower() + "s" # standards are defined in the data/standards folder
		options = reader.read_standards(standard_type) # get dict of options
		material_options = options[opt_name]
		#print("Material Options: " + str(material_options))

		valid_opt = False
		
		for vrnt_idx in range(len(product)):
			variant = product[vrnt_idx]
			init_variant = init_product[vrnt_idx]

			valid_opt = False

			materials = ''

			if len(variant) > mat_idx:
				handle = generate_handle(variant1) #variant1[1].strip().lower()
				materials = variant[mat_idx].strip().lower()

			materials_fmla = '\"\"' # init option fmla so if no value given it is empty quotes
			if materials != '' and materials != 'n/a':
				# format materials string by correcting typos and replacing invalid characters
				materials = re.sub('\"','\",CHAR(34),\"', materials) # if something like "s" spring is given as material, then fmla will incorrectly interpret that as closing string

				materials_fmla = "\"" + materials + "\""
			#print("Materials Formula: " + colors_fmla)
			
			option_data = generate_options(variant, init_variant)

			if len(option_data) > 0:
				option_names = option_data[0]
				#print("Option Names: " + str(option_names))
				option_values = option_data[1]
				#print("Option Values: " + str(option_values))

			# get the value of the size option, if there is one
			opt_idx = 0
			for current_opt_name in option_names:
				if current_opt_name == opt_name:
					#print("Valid Opt: " + opt_name)
					valid_opt = True
					break
				opt_idx += 1


			#print("Opt Idx: " + str(opt_idx))
			if valid_opt:
				opt_value = option_values[opt_idx]
				#print("Option Value: " + opt_value)

				opt_fmla = materials_fmla

				material_options[opt_value] = opt_fmla
				
		#print("Populated Option Values: " + opt_name + ": " + str(material_options))

		# now we have populated all material values for this product
		# so create material fmla by looping through materials and printing those with valid values
		if valid_opt:
			#print("Materials: ")
			opt_idx = 0
			for material_name, material_value in material_options.items():
				if material_value != '':
					variant_material_fmla = material_value
					#print(variant_material_fmla)
					if opt_idx == 0:
						final_materials_fmla += "," + variant_material_fmla
					else:
						final_materials_fmla += ",\", or \"," + variant_material_fmla
					opt_idx += 1
			#print()
			final_materials_fmla += ",\". \""
		else:
			final_materials_fmla += "," + materials_fmla + ",\". \""
			
	#print("Materials Formula: " + final_materials_fmla + "\n")

	return final_materials_fmla

def determine_given_finishes(product_details):
	given_finishes = True

	for variant_details in product_details:
		finishes = ''

		if len(variant_details) > 6:
			finishes = variant_details[6].strip()

			if finishes == '' or finishes.lower() == 'n/a':
				# no finishes given
				given_finishes = False

		else:
			given_finishes = False

	return given_finishes

def generate_finishes_fmla(product):

	final_finishes_fmla = "\"\"" # init fmla for this part of the description
	if determine_given_finishes(product): # if we are NOT given finishes we do not include finishes in the description
		final_finishes_fmla = "\"Finishes: \"" # only if at least 1 of the variants has finishes

		# for now, only handle cases where all variants have same material
		variant1 = product[0]

		finishes = ''

		if len(variant1) > 6:
			handle = generate_handle(variant1) #variant1[1].strip().lower()
			finishes = variant1[6].strip().lower()

		finishes_fmla = '\"\"' # init option fmla so if no value given it is empty quotes
		if finishes != '' and finishes != 'n/a':
			finishes_fmla = "\"" + finishes + "\""

		final_finishes_fmla += "," + finishes_fmla + ",\". \""

	return final_finishes_fmla

def determine_given_dimensions(product_details):
	given_dims = True

	for variant_details in product_details:
		width = depth = height = ''

		if len(variant_details) > 7:
			width = variant_details[7].strip()

			if len(variant_details) > 8:
				depth = variant_details[8].strip()

				if len(variant_details) > 9:
					height = variant_details[9].strip()
		else:
			given_dims = False

		if width == '' or width.lower() == 'n/a':
			# no width given but maybe other dims given
			if depth == '' or depth.lower() == 'n/a':
				# no width or depth given but maybe height given
				if height == '' or height.lower() == 'n/a':
					given_dims = False

	return given_dims

# product input includes all variants of product
# def set_option_values(product, opt_name):
#
# 	standard_type = opt_name.lower() + "s"
#
# 	options = reader.read_standards(standard_type)
#
# 	valid_opt = False
#
# 	for variant in product:
#
# 		dim_fmla = ''
#
# 		if opt_name == 'Size':
# 			width = depth = height = ''
#
# 			if len(variant) > 5:
# 				handle = variant[1].strip()
# 				#print("Handle: " + handle)
# 				width = variant[5].strip()
#
# 				if len(item_details) > 6:
# 					depth = variant[6].strip()
#
# 					if len(item_details) > 7:
# 						height = variant[7].strip()
#
# 			width_fmla = depth_fmla = height_fmla = '\"\"'
# 			if width != '' and width != 'n/a':
# 				width_fmla = "\"" + width + "\",CHAR(34),\" W \""
# 			if depth != '' and depth != 'n/a':
# 				depth_fmla = "\"" + depth + "\",CHAR(34),\" D \""
# 			if height != '' and height != 'n/a':
# 				height_fmla = "\"" + height + "\",CHAR(34),\" H\""
#
# 			dim_fmla = width_fmla + ",\"x \"," + depth_fmla + ",\"x \"," + height_fmla + ",\". \""
#
# 		option_data = generate_options(variant)
# 		option_names = []
# 		option_values = []
# 		if len(option_data) > 0:
# 			option_names = option_data[0]
# 			option_values = option_data[1]
#
# 		# get the value of the size option, if there is one
# 		opt_idx = 0
# 		for current_opt_name in option_names:
# 			if current_opt_name == opt_name:
# 				valid_opt = True
# 				break
# 			opt_idx += 1
#
# 		if valid_opt:
# 			opt_value = option_values[opt_idx]
#
# 			opt_fmla = ''
# 			if opt_name == 'Size':
# 				opt_fmla = dim_fmla
# 			elif opt_name == 'Color':
# 				opt_fmla = color_fmla
#
# 			options[opt_value] = opt_fmla
#
# 	print(opt_name + ": " + str(options))
#
# 	return options

def determine_valid_option(option_values):
	valid_opt = False
	for dims in option_values.values():
		if dims != '':
			valid_opt = True

	return valid_opt

def get_variant_indices_by_size(product_details):
	#print("\n=== Get Variant Indices by Size ===\n")

	areas = []
	widths = []

	for variant in product_details:
		width = depth = height = ''

		if len(variant) > 7:
			handle = generate_handle(variant) #variant[1].strip().lower()
			#print("Handle: " + handle)
			width = variant[7].strip()

			if len(variant) > 8:
				depth = variant[8].strip()

				if len(variant) > 9:
					height = variant[9].strip()

		blank_width = blank_depth = blank_height = True
		if width != '' and width != 'n/a':
			blank_width = False
		if depth != '' and depth != 'n/a':
			blank_depth = False
		if height != '' and height != 'n/a':
			blank_height = False

		if not blank_width and not blank_depth:
			area = int(width) * int(depth)
			areas.append(area)
			widths.append(int(width))

		elif blank_width:
			widths.append(0) # we need 0 returned for width if empty

		elif blank_depth:
			# if width contains multiple ft symbols and depth is blank, take digits before first foot symbol as width and digits after as depth
			if re.search('\'\s*\d+\'',width):
				print("Format Notice: Measurement contains improper sequence of two separate feet measurements!")
				dims = width.split('\'')
				width = dims[0].rstrip('\'')
				depth = dims[2].rstrip('\'')
				widths.append(int(width))
		else:
			widths.append(0) # we need 0 returned for width if empty

	#print("Areas: " + str(areas))

	#areas_array = np.array(areas)
	#print("Widths: " + str(widths))
	num_widths = len(widths)
	#print("Num Widths: " + str(num_widths))

	widths_array = np.array(widths)

	sorted_indices = np.argsort(widths_array)
	sorted_indices = np.flip(sorted_indices)
	#print("Sorted Indices: " + str(sorted_indices))

	#print("\n=== Got Variant Indices by Size ===\n")

	return sorted_indices

def get_sorted_indices(product):

	#print("\n=== Get Variant Indices by Size ===\n")

	areas = []
	widths = []

	for variant in product:
		width = depth = height = ''

		if len(variant) > 7:
			handle = generate_handle(variant) #variant[1].strip().lower()
			#print("Handle: " + handle)
			width = variant[7].strip()

			if len(variant) > 8:
				depth = variant[8].strip()

				if len(variant) > 9:
					height = variant[9].strip()

		blank_width = blank_depth = blank_height = True
		if width != '' and width != 'n/a':
			blank_width = False
		if depth != '' and depth != 'n/a':
			blank_depth = False
		if height != '' and height != 'n/a':
			blank_height = False

		if not blank_width:
			meas_type = reader.determine_measurement_type(width, handle)

		if not blank_width and not blank_depth:
			width = reader.format_dimension(width, handle)
			depth = reader.format_dimension(depth, handle)
			width = re.sub("”","",width)
			depth = re.sub("”","",depth)
			
			width_float = float(width)
			width_int = round(width_float,0)

			depth_float = float(depth)
			depth_int = round(depth_float,0)

			area = width_int * depth_int
			areas.append(area)
			widths.append(width_int)

		elif blank_width:
			widths.append(0) # we need 0 returned for width if empty

		elif blank_depth:
			# if width contains multiple ft symbols and depth is blank, take digits before first foot symbol as width and digits after as depth
			if re.search('\'\s*\d+\'',width):
				dims = width.split('\'')
				width = dims[0].rstrip('\'')
				depth = dims[1].rstrip('\'')

				width = reader.format_dimension(width, handle)
				depth = reader.format_dimension(depth, handle)

				width_float = float(width)
				width_int = round(width_float,0)
				widths.append(width_int)

			if meas_type == 'round' or meas_type == 'square':
				width = reader.format_dimension(width, handle)
				depth = width

				width_float = float(width)
				width_int = round(width_float,0)

				depth_float = float(depth)
				depth_int = round(depth_float,0)

				area = width_int * depth_int
				areas.append(area)
				widths.append(width_int)
		else:
			widths.append(0) # we need 0 returned for width if empty

	#print("Areas: " + str(areas))
	#print("Widths: " + str(widths))

	#areas_array = np.array(areas)
	widths_array = np.array(widths)

	sorted_indices = np.argsort(widths_array)
	#sorted_indices = np.flip(sorted_indices) removed flip b/c thought better to have larger first for upsell but actually better to have smaller first b/c then customer is willing to explore options (otherwise high price is deterrent to even looking)
	#print("Sorted Indices: " + str(sorted_indices))

	#print("\n=== Got Variant Indices by Size ===\n")

	return sorted_indices

def sort_variants_by_size(product):

	variant1 = product[0]
	handle = generate_handle(variant1) #variant1[handle_idx]
	#print("=== Sort Variants by Size: " + handle + " ===")

	sorted_indices = get_sorted_indices(product) # numpy array
	num_widths = sorted_indices.size
	#print("Num Widths: " + str(num_widths))

	sorted_variants = product
	num_variants = len(product)
	#print("Num Variants: " + str(num_variants))

	# only sort variants if we have valid values for sorting
	if num_variants == num_widths:
		sorted_variants = []
		for idx in range(num_variants):
			#print("Index: " + str(idx))
			sorted_idx = sorted_indices[idx]
			#print("Sorted Index: " + str(sorted_idx))
			sorted_variant = product[sorted_idx]
			sorted_variants.append(sorted_variant)
	else:
		print("Warning for " + handle + ": Num Variants != Num Widths (" + str(num_variants) + " != " + str(num_widths) + ") while sorting variants!")

	#for variant in sorted_variants:
		#print("Sorted Variant: " + str(variant))
	return sorted_variants

def generate_product_dims_fmla(product, init_product, bundle=False):
	if determine_given_dimensions(product): # if we are NOT given dimensions we do not include dimensions in the description
		dimensions_fmla = '' # only if at least 1 of the variants has dimensions do we add dim fmla
		if bundle:
			#print("found bundle while generating prod dims fmla")
			dimensions_fmla += "\"Dimensions (in): \"" 
		else:
			dimensions_fmla = "\"Dimensions (in): \"" 
		#print("dimensions_fmla: " + dimensions_fmla)

		#sizes = set_option_values(product, 'Size')
		opt_name = 'Size' # choose what option we want to show
		standard_type = opt_name.lower() + "s" # standards are defined in the data/standards folder
		options = reader.read_standards(standard_type) # get dict of options
		size_options = options[opt_name]
		#print("Size Options: " + str(size_options))

		valid_opt = False

		# sort variants
		#print("Sort Init Variants")
		init_sorted_variants = sort_variants_by_size(init_product)
		#print("Sort Variants")
		sorted_variants = sort_variants_by_size(product)

		type = ''

		for vrnt_idx in range(len(sorted_variants)):
			variant = sorted_variants[vrnt_idx]
			init_variant = init_sorted_variants[vrnt_idx]

			vrnt_sku = variant[sku_idx]

			valid_opt = False

			type = generate_product_type(variant)
			
			width = depth = height = ''
				
			if len(variant) > 7:
				handle = generate_handle(variant) #variant[1].strip().lower()
				#print("Handle: " + handle)
				width = variant[7].strip()
				#print("width: " + width)

				if len(variant) > 8:
					depth = variant[8].strip()
					#print("depth: " + depth)

					if len(variant) > 9:
						height = variant[9].strip()
						#print("height: " + height)

			blank_width = blank_depth = blank_height = True
			if width != '' and width != 'n/a':
				blank_width = False
			if depth != '' and depth != 'n/a':
				blank_depth = False
			if height != '' and height != 'n/a':
				blank_height = False

			dim_fmla = ''
			width_fmla = depth_fmla = height_fmla = '\"\"' # init option fmla so if no value given it is empty quotes
			if not blank_width:
				width_fmla = "\"" + width + "\",CHAR(34),\" W\""
				dim_fmla = width_fmla
			if not blank_depth:
				depth_fmla = "\"" + depth + "\",CHAR(34),\" D\""
				if blank_width:
					dim_fmla = depth_fmla
				else:
					dim_fmla += ",\" x \"," + depth_fmla
			if not blank_height:
				height_fmla = "\"" + height + "\",CHAR(34),\" H\""
				if blank_width and blank_height:
					dim_fmla = height_fmla
				else:
					dim_fmla += ",\" x \"," + height_fmla

			dim_fmla += ",\". \"" # end with period and space

			option_data = generate_options(variant, init_variant)
			#print("Option Data: " + str(option_data))
			option_names = []
			option_values = []
			if len(option_data) > 0:
				option_names = option_data[0]
				option_values = option_data[1]

			# order option values from large to small, and correspond with dim_fmla

			# get the value of the size option, if there is one
			opt_idx = 0
			for current_opt_name in option_names:
				if current_opt_name == opt_name:
					valid_opt = True
					break
				opt_idx += 1

			if valid_opt:
				opt_value = option_values[opt_idx] # option value is dictionary key

				opt_fmla = dim_fmla

				# before assigning to options dict, could sort dims but could also sort after by checking if custom dims and storing in separate array to sort
				size_options[opt_value] = opt_fmla

		#print(opt_name + ": " + str(size_options))
		
		# similar to looping through sizes to list sizes in descrip, we loop through pieces in a set

		# now we have populated all size values for this product
		# so create dim fmla by looping through sizes and printing those with valid values
		if valid_opt:
			#print("Dimensions: ")

			# reorder custom dims from large to small

			for size, dims in size_options.items():
				if dims != '':
					#print("Dims: " + dims)
					if type == 'rugs':
						variant_dim_fmla = dims # do not add quote-comma to dims b/c already there
						#print("variant_dim_fmla: " + variant_dim_fmla)
						dimensions_fmla += ",CHAR(10)," + variant_dim_fmla
					else:
						size_fmla = "\"" + size + ": \","
						#print("size_fmla: " + size_fmla)
						variant_dim_fmla = size_fmla + dims
						#print(variant_dim_fmla)
						dimensions_fmla += ",CHAR(10)," + variant_dim_fmla

			#print()
		else:
			dimensions_fmla += "," + dim_fmla
		
	else:
		print("Warning: No dimensons given for product!")
			
	#print("dimensions_fmla: " + dimensions_fmla)
	return dimensions_fmla

def generate_dimensions_fmla(product, init_product, collection=[]):

	variant1 = init_product[0]
	handle = generate_handle(variant1) #variant[1].strip().lower()
	
	#print("\n===Generate Dimensions Formula for " + handle + "===\n")
	
	dimensions_fmla = "" # init fmla for this part of the description
	
	type = generate_product_type(variant1) # take any variant in the product to determine if it is a bundle
	
	# sets of items list sizes of included pieces
	if re.search("set",type): 
		#print("found set of type: " + type)
		
		if len(variant1) > width_idx: # sets have no dims so check
			width = variant1[width_idx].strip()
			#print("width: " + width)
			if width != 'n/a':
				print("Warning for " + handle + ": Check dimensions given for set!")
				
		
		# instead of looping through all variants in a product, 
		# we loop through all variants in collection bc we want to match skus of other products

		# we want to show the titles of pieces included even if no dims:
		# Table
		# Chair
		# Dimensions (in): 1" x 1" x 1". 
		# Bench
		if len(collection) > 0:
			# we need to get info check skus for each bundle vrnt
			added_skus = [] # store skus already included in descrip from previous vrnts
			added_handles = []
			for bundle_variant in product: # bc product is known bundle
				bundle_vrnt_sku = bundle_variant[sku_idx]
				#print("bundle_vrnt_sku: " + bundle_vrnt_sku)
				 
				for variant in collection:
					# see if sku of variant can be found in sku of bundle/set
					vrnt_sku = variant[sku_idx]
					#print("vrnt_sku: " + vrnt_sku)
					coll_name = variant[collection_idx]
					
					if re.search(vrnt_sku,bundle_vrnt_sku):
					
						vrnt_handle = generate_handle(variant)
						#print("vrnt_handle: " + vrnt_handle)
					
						if vrnt_handle not in added_handles: # do not repeat pieces already added bc it includes all vrnts of piece in first loop
							vrnt_title = generate_title(vrnt_handle) # get vrnt title
							#print("vrnt_title: " + vrnt_title)
							vrnt_title = writer.format_title_for_dim_fmla(vrnt_title,coll_name)
							if len(added_handles) != 0:
								dimensions_fmla += ",CHAR(10),CHAR(10),"
							dimensions_fmla += "\"" + vrnt_title + ": \",CHAR(10)," # add title to dim fmla
							
							# iso prod and list vrnt dims
							vrnt_prod = isolator.isolate_product_in_collection(vrnt_handle,collection)
							#print("vrnt_prod: " + str(vrnt_prod))
							dimensions_fmla += generate_product_dims_fmla(vrnt_prod, vrnt_prod, True)
							#print("dimensions_fmla: " + dimensions_fmla)
							
							added_skus.append(vrnt_sku)
							added_handles.append(vrnt_handle)
		else:
			print("Warning: No collection given to determine piece dimensions in bundle!")
	
	# for non-sets
	else:
		dimensions_fmla = generate_product_dims_fmla(product, init_product)
	
	#print("Dimensions Formula: " + dimensions_fmla + "\n")

	return dimensions_fmla

def generate_features_fmla(product):
	features_fmla = "\"\""

	for variant in product:

		features = vendor = ''

		if len(variant) > 11:
			handle = generate_handle(variant) #variant[1].strip().lower()
			#print("Handle: " + handle)
			features = variant[11].strip()
			#print("Features: " + features + "\n")
			if len(variant) > gen_vendor_idx:
				vendor = variant[gen_vendor_idx]
			#print("vendor: " + vendor)

		if features != '' and features != 'n/a':
			# need better way to check if there are no proper nouns that should stay capitalized, b/c too blunt to lowercase everything
			#features = features.lower()
			#features = capitalize_sentences(features).strip()
			#print("Capped Features: " + features)

			features = re.sub('\"','\",CHAR(34),\"', features) # if quotes found in features, such as for dimensions, then fmla will incorrectly interpret that as closing string
			features = re.sub(' •','. •', features) # add periods at end of lines
			features = re.sub(';',',', features) # add periods at end of lines
			if features[-1] != '.':
				#print("Last character: " + features[-1])
				features += '. '
				
			#print("Features: " + features + "\n")
			
			vendors_no_bullets = ['Ashley','Homelegance','Coaster'] # vendors with lists that do not have bullets so rely on capital letter as newline
				
			if vendor == 'Ashley' or vendor == 'Homelegance' or vendor == 'Coaster':
				# If capital letter, start new line before printing
				# Example sentence one Example sentence two
				# replace the character preceding capital letter w/ newline char
				# replace common caps, like abbrevs and proper nouns w/ lowercase 
				# temp make common caps lowercase so not newline
				common_caps = ["LED","UL","Listed","USB","Next-Gen Durapella","Easy View"]
				for capital in common_caps:
					#print("before: " + features)
					#print("capital: " + features)
					features = re.sub(capital,capital.lower(),features)
					
					#print("after: " + features)
				
				features = re.sub("CHAR","char",features) # temp make CHAR cmd lowercase so not newline
							
				features = re.sub('(?=[A-Z])','\",CHAR(10),\"• \",\"', features) # make newline before uppercase
				#print("subbed in bullets: " + features)
				
				# make common caps capital again
				for capital in common_caps:
					#print("before: " + features)
					#print("capital: " + features)
					features = re.sub(capital.lower(),capital,features)
				
				features = re.sub("char","CHAR",features) # make char capital again
			else:
				features = re.sub('•','\",CHAR(10),\"• \",\"', features) # bullet point indicates new line
				features = re.sub('    ','\",CHAR(10),\"• \",\"', features) # 4 spaces indicates new line
				features = re.sub('ï|Ï','\",CHAR(10),\"• \",\"', features) # ï character indicates new line (for Coaster)
				features = re.sub('--','\",CHAR(10),\"• \",\"', features) # double dash indicates new line
			
			# replace letter after bullet with capitalized version	
			#print("Features: " + features + "\n")
						
			features_fmla = "\"" + features + "\""
		#print("Features Formula: " + features_fmla)

	return features_fmla
	
def generate_all_item_descriptions(collections):

	all_item_descrips = []
	
	vendor = "Global"
	keyword = "description"
	data_type = "raw data"
	item_descrip_idx = determiner.determine_field_idx(vendor, keyword, data_type)

	prod_row_len = 20
	for coll in collections:
		product_descrip = ''
		for row_idx in range(len(coll)):
			
			row = coll[row_idx]
			# check if next row is less than 20 to determine if pkg row
			# if pkg then do not include barcode
			if determiner.determine_pkg_row(coll, row_idx):
				product_descrip = row[item_descrip_idx]
			elif not determiner.determine_pkg_row(coll, row_idx):
				item_descrip = row[item_descrip_idx]
				if item_descrip == '':
					item_descrip = 'n/a'
					
				# only include product descrip for certain product types where item descrip is too vague like "legs" but legs for what?
				if not re.search(product_descrip.lower(),item_descrip.lower()):
					#item_handle = generate_handle_from_x()
					#item_type = generate_product_type_from_handle(handle)
					# if it is a vague item descrip add product descrip for clarity
					if determiner.determine_vague_item_descrip(item_descrip):
						item_descrip = product_descrip + " " + item_descrip
					
				all_item_descrips.append(item_descrip)
				
	return all_item_descrips
	
def generate_tot_cart_wt(prod):
	vendor = "Global"
	data_type = "raw data"
	keyword = "carton weight"
	item_cart_wt_idx = determiner.determine_field_idx(vendor, keyword, data_type)
	
	tot_cart_wt = 0.0
	
	for row_idx in range(len(prod)):
		row = prod[row_idx]
	
		if not determiner.determine_pkg_row(prod, row_idx):
			item_cart_wt = row[item_cart_wt_idx]
			if item_cart_wt == '':
				tot_cart_wt = 0.0
				break
			
			tot_cart_wt += float(item_cart_wt)
			
	#print("tot_cart_wt: " + str(tot_cart_wt))
	return tot_cart_wt
	
def generate_all_item_asm_weights(collections):

	print("\n=== Generate All Item Assembled Weights ===\n")

	all_item_weights = []
	
	vendor = "Global"
	keyword = "assembled weight"
	data_type = "raw data"
	item_asm_wt_idx = determiner.determine_field_idx(vendor, keyword, data_type)
	
	keyword = "carton weight"
	item_cart_wt_idx = determiner.determine_field_idx(vendor, keyword, data_type)

	prod_row_len = 20
	for coll in collections:
	
		prods_in_coll = isolator.isolate_prods_in_coll(coll)
		
		product_wt = 'n/a' 
		for prod in prods_in_coll:
	
			for row_idx in range(len(prod)):
			
				row = prod[row_idx]
				
				coll_sku_idx = 0
				coll_sku = row[coll_sku_idx]
				
				# check if next row is less than 20 to determine if pkg row
				# if pkg then do not include barcode
				if determiner.determine_pkg_row(prod, row_idx):
					product_wt = row[item_asm_wt_idx]
					#print("product_wt: " + product_wt)
				elif not determiner.determine_pkg_row(prod, row_idx):
					item_asm_wt = 'n/a'
					if len(row) > item_asm_wt_idx:
						item_asm_wt = row[item_asm_wt_idx]
						if item_asm_wt != '':
							item_wt = item_asm_wt
						else:
							item_cart_wt = 'n/a'
							if len(row) > item_cart_wt_idx:
								item_cart_wt = row[item_cart_wt_idx]
								if item_asm_wt == '' and item_cart_wt == '':
									if product_wt == '' or product_wt == 'n/a':
										item_wt = 'n/a' 
									else:
										item_wt = str(round(float(product_wt)/(len(prod)-1), 1)) # we actually can determine still if the product has an asm wt and we divide into equal segments
								if item_cart_wt != '':
									if product_wt != '' and product_wt != 'n/a':
										# we can use carton wt ratio to get item asm wt
										tot_cart_wt = generate_tot_cart_wt(prod)
										if tot_cart_wt == 0.0:
											item_wt = 'n/a' 
										else:
											wt_rate = float(item_cart_wt) / tot_cart_wt
											#print("wt_rate: " + str(wt_rate))
											item_wt = str(round(float(product_wt) * wt_rate, 1))
									else:
										item_wt = item_cart_wt
							else:
								print("Warning at row index " + str(row_idx) + ": " + coll_sku + " pkg row does not have cart wt!")
					else:
						print("Warning at row index " + str(row_idx) + ": " + coll_sku + " pkg row does not have asm wt!")
				
								
					#print("item_wt: " + item_wt)
					all_item_weights.append(item_wt)
				
	return all_item_weights
	
def generate_all_item_colors(collections):

	all_item_colors = []
	
	vendor = "Global"
	keyword = "color"
	data_type = "raw data"
	item_color_idx = determiner.determine_field_idx(vendor, keyword, data_type)

	prod_row_len = 20
	for coll in collections:
	
		prods_in_coll = isolator.isolate_prods_in_coll(coll)
		
		product_color = 'n/a'
		for prod in prods_in_coll:
	
			for row_idx in range(len(prod)):
			
				row = prod[row_idx]
				# check if next row is less than 20 to determine if pkg row
				# if pkg then do not include barcode
				if determiner.determine_pkg_row(prod, row_idx):
					product_color = row[item_color_idx]
					if product_color == '':
						product_color = 'n/a'
				elif not determiner.determine_pkg_row(prod, row_idx):
					item_color = product_color
					if len(row) > item_color_idx:
						item_color = row[item_color_idx]
						if item_color == '':
							item_color = product_color # the line might be long enough bc later vals so double check color not blank
					
					item_color = re.sub(";",",",item_color) # remove semicolons bc we use as data delimiter
					all_item_colors.append(item_color)
				
	return all_item_colors
	
def generate_all_item_materials(collections):

	all_item_materials = []
	
	vendor = "Global"
	keyword = "material"
	data_type = "raw data"
	item_material_idx = determiner.determine_field_idx(vendor, keyword, data_type)

	prod_row_len = 20
	for coll in collections:
	
		prods_in_coll = isolator.isolate_prods_in_coll(coll)
		product_material = 'n/a'
		for prod in prods_in_coll:
	
			for row_idx in range(len(prod)):
			
				row = prod[row_idx]
				# check if next row is less than 20 to determine if pkg row
				# if pkg then do not include barcode
				if determiner.determine_pkg_row(prod, row_idx):
					product_material = row[item_material_idx]
					if product_material == '':
						product_material = 'n/a'
					
				elif not determiner.determine_pkg_row(prod, row_idx):
					item_material = product_material
					if len(row) > item_material_idx:
						item_material = row[item_material_idx]
						if item_material == '':
							item_material = product_material # the line might be long enough bc later vals so double check color not blank
					
					item_material = re.sub(";",",",item_material) # remove semicolons bc we use as data delimiter
					all_item_materials.append(item_material)
				
	return all_item_materials
	
def generate_all_item_features(collections):

	all_item_features = []
	
	vendor = "Global"
	keyword = "Comments/Features"
	data_type = "raw data"
	item_feature_idx = determiner.determine_field_idx(vendor, keyword, data_type)

	prod_row_len = 20
	for coll in collections:
	
		prods_in_coll = isolator.isolate_prods_in_coll(coll)
		product_features = 'n/a'
		for prod in prods_in_coll:
	
			for row_idx in range(len(prod)):
			
				row = prod[row_idx]
				# check if next row is less than 20 to determine if pkg row
				# if pkg then do not include barcode
				if determiner.determine_pkg_row(prod, row_idx):
					product_features = row[item_feature_idx]
					if product_features == '':
						product_features = 'n/a'
				elif not determiner.determine_pkg_row(prod, row_idx):
					item_features = product_features
					if len(row) > item_feature_idx:
						item_features = row[item_feature_idx]
						if item_features == '':
							item_features = product_features # the line might be long enough bc later vals so double check color not blank
					
					item_features = re.sub(";",",",item_features) # remove semicolons bc we use as data delimiter
					all_item_features.append(item_features)
				
	return all_item_features
	
def generate_all_item_costs(collections):

	all_item_costs = []
	
	vendor = "Global"
	keyword = "base price"
	data_type = "raw data"
	item_cost_idx = determiner.determine_field_idx(vendor, keyword, data_type)

	prod_row_len = 20
	for coll in collections:
	
		prods_in_coll = isolator.isolate_prods_in_coll(coll)
		
		product_cost = item_cost = 'n/a'
		for prod in prods_in_coll:
	
			for row_idx in range(len(prod)):
			
				row = prod[row_idx]
				
				# check if next row is less than 20 to determine if pkg row
				# if pkg then do not include barcode
				if determiner.determine_pkg_row(prod, row_idx):
					product_cost = row[item_cost_idx]
					if product_cost == '':
						product_cost = 'n/a'
						print("Warning: no pkg cost!!!")
				elif not determiner.determine_pkg_row(prod, row_idx):
					if len(row) > item_cost_idx:
						item_cost = row[item_cost_idx]
						if item_cost == '':
							item_cost = product_cost # the line might be long enough bc later vals so double check color not blank
					elif product_cost != 'n/a':
						item_cost = str(round(float(product_cost)/(len(prod)-1), 2))
					
					all_item_costs.append(item_cost)
				
	return all_item_costs
	
def generate_all_item_widths(all_assembled_dims, all_cart_dims):

	all_widths = []
	
	for item_idx in range(len(all_assembled_dims)):
		
		assembled_dims = all_assembled_dims[item_idx]
		cart_dims = all_cart_dims[item_idx]
		
		if assembled_dims != 'n/a' and assembled_dims != '':
			if re.search("headboard",assembled_dims) and re.search("fotboard",assembled_dims):
				dims = cart_dims
			else:
				dims = assembled_dims
		else:
			dims = cart_dims
			
		width = '0'
		dim_data = re.split('\s*x\s*', dims.lower())
		width_idx = 0
		if len(width) > width_idx:
			width = dim_data[width_idx]
			if width == '':
				width = '0'
			else:
				width = re.sub('w|l|\"','',width.lower()) # isolate meas val from other identifying chars such as unit 
		
		all_widths.append(width)
	
	return all_widths
	
def generate_all_item_depths(all_assembled_dims, all_cart_dims):

	all_depths = []
	
	for item_idx in range(len(all_assembled_dims)):
		assembled_dims = all_assembled_dims[item_idx]
		cart_dims = all_cart_dims[item_idx]
		
		if assembled_dims != 'n/a' and assembled_dims != '':
			if re.search("headboard",assembled_dims) and re.search("fotboard",assembled_dims):
				dims = cart_dims
			else:
				dims = assembled_dims
		else:
			dims = cart_dims
			
		depth = '0'
		if re.search("x", dims.lower()):
			dim_data = re.split('\s*x\s*', dims.lower())
			#print("dim_data: " + str(dim_data))
			depth = dim_data[1]
			depth = re.sub('d|w|\"','',depth.lower()) # isolate meas val from other identifying chars such as unit 
		else:
			print("Warning no x found in volume dims! '" + dims.lower() + "'")
		
		all_depths.append(depth)
	
	return all_depths
	
def generate_all_item_heights(all_assembled_dims, all_cart_dims):

	all_hts = []
	
	for item_idx in range(len(all_assembled_dims)):
		assembled_dims = all_assembled_dims[item_idx]
		cart_dims = all_cart_dims[item_idx]
		
		if assembled_dims != 'n/a' and assembled_dims != '':
			if re.search("headboard",assembled_dims) and re.search("fotboard",assembled_dims):
				dims = cart_dims
			else:
				dims = assembled_dims
		else:
			dims = cart_dims
			
		ht = '0'
		if re.search("x", dims.lower()):
			dim_data = re.split('\s*x\s*', dims.lower())
			ht_idx = 2
			if len(dim_data) > ht_idx:
				ht = dim_data[ht_idx]
				ht = re.sub('h|\"','',ht.lower()) # isolate meas val from other identifying chars such as unit 
		else:
			print("Warning no x found in volume dims! '" + dims.lower() + "'")
		
		all_hts.append(ht)
	
	return all_hts
	
def generate_all_item_weights(all_assembled_wts, all_cart_wts):

	all_wts = []
	
	for item_idx in range(len(all_assembled_wts)):
		assembled_wt = all_assembled_wts[item_idx]
		cart_wt = all_cart_wts[item_idx]
		
		if assembled_wt != 'n/a':
			wt = assembled_wt
		else:
			wt = cart_wt
			
		all_wts.append(wt)
	
	return all_wts
	
def generate_price(item, coll):

	tot_cost = gen_tot_cost(coll)
	
	tot_price = generate_variant_price(tot_cost, item)
	
	part_cost = item[cost_idx]

	rate = part_cost / tot_cost

	price = tot_price * rate
	
	return price
	
def generate_all_prices(all_details):

	all_prices = []
	
	field = "coll_name"
	collections = isolator.isolate_groups(all_details,field) # given coll name 

	for coll in collections:
	
		for item in coll:
	
			price = generate_price(item, coll)
			
			all_prices.append(price)
	
	return all_prices
	
def generate_all_item_skus(vendor, all_data):
	
	all_item_skus = []
	
	if vendor == "Klaussner" or vendor == "Comfort Design":
		data_type = "raw data"
		keyword = "style #"
		design_code_idx = determiner.determine_field_idx(vendor, keyword, data_type)
		keyword = "code"
		piece_code_idx = determiner.determine_field_idx(vendor, keyword, data_type)
	
		for data in all_data:
			design_code = data[design_code_idx]
			piece_code = data[piece_code_idx]
		
			sku = design_code + " " + piece_code
		
			all_item_skus.append(sku)
	else:
		print("Cannot Generate Item SKUs for vendor other than Klaussner!") # bc usually sku comes in one field, except for klaussner
	
	return all_item_skus
	
def generate_all_materials(vendor, all_skus):

	all_materials = []
	
	if vendor == "Klaussner" or vendor == "Comfort Design":
		for sku in all_skus:
		
			sku_codes = sku.split(' ')
			design_code = sku_codes[0].lower() # style num
	
			material = ''
			if re.search("m\\b",design_code):
				material = "Cotton" #'Pattern Match'
			elif re.search("\\bc?lv",design_code):
				material = "Hopkins Leather"
			elif re.search("\\bcl",design_code):
				material = "Leather"
			elif re.search("\\bl",design_code):
				material = "Vintage Leather"
			#elif re.search("\\bl?o?v?\\d{2}[12]4\\d",design_code):
				#material = "Double Needle"
			elif re.search("(air|dream|enso)",design_code):
				material = "n/a"
			elif re.search("\d{3}-\d{3}",design_code):
				material = "Wood"
			else:
				material = "Cotton"
				
			all_materials.append(material)
	else:
		print("Warning: Cannot generate materials for vendor other than Klaussner!")
	
	return all_materials
	
def generate_all_finishes(vendor, all_skus):

	all_finishes = []
	
	if vendor == "Klaussner" or vendor == "Comfort Design":
		for sku in all_skus:
		
			sku_codes = sku.split(' ')
			design_code = sku_codes[0].lower() # style num
	
			finish = ''
			if re.search("m\\b",design_code):
				finish = 'Pattern Match'
			elif re.search("\\bl?o?v?\\d{2}[12]4\\d",design_code):
				finish = "Double Needle"
			elif re.search("(air|dream|enso)",design_code):
				finish = "n/a"
			elif re.search("\d{3}-\d{3}",design_code):
				finish = "n/a"
			else:
				finish = "Chopper"
				
			all_finishes.append(finish)
	else:
		print("Warning: Cannot generate materials for vendor other than Klaussner!")
	
	return all_finishes
	
def generate_mapped_item_name(part_sku,zoho_items):
	
	mapped_item_name = ''
	
	for item in zoho_items:
		item_data = re.split(";",item)
		
		item_sku = item_data[sku_idx]
		item_name = item_data[item_name_idx]
		
		if part_sku == item_sku:
			mapped_item_name = item_name
			break
	
	return mapped_item_name
	
# input eg D637-00 * 2
def generate_mapped_qty(part_sku_with_qty):
	
	qty = '1'
	
	part_sku_data = re.split("\s\*\s",part_sku_with_qty)
	if len(part_sku_data) == 2:
		qty = part_sku_data[1]
	
	return qty
	
def generate_template_suffix(vendor):
	correct_spelling = ["Nuevo","Homelegance","Craftmaster"]
	wrong_spelling = ["Nuveo","HomeElegance","CraftMaster"]
	
	vendor_spelling = vendor 
	for idx in range(len(correct_spelling)):
		spelling = correct_spelling[idx]
		if vendor == spelling:
			vendor_spelling = wrong_spelling[idx]
			break
			
	vendor_spelling = re.sub("\s","",vendor_spelling) # remove spaces
	
	template_suffix = vendor_spelling
	
	return template_suffix

# helper functions
def roundup(x):
	 return int(math.ceil(x / 100.0)) * 100

def rounddown(x):
	 return int(math.floor(x / 100.0)) * 100
	 
	 
	 
# ====== Principle Generator ======

principle_num_idx = 0
principle_content_idx = 1

def generate_principle_num(principle_num_id):

	print("\n=== Generate Principle Num for Principle " + principle_num_id + " ===\n")

	print("principle_num_id: " + principle_num_id) # 1-1-1

	principle_num = '' 
	
	if len(principle_num_id) > 0:
	
		principle_num = re.sub('-','.',principle_num_id) + '.' 
	
	print("principle_num: " + principle_num) # 1.1.1.
	
	return principle_num

def generate_principle_num_id(principle_num):

	print("principle_num: " + principle_num) # 1.1.1.

	principle_num_id = '' 
	
	if len(principle_num) > 0:
	
		principle_num_id = re.sub('\.','-',principle_num) # 1-1-1-
		principle_num_id = principle_num_id.rstrip('-')
	
	print("principle_num_id: " + principle_num_id) # 1-1-1
	
	return principle_num_id

def generate_principle_id(principle, chapter_num):

	print("\n=== Generate Principle ID ===")

	print("principle: " + principle)
	
	principle_id = ''
	
	principle_data = isolator.isolate_principle_num_and_content(principle)
	
	principle_num = principle_data[principle_num_idx]
	
	principle_num_id = generate_principle_num_id(principle_num) # replace dots with hyphens in principle num to match id format, and remove last dot
	
	if len(principle_num) > 0:
		principle_id = "gob-p" + str(chapter_num) + "-" + principle_num_id # example: gob-p1-1 is in the first chapter and is the first principle in that chapter
	
	print("principle_id: " + principle_id)
	
	return principle_id


def generate_all_principle_ids(desired_principles):

	print("\n=== Generate All Principle IDs ===")

	all_principle_ids = []
	
	desired_chapters = isolator.isolate_chapters(desired_principles)
	title = 'desired_chapters'
	writer.display_list_with_title(desired_chapters, title)

	for chapter_idx in range(len(desired_chapters)):
		chapter = desired_chapters[chapter_idx]
		chapter_num = chapter_idx + 1

		for principle in chapter:
	
			principle_id = generate_principle_id(principle, chapter_num)
	
			all_principle_ids.append(principle_id)
	
	return all_principle_ids
	
def italicize_keywords(principle_content,keywords):

	#print("\n=== Italicize Keywords ===\n")

	italicized_content = principle_content
	
	for key in keywords:
		#print("key: " + key)
		italicized_key = "<em>" + key + "</em>"
		if re.search(key,principle_content):
			#print("found key " + key + " in content: " + principle_content)
			italicized_content = re.sub(key,italicized_key,principle_content)
			#print("italicized_content: " + italicized_content)

	#print("italicized_content: " + italicized_content)
	return italicized_content

def generate_principle_content(principle_data, principle_content_idx, demo_principles, principle_id):

	#print("\n=== Generate Principle Content ===\n")

	principle_content = ''
	
	if len(principle_data) > principle_content_idx:
		principle_content = principle_data[principle_content_idx] #'Sun Zi said: The Game of Business is of vital importance to the Company. '

		# make link if demo principle
		if determiner.determine_demo_principle(demo_principles, principle_id):
			chap_num = re.sub('(gob-p|-\d.*$)','',principle_id)
			#print("chap_num: " + chap_num)
		
			principle_num_id = re.sub('gob-p\d+-','',principle_id)
			princ_num = generate_principle_num(principle_num_id)
			
			fig_num = demo_principles.index(chap_num + "-" + princ_num) + 1
			
			principle_content = "<a href=\"#gob-fig" + str(fig_num) + "\" title=\"Figure " + str(fig_num) + ": Updating the Art of War Principle " + chap_num + "-" + princ_num + "\">" + principle_content + "</a>"

		# italicize keywords
		keywords = ["Moral Influence","Nature","Terrain","Domain","Leadership","Method","Accessible","Constricted"]
		principle_content = italicize_keywords(principle_content,keywords)
		#print("principle_content: " + principle_content)

	return principle_content

# raw data is ch titles, overviews, and principles
# valid data is ch titles, sections, and principle data
def generate_all_ch_valid_data(all_ch_titles, all_ch_overviews, all_ch_principles):

	#print("\n=== Generate All Chapter Valid Data ===\n")
	#print("all_ch_titles: " + str(all_ch_titles))

	all_ch_valid_data = []
	
	# no. ch titles tells us no. chapters
	for ch_idx in range(len(all_ch_titles)):
		ch_title = all_ch_titles[ch_idx]
		ch_overview = all_ch_overviews[ch_idx]
		ch_principles = all_ch_principles[ch_idx]
	
		#ch1_raw_principles = ['1. Sun Tzu said: The art of war is of vital importance to the State.','2. Principle two.']
		#ch1_raw_principles = isolator.isolate_all_principle_data(ch1)
		#ch1_raw_principles = all_ch_principles[0]
		
		ch_principle_data = []
		for raw_princ in ch_principles:
			principle_data = isolator.isolate_principle_num_and_content(raw_princ)
			ch_principle_data.append(principle_data)
	
		# we have list of all principles in sections, and then all sections in chapters
	
		# create groups. group means that it should be on the same page if it can be. a group could extend multiple pages, and a section could (often) only have 1 group.
		#ch1_overview = all_ch_overviews[0]
		
		ch_num = ch_idx + 1
		
		all_ch_section_groups = reader.read_json('chapter section groups') # read json
		#print("all_ch_section_groups: " + str(all_ch_section_groups))
		
		ch_section_groups = all_ch_section_groups[str(ch_num)] # groups within section within chapter
		#print("ch_section_groups: " + str(ch_section_groups))
		
		# gather chapter sections
		chs1 = ch_overview # first section of ch is always overview
		ch_sections = [chs1]
		
		# section_groups are the groups within the section, in form "1":["1.","2.",...]
		for section_groups_dict in ch_section_groups.values():
		
			section_groups = []
			
			for group in section_groups_dict.values():
			
				section_groups.append(group)
			
			ch_sections.append(section_groups)
		#print("ch_sections: " + str(ch_sections))
		
		# ch1s2g1 = ['1.','2.','2.1.','2.2.','2.3.','2.3.1.','2.3.2.','2.3.3.','2.3.4.','2.4','2.4.1.','2.4.2.','2.4.3.','2.4.4.','2.4.5.','2.5.','2.5.1.','2.5.2.','2.5.3.','2.5.4.','3.']
# 		ch1s2 = [ch1s2g1]
# 		ch1s3g1 = ['4.','4.1.','4.2.','4.3.','4.4.','4.5.','4.6.','4.7.','5.']
# 		ch1s3 = [ch1s3g1]
# 		ch1s4g1 = ['6.','7.','8.']
# 		ch1s4g2 = ['9.','9.1.','9.1.1.','9.1.2.','9.1.3.','9.1.4.']
# 		ch1s4g3 = ['10.','10.1.','10.2.','10.3.','10.4.','11.','11.1.','11.2.','11.3.','11.4.']
# 		ch1s4g4 = ['12.','13.']
# 		ch1s4 = [ch1s4g1, ch1s4g2, ch1s4g3, ch1s4g4]
# 		ch1s5g1 = ['14.']
# 		ch1s5 = [ch1s5g1]
# 	
# 		ch_sections = [chs1,ch1s2,ch1s3,ch1s4,ch1s5]
	
		#all_ch_sections = [ch1_sections]
		#all_ch_sections = []
		#for ch_section in ch_sections:
			#all_ch_sections.append(ch_section)
	
		#ch1_title = all_ch_titles[0]
		
		ch_data = [ch_title, ch_sections, ch_principle_data] # generate ch data based on given  divisions of ch, sec, and grp

		all_ch_valid_data.append(ch_data)
	
	return all_ch_valid_data
	
def generate_runs(content, paragraph, keywords):

	#print("\n=== Generate Runs ===\n")
	
	#print("content: " + content)
	#print("keywords: " + str(keywords))

	split_content = writer.split_docx_keywords(content, keywords)

	# create run for each piece of content
	for content_idx in range(len(split_content)):
		content = split_content[content_idx]
		#print("content: " + content)
		final_content = content
		# if content_idx == 0: # if we want first line to be indented
# 			final_content = "\t" + content
		run = paragraph.add_run(final_content)
		run.italic = determiner.determine_keyword(content, keywords)
	
def generate_section_header(document, doc_section, book_title, ch_name, section_title='', ch_num=''):

	#print("\n=== Generate Section Header ===\n")
	
	#print("book_title: " + book_title)
	
	full_book_title = book_title
	if book_title == 'aw':
		full_book_title = "The Art of War"
	elif book_title == 'aw,gb':
		full_book_title = "The Art of War | Game of Business"
	elif re.search("gb",book_title):
		full_book_title = "Game of Business"

	doc_section.header_distance = Inches(0.25)

	doc_header = doc_section.header
	doc_header.is_linked_to_previous = False
	header_block = doc_header.paragraphs[0]
	header_block.style = document.styles['Header Footer']
	
	# add tab stop so text can be aligned left and right 
	tab_stops = header_block.paragraph_format.tab_stops
	#print("tab_stops: " + str(tab_stops))
	tab_stop = tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
	#tab_stop.position = Inches(6.5)
	#tab_stop.alignment = WD_TAB_ALIGNMENT.RIGHT
	
	book_title_run = header_block.add_run(full_book_title)
	book_title_run.italic = True
	
	ch_num_name_run = header_block.add_run()
	if ch_num != '':
		ch_num_name_string = ", " + ch_num + ". " + ch_name
	else:
		ch_num_name_string = ", " + ch_name
	ch_num_name_run.text = ch_num_name_string
	
	if section_title != '':
	
		if re.search(",",book_title):
			run = header_block.add_run("\n")
		else:
			run = header_block.add_run("\t")
	
		#keywords = ["Sun Zi\'s Art of War"]
		raw_keywords = reader.extract_data("gb-keywords","Supplements","tsv")
		keywords = []
		for key in raw_keywords:
			keywords.append(key[0])
		
		# if we have sect title determine if it has keywords and needs runs
		# if determiner.determine_has_keywords(section_title,keywords):
# 		
# 			split_content = writer.split_docx_keywords(section_title, keywords)
# 			
# 			# create run for each piece of content
# 			for content_idx in range(len(split_content)):
# 				content = split_content[content_idx]
# 				print("content: " + content)
# 				if content_idx == 0:
# 					content = "\t" + content
# 				ch_sect_run = header_block.add_run(content)
# 				ch_sect_run.italic = determiner.determine_keyword(content, keywords)
# 				
# 		else:
# 			ch_sect_run = header_block.add_run()
# 			
# 			ch_sect_string = "\t" + section_title
# 			
# 			ch_sect_run.text = ch_sect_string
			
		generate_runs(section_title, header_block, keywords)
		
def generate_aw_cover_page(document):
	
	cover_paragraph = document.add_paragraph()
	cover_paragraph.style = document.styles['GB Cover']
	
	the_run = cover_paragraph.add_run("The\n")
	the_run.font.size = Pt(48)
	
	aw_run = cover_paragraph.add_run("Art of War\n\n")
	
	by_run = cover_paragraph.add_run("by\n")
	by_run.font.size = Pt(20)
	author_run = cover_paragraph.add_run("Sun Zi")
	author_run.font.size = Pt(48)
	
	document.add_section(start_type=WD_SECTION.NEW_PAGE)	
	
def generate_gb_cover_page(document):
	
	cover_paragraph = document.add_paragraph()
	cover_paragraph.style = document.styles['GB Cover']
	
	gb_run = cover_paragraph.add_run("Game of Business\n\n")
	
	from_run = cover_paragraph.add_run("Transformed from\n")
	from_run.font.size = Pt(20)
	author_run = cover_paragraph.add_run("Sun Zi\'s\n")
	author_run.font.size = Pt(48)
	author_run = cover_paragraph.add_run("Art of War")
	author_run.font.size = Pt(48)
	
	document.add_section(start_type=WD_SECTION.NEW_PAGE)	
	
# def generate_aw_gb_cover_page(document):
# 	
# 	cover_paragraph = document.add_paragraph()
# 	cover_paragraph.style = document.styles['GB Cover']
# 	
# 	compare_run = cover_paragraph.add_run("Comparison of\n\n")
# 	compare_run.font.size = Pt(20)
# 	
# 	aw_run = cover_paragraph.add_run("the Art of War\n\n")
# 	aw_run.font.size = Pt(56)
# 	aw_run.italic = True
# 	
# 	and_run = cover_paragraph.add_run("and\n\n")
# 	and_run.font.size = Pt(20)
# 	
# 	gb_run = cover_paragraph.add_run("Game of Business\n\n")
# 	gb_run.font.size = Pt(48)
# 	gb_run.italic = True
# 	
# 	by_run = cover_paragraph.add_run("by\n")
# 	by_run.font.size = Pt(20)
# 	author_run = cover_paragraph.add_run("Game of Business")
# 	author_run.font.size = Pt(24)
# 	
# 	document.add_section(start_type=WD_SECTION.NEW_PAGE)	

def generate_book_cover_page(document, book_title):

	if book_title == 'aw':
		generate_aw_cover_page(document)
	elif re.search("gb",book_title):
		generate_gb_cover_page(document)
	# elif book_title == 'aw,gb':
# 		generate_aw_gb_cover_page(document)
	
# ToC same for all books bc takes data from headings
def generate_book_toc(document, book_title='aw'):

	toc_heading = document.add_paragraph("Table of Contents", style="GB TOC Heading")
	
	# insert toc: generate the single line that inserts the TOC field
	toc_paragraph = document.add_paragraph()
	run = toc_paragraph.add_run()
	fldChar = OxmlElement('w:fldChar') # creates a new element
	fldChar.set(qn('w:fldCharType'), 'begin') # sets attribute on element
	instrText = OxmlElement('w:instrText')
	instrText.set(qn('xml:space'), 'preserve') # sets attribute on element
	if book_title == 'aw':
		instrText.text = 'TOC \\o "1-3" \\h \\z \\u' # change 1-3 depending on heading levels you need
	elif re.search("gb",book_title):
		instrText.text = 'TOC \\o "1-1" \\h \\z \\u' # change 1-3 depending on heading levels you need
	
	fldChar2 = OxmlElement('w:fldChar')
	fldChar2.set(qn('w:fldCharType'), 'separate')
	fldChar3 = OxmlElement('w:t')
	fldChar3.text = "Right-click to update field."
	fldChar2.append(fldChar3)
	
	fldChar4 = OxmlElement('w:fldChar')
	fldChar4.set(qn('w:fldCharType'), 'end')
	
	r_element = run._r
	r_element.append(fldChar)
	r_element.append(instrText)
	r_element.append(fldChar2)
	r_element.append(fldChar4)
	p_element = toc_paragraph._p
	
def generate_book_intro(document, book_title):

	intro_section = document.add_section(start_type=WD_SECTION.NEW_PAGE)
	intro_section.header_distance = Inches(0.25)
	
	ch_name = "Introduction"
	intro_heading = document.add_heading(ch_name)
	intro_heading.style = document.styles["GB Heading 1"]
	intro_heading.paragraph_format.space_after = Pt(6)
	
	generate_section_header(document, intro_section, book_title, ch_name)
	
	intro_paragraphs = reader.extract_data("gb-intro","intro","tsv")
	#print("intro_paragraphs: " + str(intro_paragraphs))
	
	raw_keywords = reader.extract_data("gb-keywords","Supplements","tsv")
	keywords = []
	for key in raw_keywords:
		keywords.append(key[0])
	#keywords = ["Sun Zi\'s Art of War"] # needed to gen runs. depends on context
	
	for p in intro_paragraphs:
		# if paragraph ends with number, take it as reference and make it superscript
		p_string = p[0]
		p_string = p_string.strip()
		# sub special characters
		p_string = reader.sub_special_characters(p_string)
		ref_nums = ''
		p_without_ref_nums = p_string
		if re.search("\d$",p_string):
			p_parts = p_string.split(".")
			ref_nums = p_parts[len(p_parts)-1]
			#print("ref_nums: " + ref_nums)
			p_without_ref_nums = re.sub("\d+(,\d+)*$","",p_string)
			#print("p_without_ref_nums: " + p_without_ref_nums)
			
		#intro_paragraph = document.add_paragraph(p_without_ref_nums)
		intro_paragraph = document.add_paragraph("\t",style="GB Intro")
		
		generate_runs(p_without_ref_nums, intro_paragraph, keywords) # gen runs to italicize keywords
	
def generate_valid_appendix_table(raw_ap_table, book_title='aw'):

	print("\n=== Generate Valid Appendix Table ===\n")

	print("raw_ap_table: " + str(raw_ap_table))
	
	# orig format only had aw app table, and we expanded so input in format [aw, gb]
	# eg ["Generals","c2",c3","Leaders","c5","c6"],["c1","c2","c3","c4"]

	valid_ap_table = raw_ap_table
	
	num_total_cols = len(valid_ap_table[0])
	print("num_total_cols: " + str(num_total_cols))
	
	merge_text = ''
	
	if num_total_cols > 3: # split aw and gb
	
		gb_idx = 3 
	
		for row in valid_ap_table:
		
			print("input row: " + str(row))
		
			num_cols = len(row)
			print("num_cols: " + str(num_cols)) # it's important that the number of cols differs by 1 only bc this is a standard merge format
		
			# take the first element of the first row and other rows of that length
			if num_cols == num_total_cols:
				aw_merge_text = row[0]
				print("aw_merge_text: " + aw_merge_text)
				num_books = 2 # 2 bc 2 books included with input
				num_book_cols = int(num_cols / num_books)
				print("num_book_cols: " + str(num_book_cols))
				gb_idx = num_book_cols
				print("gb_idx: " + str(gb_idx))
				gb_merge_text = row[gb_idx]
				print("gb_merge_text: " + gb_merge_text)
			
			else:
				print("row before insert: " + str(row))
				row.insert(0,aw_merge_text)
				print("row after aw insert: " + str(row))
				#row.insert(gb_idx,gb_merge_text)
				row[gb_idx] = gb_merge_text # cell already exists with blank entry bc format of input
				print("row after gb insert: " + str(row))
				
			print("row: " + str(row))
			
		# split aw and gb	
		valid_ap_table = isolator.isolate_book_appendix_data(valid_ap_table, book_title)
	
	else:
		for row in valid_ap_table:
		
			num_cols = len(row)
			#print("num_cols: " + str(num_cols)) # it's important that the number of cols differs by 1 only bc this is a standard merge format
		
			# take the first element of the first row and other rows of that length
			if num_cols == num_total_cols:
				merge_text = row[0]
				#print("merge_text: " + merge_text)
			
			else:
		
				row.insert(0,merge_text)
			
	# with aw and gb format: ["Generals","c2",c3","Leaders","c5","c6"],["Generals","c2",c3","Leaders","c5","c6"]
	#print("valid_ap_table: " + str(valid_ap_table))
	
	return valid_ap_table
	
def generate_all_valid_appendix_tables(raw_ap_tables, book_title='aw'):

	valid_ap_tables = []

	for raw_ap_table in raw_ap_tables:
	
		appendix_table = generate_valid_appendix_table(raw_ap_table, book_title)
	
		valid_ap_tables.append(appendix_table)
		
	return valid_ap_tables
	
# table overview data in form Table A14.1-5: Strategies for example
def generate_table_nums(all_table_overview_data):

	#print("\n=== Generate Table Nums ===\n")

	table_nums = []
	
	for table_overview_data in all_table_overview_data:
		ap_table_id = table_overview_data.split(":")[0] # 
		if re.search("\.",ap_table_id): # multiple tables, eg Table A14.1-5
			# last num is no. tables. for now only single digit.
			table_num_range = ap_table_id.split("A")[1] # eg 14.1-5
			#print("table_num_range: " + table_num_range)
			overall_table_num = table_num_range.split(".")[0]
			#print("overall_table_num: " + overall_table_num)
			num_partial_tables = int(table_num_range[-1]) # eg 5
			for num_partial_table in range(num_partial_tables):
				table_num = overall_table_num + "." + str(num_partial_table+1)
				table_nums.append(table_num)
				
		else: # one table, eg Table A1
			table_num = ap_table_id.split("A")[1]
			table_nums.append(table_num)
	
	#print("table_nums: " + str(table_nums))
	return table_nums
	
def generate_table_captions(all_table_overview_data, table_type='appendix'):

	#print("\n=== Generate Table Captions ===\n")

	table_captions = []
	
	for table_overview_data_idx in range(len(all_table_overview_data)):
		table_overview_data = all_table_overview_data[table_overview_data_idx]
		
		# table overview groups partial tables into sets so we need to split each partial table into a separate caption
		if table_type == 'appendix':
			if re.search("\.",table_overview_data): # multiple tables, eg Table A14.1-5: Example title
				ap_table_id = table_overview_data.split(":")[0] # Table A14.1-5
				ap_table_title = table_overview_data.split(":")[1] # Example title
				# last num is no. tables. for now only single digit.
				table_num_range = ap_table_id.split("A")[1] # eg 14.1-5
				#print("table_num_range: " + table_num_range)
				overall_table_num = table_num_range.split(".")[0]
				#print("overall_table_num: " + overall_table_num)
				num_partial_tables = int(table_num_range[-1]) # eg 5
				for num_partial_table in range(num_partial_tables):
					table_caption = "Table A" + overall_table_num + "." + str(num_partial_table+1) + ":" + ap_table_title
					table_captions.append(table_caption)
				
			else: # one table, eg Table A1: Example title
				table_caption = table_overview_data
				table_captions.append(table_caption)
		elif table_type == 'conclusion':
			table_num = table_overview_data_idx + 1
			table_caption = "Table " + str(table_num) + ": " + table_overview_data[0]
			table_captions.append(table_caption)
	
	#print("table_captions: " + str(table_captions))
	
	return table_captions
	
# for appendix tables
def generate_comparison_table_captions(all_table_overview_data, book_title='aw'):

	#print("\n=== Generate Comparison Table Captions ===\n")

	table_captions = []
	
	for table_overview_data_idx in range(len(all_table_overview_data)):
		table_overview_data = all_table_overview_data[table_overview_data_idx]
		
		# table overview groups partial tables into sets so we need to split each partial table into a separate caption
		if re.search("\.",table_overview_data): # multiple tables, eg Table A14.1-5: Example title
			ap_table_id = table_overview_data.split(":")[0] # Table A14.1-5
			ap_table_title = table_overview_data.split(":")[1] # Example title
			# last num is no. tables. for now only single digit.
			table_num_range = ap_table_id.split("A")[1] # eg 14.1-5
			#print("table_num_range: " + table_num_range)
			overall_table_num = table_num_range.split(".")[0]
			#print("overall_table_num: " + overall_table_num)
			num_partial_tables = int(table_num_range[-1]) # eg 5
			for num_partial_table in range(num_partial_tables):
				table_caption = "Table (" + book_title.upper() + ", A" + overall_table_num + "." + str(num_partial_table+1) + "):" + ap_table_title
				table_captions.append(table_caption)
			
		else: # one table, eg Table A1: Example title
			ap_table_id = table_overview_data.split(":")[0] # Table A14
			table_num = ap_table_id.split("A")[1] # eg 14
			ap_table_title = table_overview_data.split(":")[1] # Example title
			table_caption = "Table (" + book_title.upper() + ", A" + table_num + "):" + ap_table_title
			table_captions.append(table_caption)
	
	#print("table_captions: " + str(table_captions))
	
	return table_captions
	
def generate_appendix_tables(document, book_title, all_ap_table_overview_data=[]):

	appendix_tables = []
	
	ap_table_nums = generate_table_nums(all_ap_table_overview_data) #['1','2','3','4','5','6','7','8','9','10','11','12','13']
	raw_ap_tables = reader.extract_appendix_tables(book_title, "tsv", ap_table_nums) # lines of document init format
	
	appendix_tables = generate_all_valid_appendix_tables(raw_ap_tables, book_title) # output is for a single book
	
	#table1 = appendix_tables[0]
	# display table with appendix data
	#ap_table1 = writer.display_appendix_table(document, table1)
	
	ap_table_captions = generate_table_captions(all_ap_table_overview_data)
	writer.display_all_appendix_tables(document, appendix_tables, ap_table_captions)
	
	# num_rows = len(appendix_tables[0])
# 	ap_table = document.add_table(rows=num_rows,cols=2)
# 	
# 	for row_idx in range(num_rows):
# 		table_overview = aw_table_overview_data[table_idx]
# 		print("table_overview: " + table_overview)
# 		table_overview_parts = table_overview.split(":")
# 		table_num = table_overview_parts[0] + ":"
# 		table_title = table_overview_parts[1].strip()
# 		num_cell = ap_tables_overview.cell(table_idx,0)
# 		title_cell = ap_tables_overview.cell(table_idx,1)
# 		num_cell.text = table_num
# 		title_cell.text = table_title
# 		num_cell.width = Inches(1.75)
# 		title_cell.width = Inches(4.75)
# 		
# 		num_paragraph = num_cell.paragraphs[0]
# 		title_paragraph = title_cell.paragraphs[0]
# 		num_paragraph.style = document.styles['Tables Overview']
# 		title_paragraph.style = document.styles['Tables Overview']
	
	#writer.display_appendix_tables(document, appendix_tables)
	
	# now we have valid appendix tables in list, so display them
	
	# for ap_table_data in appendix_tables: # ap_table_data=[[Your Actual Condition, Strategy], [Capable, Seem Incapable], [Able to advance, Seem unable to advance]]
# 	
# 		print("ap_table_data: " + str(ap_table_data))
# 		
# 		# add caption
# 		# same as used in list of tables
# 	
# 		num_rows = len(ap_table_data)
# 		print("num_rows: " + str(num_rows))
# 		num_cols = len(ap_table_data[0])
# 		print("num_cols: " + str(num_cols))
# 	
# 		ap_table = document.add_table(rows=num_rows, cols=num_cols)
# 		for row_idx in range(num_rows):
# 			row = ap_table_data[row_idx]
# 			for col_idx in range(num_cols):
# 				cur_cell = ap_table.cell(row_idx,col_idx)
		
def generate_comparison_appendix_tables(document, book_title, all_books_table_ov_data=[]):

	print("\n=== Generate Comparison Appendix Tables: " + book_title + " ===\n")

	appendix_tables = []
	
	all_book_titles = book_title.split(",")
	
	all_books_appendix_tables = []
	all_books_ap_table_captions = []
	
	for book_idx in range(len(all_book_titles)):
		all_ap_table_overview_data = all_books_table_ov_data[book_idx]
		title = all_book_titles[book_idx]
	
		ap_table_nums = generate_table_nums(all_ap_table_overview_data) #['1','2','3','4','5','6','7','8','9','10','11','12','13']
		raw_ap_tables = reader.extract_appendix_tables(ap_table_nums, "tsv") # lines of document init format. gets all tables for all books for ap_table_nums given
	
		appendix_tables = generate_all_valid_appendix_tables(raw_ap_tables, title) # output is for a single book
		all_books_appendix_tables.append(appendix_tables)
		ap_table_captions = generate_comparison_table_captions(all_ap_table_overview_data, title)	
		print("ap_table_captions: " + str(ap_table_captions))
		all_books_ap_table_captions.append(ap_table_captions)
		print("all_books_ap_table_captions: " + str(all_books_ap_table_captions))
		
	writer.display_all_comparison_appendix_tables(document, all_books_appendix_tables, all_books_ap_table_captions)
	
def generate_principles_appendix(document, book_title):

	print("\n=== Generate Principles Appendix ===\n")

	ap_section = document.add_section(start_type=WD_SECTION.NEW_PAGE)
	ap_section.header_distance = Inches(0.25)
	
	ch_name = "Appendix"
	
	ap_heading = document.add_heading(ch_name, level=1)
	ap_heading.style = document.styles['GB Heading 1']
	
	origin_section = ap_section
	
	if re.search("gb",book_title) or book_title.lower() == 'game of business':
	
		# add why, how, what
		# add why sections
		why_heading = document.add_heading("Why",level=2) # add why heading
		why_heading.style = document.styles['GB Heading 2']
	
		section_title = "Why Game of Business was Written"
		generate_section_header(document, ap_section, book_title, ch_name, section_title)
	
		why_gb_written_heading_start = "Why "
		why_gb_written_heading = document.add_heading(why_gb_written_heading_start,level=3)
		why_gb_written_heading.style = document.styles['GB Heading 3']
		gb_title_run = why_gb_written_heading.add_run("Game of Business")
		gb_title_run.italic = True
		why_gb_written_heading_end_run = why_gb_written_heading.add_run(" was Written")
	
		document.add_paragraph("Game of Business was written to—")
		document.add_paragraph("Demonstrate the profitable ways of conducting business, ",style="List Bullet")
		document.add_paragraph("Guide business decisions, and ",style="List Bullet")
		document.add_paragraph("Improve company leadership. ",style="List Bullet")
		document.add_paragraph("Game of Business presents principles and strategies that were discovered by potentially the most experienced, known strategist, Sun Zi. Consequently, it will acquaint you with the dangers of business as you are guided through the most fun and useful way to maximize the likelihood of your company’s success. ")
	
		why_apply_gb_section = document.add_section(start_type=WD_SECTION.NEW_PAGE)
		why_apply_gb_section.header_distance = Inches(0.25)
		section_title = "Why Game of Business' Principles are Worth Applying"
		generate_section_header(document, why_apply_gb_section, book_title, ch_name, section_title)
	
		why_apply_gb_heading_start = "Why "
		why_apply_gb_heading = document.add_heading(why_gb_written_heading_start,level=3)
		why_apply_gb_heading.style = document.styles['GB Heading 3']
		gb_title_run = why_apply_gb_heading.add_run("Game of Business")
		gb_title_run.italic = True
		why_apply_gb_heading_end_run = why_apply_gb_heading.add_run("\' Principles are Worth Applying")
	
		document.add_paragraph("Game of Business’ principles are derived from innumerable years of strategically progressing towards a balance of time, freedom, and energy. ")
		document.add_paragraph("Sun Zi had an exceptionally well-balanced life, because he was able to get the goods and services that he needed. At a time when violence was the primary way of getting needed goods and services, he became a master of military strategy. He then conveyed his successful strategies in Art of War, so his descendants could use its framework to master basic strategy for military. ")
		document.add_paragraph("Technological progress passed a threshold at which the non-violent exchange of goods and services, enabled through business, yields a greater payoff than the violent alternative. Accordingly, Game of Business transforms the framework of Sun Zi’s Art of War, so it can be used to master basic strategy for business. So, if your ultimate goal is to balance your time, freedom, and energy, applying the principles in Game of Business will help you maximize your success. ")
	
		# add how sections
		how_section = document.add_section(start_type=WD_SECTION.NEW_PAGE)
		how_section.header_distance = Inches(0.25)
	
		how_heading = document.add_heading("How",level=2) # add how heading
		how_heading.style = document.styles['GB Heading 2']
	
		section_title = "How Game of Business is Written"
		generate_section_header(document, how_section, book_title, ch_name, section_title)
	
		how_gb_written_heading = document.add_heading(section_title,level=3) # add how gb written heading
		how_gb_written_heading.style = document.styles['GB Heading 3']
	
		document.add_paragraph("Sun Zi\'s Art of War is transformed to Game of Business by the following process: ")
		document.add_paragraph("Compare Sun Zi\'s Art of War translations. ", style="List Number")
		document.add_paragraph("Organize the principles of Sun Zi\'s Art of War systematically. ", style="List Number")
		document.add_paragraph("Replace references to the military domain with references to their counterparts in the business domain. ", style="List Number")
	
		document.add_paragraph("The philosophy (i.e. world model) of Sun Zi\s Art of War has two main parts: ")
		document.add_paragraph("Knowledge base specific to the military domain; ",style="List Bullet")
		world_model_p2 = document.add_paragraph("General, fixed reasoning system that can be applied to decisions in all domains.",style="List Bullet")
		ref_run = world_model_p2.add_run("11")
		ref_run.font.superscript = True
	
		document.add_paragraph("Game of Business is written as if the Art of War’s author is speaking directly with you. In this way, it aims to optimize the reader’s ability to learn, apply, and extend its philosophy. ")
	
		document.add_page_break()
	
		source_translations_intro_p = document.add_paragraph("Game of Business logically transforms the principles of Sun Zi\'s Art of War, based on the following ten (10) commentated translations (Tables [1,2]): ")
		source_translations_intro_p.paragraph_format.space_after = Pt(6)
	
		data_type = "gb-source"
		input_type = "all sources"
		extension = "tsv"
		all_source_data = reader.extract_data(data_type, input_type, extension)
	
		num_rows = len(all_source_data)
		num_cols = 2 # 1 for list num and 1 for list item
		source_table = document.add_table(num_rows,num_cols)
		for source_idx in range(len(all_source_data)):
			source = all_source_data[source_idx]
			num = source_idx + 1
			num_cell = source_table.cell(source_idx,0) # number 
			info_cell = source_table.cell(source_idx,1) # info 
			num_cell.text = str(num) + "."
			#info_cell.text = source
			writer.format_cell_paragraph(info_cell.paragraphs[0], source[0])
			num_cell.width = Inches(0.35)
			info_cell.width = Inches(6)
		
			p = info_cell.paragraphs[0]
			p.paragraph_format.space_after = Pt(6)
	
		input_type = "primary sources"
		include_field_title = True
		primary_source_data = reader.extract_data(data_type, input_type, extension, include_field_title)
	
		caption = "Table 1: Differences between primary sources (by publication date)"
		writer.display_nested_table(document, primary_source_data, caption)
	
		input_type = "secondary sources"
		secondary_source_data = reader.extract_data(data_type, input_type, extension, include_field_title)
	
		caption = "Table 2: Differences between secondary sources (by publication date)"
		writer.display_nested_table(document, secondary_source_data, caption)
	
		how_apply_gb_section = document.add_section(start_type=WD_SECTION.NEW_PAGE)
		how_apply_gb_section.header_distance = Inches(0.25)
	
		section_title = "How to Apply Game of Business\' Principles"
		generate_section_header(document, how_apply_gb_section, book_title, ch_name, section_title)
	
		how_apply_gb_heading = document.add_heading(section_title,level=3) # add how to apply gb heading
		how_apply_gb_heading.style = document.styles['GB Heading 3']
	
		document.add_paragraph("Game of Business\' principles are divided into skill sets that can be used to guide your company’s business decisions (see Table 3), and practiced to improve its leadership. Each chapter covers one idea about business, so you can learn one skill at a time. The skills learned are most effective when applied together fluidly. ")
	
		data_type = 'gb-appendix'
		input_type = 'decide'
		decide_data = reader.extract_data(data_type, input_type, extension, include_field_title)
		caption = "Table 3: Making a Decision Without and With Game of Business"
		writer.display_nested_table(document, decide_data, caption)
	
		document.add_page_break()
	
		document.add_paragraph("A proposition is like a question in the form of a statement, so its validity is undecided. It can be assumed true or false to move forward in a decision, or it can be kept as a proposition and its validity can be assigned a degree of belief. While a proposition’s validity remains undecided, one’s degree of belief of its validity enables one to move forward in a decision.")
		document.add_paragraph("Game of Business gives measurable goals, and methods to evaluate results, so you can learn and improve your business strategy. For example, when a company fails and its leader dismissed, the cause will surely be found among the Five Dangerous Faults that may affect leaders: Recklessness, cowardice, a hasty temper, a delicacy of honor, and excessive attentiveness to associates (see chapter 8, principles ten through twelve [8-10, 8-12]).")
		document.add_paragraph("Sun Zi’s Art of War was derived through intuitive experimentation, with a practical attitude, because lives are directly at stake in war. Facts were learned by combining philosophical, deep thinking with a keen eye and critical analysis of observations. Game of Business’ power comes from its feedback loop of testing business strategy in real-time, and a desire to benefit humankind.")
	
		how_find_gb_section = document.add_section(start_type=WD_SECTION.NEW_PAGE)
		how_find_gb_section.header_distance = Inches(0.25)
	
		section_title = "How to Find a Principle in Game of Business"
		generate_section_header(document, how_find_gb_section, book_title, ch_name, section_title)
	
		how_find_gb_heading = document.add_heading(section_title,level=3) # add how to find principle in gb heading
		how_find_gb_heading.style = document.styles['GB Heading 3']
	
		document.add_paragraph("Game of Business\' principles can be reached via the table of contents, index, and page format. ")
	
		use_toc_heading = document.add_heading("Use the Table of Contents")
		use_toc_heading.style = document.styles['GB Heading 4']
		document.add_paragraph("At the beginning of Game of Business, this brief table of contents provides an overview. ")
	
		use_idx_heading = document.add_heading("Use the Index")
		use_idx_heading.style = document.styles['GB Heading 4']
		document.add_paragraph("At the end of Game of Business, this alphabetical list includes important topics and terms. ")
	
		use_format_heading = document.add_heading("Use the Page Format (Figure 1)")
		use_format_heading.style = document.styles['GB Heading 4']
		
		#ids = ['1.','2.','3.','3.1.','3.2.','4.','4.1.']
		#parts = ["Chapter number and title. ", "Header showing the topic being discussed on this page. ", "Writing containing a principle. The numeric code consists of the principle’s number in the chapter. ", "Extra space between principles shows that they are related to different topics within the chapter. The principles are spaced close together when part of the same subtopic, and noticeably farther apart when part of different subtopics. A noticeable space after a principle means that the next principle is part of a different subtopic. ", "The principles within a subtopic depend on each other to be understood. Dependent statements are placed close to each other, whereas there is additional space between independent statements. ", "Writing containing a sub-principle. The numeric code consists of the principle’s number in the chapter, and sub-principle numbers separated by points. ", "An indented principle means that it is part of a set within a principle. The principles are in a sequence contained by the overarching principle above it. Hierarchical statements are denoted by adding a second number to the principle’s code, separated by a decimal point (e.g. 4.7. refers to the seventh statement or question of principle 4). "]
		ids = ['1.','2.']
		parts = ["Chapter number and title. ", "Header showing the topic being discussed on this page. "]
		writer.display_table(document, ids, parts)
		ids = ['3.','3.1.','3.2.']
		parts = ["Writing containing a principle. The numeric code consists of the principle’s number in the chapter. ", "Extra space between principles shows that they are related to different topics within the chapter. The principles are spaced close together when part of the same subtopic, and noticeably farther apart when part of different subtopics. A noticeable space after a principle means that the next principle is part of a different subtopic. ", "The principles within a subtopic depend on each other to be understood. Dependent statements are placed close to each other, whereas there is additional space between independent statements. "]
		table_data = [ids, parts]
		writer.display_nested_table(document, table_data, '','ordered list')
		ids = ['4.','4.1.']
		parts = ["Writing containing a sub-principle. The numeric code consists of the principle’s number in the chapter, and sub-principle numbers separated by points. ", "An indented principle means that it is part of a set within a principle. The principles are in a sequence contained by the overarching principle above it. Hierarchical statements are denoted by adding a second number to the principle’s code, separated by a decimal point (e.g. 4.7. refers to the seventh statement or question of principle 4). "]
		table_data = [ids, parts]
		writer.display_nested_table(document, table_data, '','ordered list')
		# document.add_paragraph("Chapter number and title. ",style="List Number")
# 		document.add_paragraph("Header showing the topic being discussed on this page. ",style="List Number")
# 		document.add_paragraph("Writing containing a principle. The numeric code consists of the principle’s number in the chapter. ",style="List Number")
# 		document.add_paragraph("Writing containing a sub-principle. The numeric code consists of the principle’s number in the chapter, and sub-principle numbers separated by points. ",style="List Number")
	
		document.add_page_break()
	
		pic_p = document.add_paragraph()
		sample_page_pic = pic_p.add_run().add_picture("../data/images/GB-sample-page.png",width=Inches(6.5),height=Inches(7.4)) # 7.4 to maintain arbitrary aspect ratio
		fig_caption_p = document.add_paragraph("Figure 1: Sample page of Game of Business")
	
		# add what sections
		what_section = document.add_section(start_type=WD_SECTION.NEW_PAGE)
		what_section.header_distance = Inches(0.25)
		what_heading = document.add_heading("What",level=2) # add what heading
		what_heading.style = document.styles['GB Heading 2']
	
		section_title = "Subject of Game of Business"
		generate_section_header(document, what_section, book_title, ch_name, section_title)
	
		what_subject_gb_heading = document.add_heading(section_title,level=3) 
		what_subject_gb_heading.style = document.styles['GB Heading 3']
	
		document.add_paragraph("You will learn how to lead a company to success, by applying the transformed principles and strategies of Sun Zi\'s Art of War to business. ")
	
		what_limit_gb_section = document.add_section(start_type=WD_SECTION.CONTINUOUS)
		what_limit_gb_section.header_distance = Inches(0.25)
	
		section_title = "Limits of Game of Business"
		generate_section_header(document, what_limit_gb_section, book_title, ch_name, section_title)
	
		what_limit_gb_heading = document.add_heading(section_title,level=3) 
		what_limit_gb_heading.style = document.styles['GB Heading 3']
		
		document.add_paragraph("Sun Zi\'s Art of War derives conclusions by applying logical deduction to a combination of facts and assumptions: ")
		
		ids = ['1.','2.','3.']
		parts = ["The principles that could not be logically deduced as fact are qualified by a set of assumptions; ","Conclusions are then derived, based on assumptions; and ","Decisions justified, based on conclusions. "]
		writer.display_table(document, ids, parts)
		
		document.add_paragraph("Sun Zi\'s Art of War manages its assumptions by dynamically asserting and retracting them, as new information is learned. However, it does not specify the conditions nor situations that cause assumptions to be asserted and retracted. So, if two assumptions contradict each other, the contradiction cannot be resolved. ")
		document.add_paragraph("Consider Game of Business\' principle 9-4, which states, \"If you are careful of your people, and occupy stable positions, the company will be free from disorder of every kind; this will result in success.\" ")
		
		limit_p4 = document.add_paragraph("Principle 9-4 can be divided and simplified into two conditional statements: ")
		limit_p4.paragraph_format.keep_with_next = True
		ids = ['1.','2.']
		parts = ["If a company is careful and stable, it will be organized. ", "If a company is organized, it will succeed. "]
		writer.display_table(document, ids, parts)
		
		document.add_paragraph("Based on principle 9-4, if you assume that your company is organized, you can logically deduce the belief that your company will succeed. If it turns out that your company is careless and unstable, the organized assumption will be retracted, leading you to also retract the belief that your company will succeed. ")
		document.add_paragraph("Consider the following dilemma: If we say “companies succeed,” then deductive logic would be able to infer the expected conclusion when it discovers a company; however, it would fall into an inconsistency if it encounters a company that fails. In contrast, if we say, “If a company is organized, it will succeed,” deductive logic will not be able to reach the expected conclusion by discovering the company, as it would not know whether the company is organized or not. ")
		document.add_paragraph("When determining an ideal strategy, instead of using a method to resolve conflicting assumptions, Game of Business recommends assigning degrees of belief to relevant propositions. Degrees of belief are probabilities that can be revised up or down as new information is learned, to dynamically update the ideal strategy. In this way, degrees of belief alleviate the need for assumptions, or help decide which assumptions to make. ")
		limit_p8 = document.add_paragraph("For example, if you assume that a company is organized unless observed otherwise, you will be led to the tenuous belief that the company will succeed. Instead of making this assumption, you can assign a degree of belief to the company’s organization, and then use this to derive a corresponding degree of belief in the company’s success. You may initially believe that a company is organized with 99% certainty, and later lower your certainty to 50% after observing that the company took a careless and unstable action.")
		ref_run = limit_p8.add_run("11")
		ref_run.font.superscript = True
		document.add_paragraph("Game of Business provides guidance to help you act as a strategist, rather than a gambler. However, to resolve instances of circular reasoning, you must assign degrees of belief to the propositions that would otherwise have been made into assumptions. ")
		
		what_scope_gb_section = document.add_section(start_type=WD_SECTION.CONTINUOUS)
		what_scope_gb_section.header_distance = Inches(0.25)
	
		section_title = "Scope of Game of Business"
		generate_section_header(document, what_scope_gb_section, book_title, ch_name, section_title)
	
		what_scope_gb_heading = document.add_heading(section_title,level=3) 
		what_scope_gb_heading.style = document.styles['GB Heading 3']
		
		view_gb_heading = document.add_heading("View",level=4) 
		view_gb_heading.style = document.styles['GB Heading 4']
		
		document.add_paragraph("Viewing business as the logical way to exchange goods and services, Game of Business transforms the principles of Sun Zi\'s Art of War to business, to create mutually beneficial, team strategies for business. ")
		document.add_paragraph("After witnessing the devastating violence and starvation resulting from war, the Art of War\'s author, Sun Zi, appears to have been committed to achieving a mutually beneficial, thriving community. For example, principle 9-18 states, \"If instructions are consistently followed, it indicates the instructions are mutually beneficial for leaders and associates.\" ")
		document.add_paragraph("When faced with a situation in which you must choose one of several alternatives, Game of Business recommends learning all the possible alternatives and consequences, and ranking the set of possible consequences, from most to least desired. ")
		document.add_paragraph("Consider if you have to decide whether to outmaneuver or ally with an opponent. What will happen if you ally with them? Maybe over time your main product will become obsolete, and your company will no longer be sustainable. Then again, maybe profit will continue to increase. What happens if you outmaneuver them, to reach your goal before them? It may lead to great success, or the maneuvers will fail, causing harmful expense. Both actions have uncertainty over the resulting outcomes, so the best action must be carefully considered. ")
		view_p5 = document.add_paragraph("Your company will either succeed or fail, depending on your decision. However, there is uncertainty about which outcome will prevail, and the uncertainty is tied to the choice you make. As indicated by principle 9-18, success is more likely if you choose to ally with the opponent, while less likely if you choose to outmaneuver them. However, if you know that the opponent is malicious and untrustworthy, your best response may be to outmaneuver them. Similarly, if a malicious action is taken against you, your best response may be to remove the malevolent source.")
		ref_run = view_p5.add_run("13")
		ref_run.font.superscript = True
		
		outlook_gb_heading = document.add_heading("Outlook",level=4) 
		outlook_gb_heading.style = document.styles['GB Heading 4']
		
		document.add_paragraph("Game of Business\' principles are meant to be tested and refined. ")
		
		application_gb_heading = document.add_heading("Application",level=4) 
		application_gb_heading.style = document.styles['GB Heading 4']
		
		document.add_paragraph("Apply Game of Business to your business strategy, because it has a generalizable, fixed reasoning system that focuses on simplifying concepts to make practical systems. ")
		
		operation_gb_heading = document.add_heading("Operation",level=4) 
		operation_gb_heading.style = document.styles['GB Heading 4']
		
		document.add_paragraph("Use Game of Business for your company’s strategic decisions, because it will guide you to the dominant factors based on your conditions and situation. ")
		
		effect_gb_heading = document.add_heading("Effectiveness",level=4) 
		effect_gb_heading.style = document.styles['GB Heading 4']
		
		document.add_paragraph("Game of Business is effective at focusing your energy on the main tasks to achieve your business goal. Its comprehensive reasoning sharpens business strategy by increasing business awareness; this simplifies the decision process, while improving decision quality. ")
		
		opp_gb_heading = document.add_heading("Opportunity for Operation",level=4) 
		opp_gb_heading.style = document.styles['GB Heading 4']
		
		document.add_paragraph("When an opportunity arises to make a strategic business decision, find the relevant chapters in Game of Business: Planning, conducting business, progressing strategically, making tactical arrangements, allocating energy, using weaknesses and strengths, maneuvering, varying tactics, preparing your company, evaluating activities, evaluating situations, progressing fervently, and using information. ")
		
		length_gb_heading = document.add_heading("Length",level=4) 
		length_gb_heading.style = document.styles['GB Heading 4']
		
		document.add_paragraph("Game of Business has thirteen (13) chapters. ")
		
		aim_gb_heading = document.add_heading("Aim",level=4) 
		aim_gb_heading.style = document.styles['GB Heading 4']
		
		document.add_paragraph("Game of Business has three (3) main aims: ")
		ids = ['1.','2.','3.']
		parts = ["Make business a collaborative process that grows in accessibility, and improves your living standard. ","Make business success repeatable by deconstructing the process into repeatable steps. ", "Provide constructive feedback, by highlighting your company’s weaknesses and strengths, and recommending solutions. "]
		writer.display_table(document, ids, parts)
		
		gb_aw_connection_section = document.add_section(start_type=WD_SECTION.NEW_PAGE)
		gb_aw_connection_section.header_distance = Inches(0.25)
	
		section_title = "Game of Business Connection to the Art of War"
		generate_section_header(document, gb_aw_connection_section, book_title, ch_name, section_title)
	
		gb_aw_connection_heading = document.add_heading(section_title,level=3) 
		gb_aw_connection_heading.style = document.styles['GB Heading 3']
		
		gb_aw_connection_p1 = document.add_paragraph("Game of Business is derived from Sun Zi\'s Art of War, which was most likely created as a compilation of the teachings of Sun Zi (AKA Sun Wu, a military advisor who lived ~450 years before the Current Era. Sun Zi\'s Art of War was then likely studied, tested, and edited by military theorists, like Sun Bin. Its principles remain respected, partly because they aim to achieve a mutually beneficial, thriving community.")
		ref_run = gb_aw_connection_p1.add_run("5,7,9,10")
		ref_run.font.superscript = True
		document.add_paragraph("Business will always have the risk of becoming malevolent, and military will always exist for defense, so the goal of Game of Business is to minimize belligerence in favor of mutually beneficial business negotiations. To minimize the risk of business becoming combative, the system in which businesses operate must seek the optimal payoffs for all involved. So, companies must focus on customers and associates, instead of enemies or competitors. ")
	
		# add origin story section after gb background
		origin_section = document.add_section(start_type=WD_SECTION.NEW_PAGE)
		origin_section.header_distance = Inches(0.25)
	
	# add origin story
	section_title = "Origin of Sun Zi\'s Art of War"
	section_title_start = "Origin of "
	
	origin_heading = document.add_heading(section_title_start, level=2)
	origin_heading.style = document.styles['GB Heading 2']
	title_run = origin_heading.add_run("Sun Zi\'s Art of War")
	title_run.italic = True
	
	generate_section_header(document, origin_section, book_title, ch_name, section_title)
	
	#p1 = 'Sun Zi’s Art of War is the compiled teachings of military strategist, Sun Zi. “Zi” (pronounced “Tzu”) means “Master,” which is the title used by students to refer to their teacher. “Sun” is a family name. In Chinese, the family’s name is presented before the individual’s name or title; so, in English, his students would have referred to him as Master Sun. The direct translation of the original version of Sun Zi’s Art of War is Master Sun’s Military Methods. It is possible that Master Sun’s students compiled his teachings into a treatise, and called it Master Sun’s Military Methods, which later became popular in English under the title the Art of War.12 '
	#origin_paragraphs = [p1]
	origin_paragraphs = reader.extract_data("aw-origin","origin","tsv")
	#print("origin_paragraphs: " + str(origin_paragraphs))
	
	for p in origin_paragraphs:
	
		# if paragraph ends with number, take it as reference and make it superscript
		p_string = p[0]
		writer.add_formatted_paragraph(document, p_string)
	
	# add strategy tables
	
	table_list_sect = document.add_section(start_type=WD_SECTION.NEW_PAGE)
	
	section_title = "List of Tables"
	
	table_list_heading = document.add_heading(section_title, level=2)
	table_list_heading.style = document.styles['GB Heading 2']
	
	generate_section_header(document, ap_section, book_title, ch_name, section_title)
	
	# generate table overview
	book_table_ov_data = []
	all_table_overview_data = reader.extract_data("aw-tables","Tables","tsv")
	if re.search(",",book_title):
		# display comparison appendix table overview
		print("\n=== Display Comparison Appendix Table Overview ===\n")
		all_books_table_ov_data = []
		all_book_titles = book_title.split(",")
		book_table_ov_data = []
		for title in all_book_titles:
			print("title: " + title)
			book_table_ov_data = isolator.isolate_data_field(all_table_overview_data,"tables",title)
			all_books_table_ov_data.append(book_table_ov_data)
			
		num_rows = 1 # each row is new table for formatting purposes
	
		for table_idx in range(len(book_table_ov_data)):
		
			num_cols = len(all_books_table_ov_data) + 1 # no. books + 1 for table num col
		
			ap_tables_overview = document.add_table(rows=num_rows,cols=num_cols)
			ap_tables_overview.autofit = False
			ap_tables_overview.allow_autofit = False
			ap_tables_overview.alignment = WD_TABLE_ALIGNMENT.CENTER
		
			num_cell = ap_tables_overview.cell(0,0)
		
			title_cells = []
			for book_idx in range(len(all_books_table_ov_data)):
				#print("book_idx: " + str(book_idx))
				col_idx=book_idx+1 #offset 1 for num col
				title_cell = ap_tables_overview.cell(0,col_idx)
				title_cells.append(title_cell)
				
			# set cell width
			num_cell_width = 1
			num_cell.width = Inches(num_cell_width)
			#title_cell.width = Inches(4.9)
			all_titles_width = 6.5 - num_cell_width
			remaining_width = all_titles_width
			title_width = all_titles_width / len(title_cells)
			for cell_idx in range(len(title_cells)):
				#print("remaining_width: " + str(remaining_width))
				#print("princ_content_width: " + str(princ_content_width))
	
				cell = title_cells[cell_idx]
		
				if cell_idx == len(title_cells)-1:
					cell.width = Inches(remaining_width)
				else:
					cell.width = Inches(title_width)
			
					remaining_width -= title_width
		
			# write cell text
			table_overview = book_table_ov_data[table_idx]
			#print("table_overview: " + table_overview)
			table_overview_parts = table_overview.split(":")
			table_num = table_overview_parts[0]
			table_num = table_num.split(" ")[1] # take out "Table"
			num_cell.text = table_num # only need 1 table num for all books
			for book_idx in range(len(all_books_table_ov_data)):
				book_table_ov_data = all_books_table_ov_data[book_idx]
				table_overview = book_table_ov_data[table_idx]
				#print("table_overview: " + table_overview)
				table_overview_parts = table_overview.split(":")
				table_num = table_overview_parts[0] + ":"
				table_title = table_overview_parts[1].strip()
				
				cell = title_cells[book_idx]
				cell.text = table_title
		
			# style
			num_paragraph = num_cell.paragraphs[0]
			num_paragraph.style = document.styles['Tables Overview']
			for title_cell in title_cells:
				title_paragraph = title_cell.paragraphs[0]
				title_paragraph.style = document.styles['Tables Overview']
				
			writer.set_cell_border(
				num_cell,
				top={},
				bottom={},
				start={},
				end={"sz": 18, "color": "#830303", "val": "single"},
			)
			for cell in title_cells[:-1]:
				writer.set_cell_border(
					cell,
					top={},
					bottom={},
					start={},
					end={"sz": 9, "color": "#830303", "val": "dashed"},
				)
			cell_margin = 180
			writer.set_cell_margins(num_cell,end=cell_margin)
			for cell in title_cells:
				writer.set_cell_margins(cell,start=cell_margin,end=cell_margin)
	else:
		book_table_ov_data = isolator.isolate_data_field(all_table_overview_data,"tables",book_title)
	
		# display table with table overview data
		num_tables = len(book_table_ov_data) # num rows, 1 row per table
		num_cols = 2
		ap_tables_overview = document.add_table(rows=num_tables,cols=num_cols)
	
		for table_idx in range(num_tables):
			table_overview = book_table_ov_data[table_idx]
			#print("table_overview: " + table_overview)
		
			table_overview_parts = table_overview.split(":")
			table_num = table_overview_parts[0] + ":"
			table_title = table_overview_parts[1].strip()
		
			num_cell = ap_tables_overview.cell(table_idx,0)
			title_cell = ap_tables_overview.cell(table_idx,1)
		
			num_cell.text = table_num
			title_cell.text = table_title
		
			num_cell.width = Inches(1.6)
			title_cell.width = Inches(4.9)
		
			num_paragraph = num_cell.paragraphs[0]
			title_paragraph = title_cell.paragraphs[0]
			num_paragraph.style = document.styles['Tables Overview']
			title_paragraph.style = document.styles['Tables Overview']
			# num_paragraph.paragraph_format.line_spacing = 1
	# 		title_paragraph.paragraph_format.line_spacing = 1
	# 		num_paragraph.paragraph_format.space_before = Pt(0)
	# 		title_paragraph.paragraph_format.space_before = Pt(0)
	# 		num_paragraph.paragraph_format.space_after = Pt(6)
	# 		title_paragraph.paragraph_format.space_after = Pt(6)
	
	document.add_page_break()
	
	# add appendix tables
	if re.search(",",book_title):
		generate_comparison_appendix_tables(document, book_title, all_books_table_ov_data)
	else:
		generate_appendix_tables(document, book_title, book_table_ov_data)
	
	# add table A1
	#table_data_a1 = reader.extract_appendix_table_data()
	
# ref input in format ['1','Giles',...]
# output in format '1. Giles...'
def generate_reference(ref_data):

	ref = ''

	ref_num_idx = 0
	auth_last_name_idx = 1
	auth_first_name_idx = 2
	title_idx = 3
	pub_city_idx = 4
	pub_province_idx = 5
	publisher_idx = 6
	pub_year_idx = 7
	
	ref_num = ref_data[ref_num_idx]
	auth_last_name = ref_data[auth_last_name_idx]
	auth_first_name = ref_data[auth_first_name_idx]
	first_initial = auth_first_name[0]
	title = ref_data[title_idx]
	pub_city = ref_data[pub_city_idx]
	pub_province = ref_data[pub_province_idx]
	publisher = ref_data[publisher_idx]
	pub_year = ref_data[pub_year_idx]
	
	ref_part1 = ref_part2 = ''
	
	# format first and middle name
	if auth_first_name == 'n/a':
		ref_part1 = ref_num + ". " + auth_last_name
	elif re.search("\s",auth_first_name):
		# first and middle initials
		#print("found space in auth_first_name: " + auth_first_name)
		middle_name = re.split('\s',auth_first_name)[1]
		middle_initial = middle_name[0]
		ref_part1 = ref_num + ". " + auth_last_name + " " + first_initial + middle_initial
	else:
		ref_part1 = ref_num + ". " + auth_last_name + " " + first_initial
	
	if pub_province == 'n/a':
		ref_part2 = ". " + title + ". " + pub_city + ": " + publisher + "; " + pub_year + "."
	else:
		ref_part2 = ". " + title + ". " + pub_city + ", " + pub_province + ": " + publisher + "; " + pub_year + "."
	
	ref = ref_part1 + ref_part2
	
	#print("ref: " + ref)
	return ref
	
def generate_principles_references(document, book_title):

	ref_section = document.add_section(start_type=WD_SECTION.NEW_PAGE)
	
	ch_name = "References"
	
	generate_section_header(document, ref_section, book_title, ch_name)
	
	ref_heading = document.add_heading(ch_name, level=1)
	ref_heading.style = document.styles['GB Heading 1']
	
	# extract ref data
	if re.search('gb',book_title):
		book_title = 'gb'
	book_ref_data = reader.extract_data("reference data",book_title,"tsv")
	
	for ref_data in book_ref_data:
	
		ref = generate_reference(ref_data)
		
		ref_paragraph = document.add_paragraph(ref, style='GB Reference')
		#writer.display_reference(ref, document)
	
def generate_principles_index(document, book_title):

	idx_section = document.add_section(start_type=WD_SECTION.NEW_PAGE)
	
	ch_name = "Index"
	
	generate_section_header(document, idx_section, book_title, ch_name)
	
	idx_heading = document.add_heading(ch_name, level=1)
	idx_heading.style = document.styles['GB Heading 1']
	
	# insert index: generate the single line that inserts the INDEX field
	idx_paragraph = document.add_paragraph()
	run = idx_paragraph.add_run()
	fldChar = OxmlElement('w:fldChar') # creates a new element
	fldChar.set(qn('w:fldCharType'), 'begin') # sets attribute on element
	instrText = OxmlElement('w:instrText')
	instrText.set(qn('xml:space'), 'preserve') # sets attribute on element
	instrText.text = 'INDEX \\h "A" \\c "2"' # change settings depending on what you need
	
	fldChar2 = OxmlElement('w:fldChar')
	fldChar2.set(qn('w:fldCharType'), 'separate')
	fldChar3 = OxmlElement('w:t')
	fldChar3.text = "Right-click to update field."
	fldChar2.append(fldChar3)
	
	fldChar4 = OxmlElement('w:fldChar')
	fldChar4.set(qn('w:fldCharType'), 'end')
	
	r_element = run._r
	r_element.append(fldChar)
	r_element.append(instrText)
	r_element.append(fldChar2)
	r_element.append(fldChar4)
	p_element = idx_paragraph._p
	
def generate_acknowledgement(ack_data, book_title):

	#print("\n=== Generate Acknowledgement ===\n")
	#print("ack_data: " + str(ack_data))
	#print("book_title: " + book_title)

	ack = ''
	
	book_title = writer.format_book_title(book_title)
	
	ack = re.sub("<book title>",book_title,ack_data[0]) # book title mid sentence
	
	#print("ack: " + ack)
	return ack
	
def generate_principles_acknowledgements(document, book_title):

	ack_section = document.add_section(start_type=WD_SECTION.NEW_PAGE)
	
	ch_name = "Acknowledgements"
	
	ack_heading = document.add_heading(ch_name, level=1)
	ack_heading.style = document.styles['GB Heading 1']
	
	generate_section_header(document, ack_section, book_title, ch_name)
	
	# extract ack data
	aw_ack_data = reader.extract_data("acknowledgements data","aw","tsv")
	
	for ack_data in aw_ack_data:
	
		ack = generate_acknowledgement(ack_data, book_title)
		
		ack_paragraph = document.add_paragraph(style='GB Acknowledgement')
		
		keywords = [writer.format_book_title(book_title)]
		generate_runs(ack, ack_paragraph, keywords)
	
def generate_principles_back_page(document, book_title):

	# web url variable currently same for all transformations
	website_url = "https://www.gameofbusiness.net"

	back_section = document.add_section(start_type=WD_SECTION.NEW_PAGE)
	back_section.header.is_linked_to_previous = False
	
	# logo
	logo_text = 'GB'
	logo_paragraph = document.add_paragraph(logo_text, style="GB Logo")
	
	# book title paragraph
	book_title_paragraph = document.add_paragraph(style="GB Back Page")
	book_title_run = book_title_paragraph.add_run(writer.format_book_title(book_title))
	book_title_run.italic = True
	# subtitle paragraph
	subtitle_start = ''
	if re.search("gb",book_title):
		subtitle_start = 'Transformed from '
	if subtitle_start != '':
		book_title_paragraph.paragraph_format.space_after = 0
		subtitle_paragraph = document.add_paragraph(subtitle_start,style="GB Back Page")
		title_run = subtitle_paragraph.add_run('Sun Zi\'s Art of War')
		title_run.italic = True
	
	# created by paragraph
	author = 'Sun Zi'
	if re.search("gb",book_title):
		author = 'Game of Business'
	author_string = "Created by " + author + "."
	author_paragraph = document.add_paragraph(author_string,style="GB Back Page")
	
	# design and typesetting paragraph
	design_string = "Design and Typesetting created by Game of Business, LLC. "
	design_paragraph = document.add_paragraph(design_string,style="GB Back Page")
	
	# logo credit paragraph
	logo_credit_string = "Game of Business logo created by Game of Business, LLC. "
	logo_credit_paragraph = document.add_paragraph(logo_credit_string,style="GB Back Page")
	
	# info notice
	info_notice_string = "This book contains information adapted from authentic and highly regarded sources. Reasonable efforts have been made to publish reliable information, but the author and publisher cannot assume responsibility for the validity of all materials or the consequences of their use. If any copyright material has not been acknowledged please write and let us know so we may rectify in any future reprint. "
	info_notice_paragraph = document.add_paragraph(info_notice_string,style="GB Back Page")
	
	# electronic material
	electronic_material_string = "Electronic material from this work may be accessed at " + website_url + " or contact Game of Business, LLC (GB), sunzi@gameofbusiness.net. "
	electronic_material_paragraph = document.add_paragraph(electronic_material_string,style="GB Back Page")
	
	# trademark notice
	trademark_notice_paragraph = document.add_paragraph(style="GB Back Page")
	trademark_notice_title = "Trademark Notice"
	trademark_notice_title_run = trademark_notice_paragraph.add_run(trademark_notice_title + ": ")
	trademark_notice_title_run.bold = True
	trademark_notice_string = "Product or corporate names may be trademarks or registered trademarks, and are used only for identification and explanation without intent to infringe. "
	trademark_notice_run = trademark_notice_paragraph.add_run(trademark_notice_string)
	
	# pub year
	pub_year = "2018"
	pub_year_string = "Published in " + pub_year + " by "
	pub_year_paragraph = document.add_paragraph(pub_year_string,style="GB Back Page")
	pub_year_paragraph.paragraph_format.space_after = 0
	
	# publisher
	publisher = "Game of Business, LLC"
	pub_paragraph = document.add_paragraph(publisher,style="GB Back Page")
	pub_paragraph.paragraph_format.space_after = 0
	
	# website
	web_paragraph = document.add_paragraph(website_url,style="GB Back Page")
	web_paragraph.paragraph_format.space_after = 0
	
	# email
	email_string = "sunzi@gameofbusiness.net"
	email_paragraph = document.add_paragraph(email_string,style="GB Back Page")
	
	# countdown
	countdown = ""
	for num in range(10, 0, -1):
		countdown += str(num) + " "
	countdown_paragraph = document.add_paragraph(countdown,style="GB Back Page")
	
	# ISBN
	isbn_num = '978-0-9999999-0-5'
	if book_title == 'gb':
		isbn_num = '978-0-9999998-0-5'
	elif book_title == 'aw,gb':
		isbn_num = '978-0-9999997-0-5'
	isbn_string = 'ISBN-X: ' + isbn_num
	isbn_paragraph = document.add_paragraph(isbn_string,style="GB Back Page")
	
def generate_book_conclusion(document, book_title):

	print("\n=== Generate Book Conclusion ===\n")

	conc_section = document.add_section(start_type=WD_SECTION.NEW_PAGE)
	conc_section.header_distance = Inches(0.25)

	ch_name = "Conclusion"

	conc_heading = document.add_heading(ch_name,level=1)
	conc_heading.style = document.styles["GB Heading 1"]
	
	section_title = "Main Question Explored by Game of Business"

	main_q_heading = document.add_heading(section_title,level=2)
	main_q_heading.style = document.styles["GB Heading 2"]
	
	generate_section_header(document, conc_section, book_title, ch_name, section_title)
	
	document.add_paragraph("If your company is threatened by an opponent, what is your company’s best response? ")
	
	data_type = "gb-conc"
	input_type = "conc - main question"
	extension = "tsv"
	conc_main_question_paragraphs = reader.extract_data(data_type,input_type,extension)
	
	document.add_paragraph(conc_main_question_paragraphs[0]) # intro to section
	document.add_paragraph(conc_main_question_paragraphs[1]) # table 1 explanation 
# 	
# 	data_type = "gb-conc"
# 	input_type = "conc - table captions"
# 	table_captions_data = reader.extract_data(data_type,input_type,extension)
# 	print("table_captions_data: " + str(table_captions_data))
# 	table_captions = generate_table_captions(table_captions_data,'conclusion')
# 	print("table_captions: " + str(table_captions))
	
	data_type = "gb-conc"
	input_type = "conc - table 1"
	table1_data = reader.extract_data(data_type,input_type,extension)
	#print("table1_data: " + str(table1_data))
	
	for row in table1_data:
		id = row[0]
		princ = row[1]
		reason = row[2]
		
		p_string = "Principle " + id + " states, \"" + princ + "\" " + reason
		#print("p_string: " + p_string)
		
		document.add_paragraph(p_string,style="List Bullet")
		
	document.add_paragraph(conc_main_question_paragraphs[2]) # table 2 explanation
	
	input_type = "conc - table 2"
	table2_data = reader.extract_data(data_type,input_type,extension)
	#print("table2_data: " + str(table2_data))
	
	for row in table2_data:
		id = row[0]
		princ = row[1]
		reason = row[2]
		
		p_string = "Principle " + id + " states, \"" + princ + "\" " + reason
		#print("p_string: " + p_string)
		
		document.add_paragraph(p_string,style="List Bullet")
		
	q_section = document.add_section(start_type=WD_SECTION.NEW_PAGE)
	q_section.header_distance = Inches(0.25)
		
	section_title = "Questions About Game of Business"
	generate_section_header(document, q_section, book_title, ch_name, section_title)
	
	gb_q_for_translators_heading_string = section_title + " for Sun Zi\'s Art of War Translators"
	
	gb_q_for_translators_heading = document.add_heading(gb_q_for_translators_heading_string,level=2)
	gb_q_for_translators_heading.style = document.styles["GB Heading 2"]
	
	input_type = "conc - GB questions for translators"
	gb_q_for_translators_data = reader.extract_data(data_type,input_type,extension)
	#print("gb_q_for_translators_data: " + str(gb_q_for_translators_data))
	
	for row in gb_q_for_translators_data:
		q = row[0]
		
		document.add_paragraph(q,style="List Bullet")
		
	
	
	gb_q_for_you_heading_string = section_title + " for You"
	
	gb_q_for_you_heading = document.add_heading(gb_q_for_you_heading_string,level=2)
	gb_q_for_you_heading.style = document.styles["GB Heading 2"]
	
	input_type = "conc - GB questions for you"
	gb_q_for_you_data = reader.extract_data(data_type,input_type,extension)
	#print("gb_q_for_you_data: " + str(gb_q_for_you_data))
	
	for row in gb_q_for_you_data:
		q = row[0]
		
		document.add_paragraph(q,style="List Bullet")
		
	feedback_string = "Send feedback to feedback@gameofbusiness.net. "
	fdbk_p = document.add_paragraph(feedback_string,style="GB Feedback")
		
		
		
		