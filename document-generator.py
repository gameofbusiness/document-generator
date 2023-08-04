# principle-document-generator.py
# generate a document of principles
# create word docx for a group of principles
#

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.dml.color import ColorFormat
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.text.tabstops import WD_TAB_ALIGNMENT, WD_TAB_LEADER

#set cell border
from docx.table import _Cell
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn

import re, os
import isolator, reader, writer, generator

# === Parameters Needed to Create Document === 
ch1_title = "1. Planning"
ch2_title = "2. Waging War"
ch_titles = [ch1_title,ch2_title]

ch1_page_num = "1"
ch2_page_num = "8"
ch_page_nums = [ch1_page_num, ch2_page_num]

st1 = 'Cover Page'
st2 = 'Table of Contents'
st3 = 'Overview'
st4 = '5 Factors'
st5 = '7 Considerations'
section_titles = [st1,st2,st3,st4,st5]

# Chapter Overviews (selected principles or parts of principles)
ch1_overview = ['o1','o2','o3','o4','o5','o6']
ch_overviews = [ch1_overview]

# === Helper Functions ===

def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
 
    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
 
    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
 
            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
 
            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))
                    
def set_cell_margins(cell: _Cell, **kwargs):
    """
    cell:  actual cell instance you want to modify

    usage:

        set_cell_margins(cell, top=50, start=50, bottom=50, end=50) if we want x=1", then y=1440, so to get x=1/4", set y=1440/4=360

    provided values are in twentieths of a point (1/1440 of an inch).
    read more here: http://officeopenxml.com/WPtableCellMargins.php
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')

    for m in [
        "top",
        "start",
        "bottom",
        "end",
    ]:
        if m in kwargs:
            node = OxmlElement("w:{}".format(m))
            node.set(qn('w:w'), str(kwargs.get(m)))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)

    tcPr.append(tcMar)

# === Main Functions ===


# def generate_aw_page_1(document):
# 
# 
# 	p1 = 'Sun Zi said: The art of war is of vital importance to the State.'
# 	p2 = 'p2'
# 	
# 	p3 = 'p3'
# 	p4 = 'p4'
# 	p5 = 'p5'
# 	
# 	section1 = [p1,p2]
# 	section2 = [p3,p4,p5]
# 	
# 	sections = [section1, section2]
# 	
# 	
# 	
# 	for section_idx in range(len(sections)): 
# 		section = sections[section_idx]
# 		header = headers[section_idx]
# 	
# 		document.add_heading(header, level=1)
# 	
# 		for principle in section:
# 	
# 			paragraph = document.add_paragraph(principle)
# 			
# 	document.add_page_break()
	
# overview is always section 1 of the chapter
def generate_ch_overview(document, ch, book_title='aw'):

	#print("\n=== Generate Chapter Overview ===\n")

	ch_title_idx = 0
	ch_title = re.sub("\s+"," ",ch[ch_title_idx])
	#print("ch_title: " + ch_title)
	
	ch_sections_idx = 1
	ch_sections = ch[ch_sections_idx]
	#print("ch_sections: " + str(ch_sections))
	
	ov_sec_idx = 0
	ch_ov_section = ch_sections[ov_sec_idx]
	#print("ch_ov_section: " + str(ch_ov_section))
	
	ch_num = isolator.isolate_ch_num(ch_title)
	
	ch_section_titles = reader.read_section_titles(book_title, ch_num)
	# section_titles_dict = reader.read_json("chapter section titles")
# 	ch_section_titles = section_titles_dict[ch_num]
	#print("ch_section_titles: " + str(ch_section_titles))
	
	# to get ch name remove numbers, dots, and leading space from ch title
	ch_name = re.sub("\d|\.","",ch_title).strip()
	#print("ch_name: " + ch_name)
	
	doc_section = document.add_section()
	doc_section.start_type = WD_SECTION.NEW_PAGE
	
	# doc_footer = doc_section.footer
# 	doc_footer.is_linked_to_previous = False
# 	footer_block = doc_footer.paragraphs[0]
# 	writer.add_page_number(footer_block.add_run())
# 	footer_block.alignment = WD_ALIGN_PARAGRAPH.RIGHT
# 	
# 	doc_section.different_first_page_header_footer = True
# 	sectPr = doc_section._sectPr
# 
# 	pgNumType = OxmlElement('w:pgNumType')
# 	pgNumType.set(ns.qn('w:start'), "0")
# 	sectPr.append(pgNumType)

	#book_title = 'The Art of War'
	section_title = 'Overview'

	generator.generate_section_header(document, doc_section, book_title, ch_name, section_title, ch_num)
	
	# doc_header = doc_section.header
# 	doc_header.is_linked_to_previous = False
# 	header_block = doc_header.paragraphs[0]
# 	header_block.style = document.styles['Header Footer']
# 	
# 	# add tab stop so text can be aligned left and right 
# 	tab_stops = header_block.paragraph_format.tab_stops
# 	#print("tab_stops: " + str(tab_stops))
# 	tab_stop = tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
# 	#tab_stop.position = Inches(6.5)
# 	#tab_stop.alignment = WD_TAB_ALIGNMENT.RIGHT
# 	
# 	header_string = book_title + ", " + ch_num + ". " + ch_name + "\t" + section_title
# 	header_block.text = header_string
	
	ch_title_heading = document.add_paragraph(ch_title, style='GB Heading 1')
	ch_title_heading_format = ch_title_heading.paragraph_format
	
	for ov in ch_ov_section:
		#print("overview point: " + ov)
		overview_point = document.add_paragraph(ov, style="GB Overview") #.add_run(ov)
		#overview_point.font.italic = True
		
		#format paragraph line spacing
		#paragraph_format = overview_point.paragraph_format
		#paragraph_format.line_spacing = 1.5
		#paragraph_format.space_before = Pt(0)
		#paragraph_format.space_after = Pt(12)
	
	#document.add_page_break()
	
def generate_ch_sections(document, ch, ch_paragraphs=[], book_title='aw'):

	#print("\n=== Generate Chapter Sections ===\n")
	
	ch_titles_idx = 0
	ch_title = ch[ch_titles_idx]
	
	ch_sections_idx = 1
	ch_sections = ch[ch_sections_idx]
	
	ch_princ_data_idx = 2
	ch_princ_data = ch[ch_princ_data_idx]

	ch_num = isolator.isolate_ch_num(ch_title)
	
	ch_section_titles = []
	if not re.search(",",book_title):
		ch_section_titles = reader.read_section_titles(book_title, ch_num) # read section titles
	# all_section_titles_dicts = reader.read_json("chapter section titles")
# 	print("all_section_titles_dicts: " + str(all_section_titles_dicts))
# 	section_titles_dict = all_section_titles_dicts[book_title]
# 	ch_section_titles = section_titles_dict[ch_num]
	#print("ch_section_titles: " + str(ch_section_titles))
	
	# to get ch name remove numbers, dots, and leading space from ch title
	ch_name = re.sub("\d|\.","",ch_title).strip()
	#print("ch_name: " + ch_name)

	for section_idx in range(len(ch_sections)):
		section = ch_sections[section_idx]
		#print("section " + str(section_idx+1) + ": " + str(section))
		
		if section_idx != 0: # skip the first one bc it is the overview already displayed
		
			doc_section = document.add_section()
			doc_section.start_type = WD_SECTION.NEW_PAGE
			
			# doc_footer = doc_section.footer
# 			doc_footer.is_linked_to_previous = True
# 			footer_block = doc_footer.paragraphs[0]
# 			writer.add_page_number(footer_block.add_run())

			#book_title = 'The Art of War'
			
			section_title = ''
			if len(ch_section_titles) > 0:
				section_title = ch_section_titles[section_idx-1] # subtract one bc first section always overview so not written
			else:
				print("Warning: No ch_section_titles so we cannot get section_title!")

			generator.generate_section_header(document, doc_section, book_title, ch_name, section_title, ch_num)
	
			added_manual_page_break = False # bc complicated to check if next principle will fit on page before deciding to add page break
			#need_page_breaks_after = ['2.3.4.'] # manually add page breaks after these principles bc complicated to check if next principle will fit on page before deciding to add page break
		
			for group_idx in range(len(section)):
				group = section[group_idx]
			
				#print("group: " + str(group))
				
				#group_table = document.add_table(rows=len(all_sub_principles), cols=2) # the no. rows in the init table = no. sub-principle sets
				
				writer.display_principle_tables(document, group, ch_princ_data, ch_paragraphs)					
					 
						#print("check if need manual page break")
						#print("princ_num: " + princ_num)
					
						# added_manual_page_break = False
# 						for need_page_break in need_page_breaks_after:
# 							if princ_num == need_page_break:
# 								#print("need manual page break")
# 								document.add_page_break()
# 								added_manual_page_break = True
# 								break
					
				# if not last group in section, bc already section break
				# also if manual page break not added. check if last member of group needed manual page break
				
				if group_idx != len(section) - 1 and not added_manual_page_break:
					#print("Add paragraph between groups.")
					document.add_paragraph()
				
		
def generate_book_chapter(document, ch, ch_paragraphs=[], book_title='aw'):

	#print("\n=== Generate Book Chapter ===\n")

	generate_ch_overview(document, ch, book_title)
	
	generate_ch_sections(document, ch, ch_paragraphs, book_title)
	
def generate_book_chapters(document, book_title, all_ch_paragraphs, all_ch_data):

	print("\n=== Generate Book Chapters ===\n")

	for ch_data_idx in range(len(all_ch_data)):
		ch_data = all_ch_data[ch_data_idx]
		#print("ch_data: " + str(ch_data))
		ch_paragraphs = all_ch_paragraphs[ch_data_idx]
	
		generate_book_chapter(document, ch_data, ch_paragraphs, book_title)

# def generate_aw_chapter(document, ch, ch_paragraphs=[]):
# 
# 	book_title = 'aw'
# 
# 	generate_ch_overview(document, ch, book_title)
# 	
# 	generate_ch_sections(document, ch, ch_paragraphs, book_title)

# overview is always section 1 of the chapter
def generate_comparison_ch_overview(document, book_title, ch_all_book_data):

	print("\n=== Generate Comparison Chapter Overview ===\n")

	ch_titles = []
	ch_ov_sections = []
	ch_names = []
	
	ch_num = '1'
	ch_name = ''
	ch_title = ''

	for cur_book_all_ch_data_idx in range(len(ch_all_book_data)):
		cur_book_all_ch_data = ch_all_book_data[cur_book_all_ch_data_idx]
		print("cur_book_all_ch_data: " + str(cur_book_all_ch_data))

		ch_title_idx = 0
		cur_ch_title = re.sub("\s+"," ",cur_book_all_ch_data[ch_title_idx])
		#print("cur_ch_title: " + cur_ch_title)
		ch_titles.append(cur_ch_title)
		
		if cur_book_all_ch_data_idx == 0:
			ch_num = isolator.isolate_ch_num(cur_ch_title)
	
		ch_sections_idx = 1
		ch_sections = cur_book_all_ch_data[ch_sections_idx]
		print("ch_sections: " + str(ch_sections))
	
		ov_sec_idx = 0
		ch_ov_section = ch_sections[ov_sec_idx]
		print("ch_ov_section: " + str(ch_ov_section))
		ch_ov_sections.append(ch_ov_section)
	
		# to get ch name remove numbers, dots, and leading space from ch title
		cur_ch_name = re.sub("\d|\.","",cur_ch_title).strip()
		print("cur_ch_name: " + cur_ch_name)
		ch_names.append(cur_ch_name)
	
	# format chapter name
	if len(ch_names) > 0:
		print("ch_names: " + str(ch_names))
		for name_idx in range(len(ch_names)):
			if name_idx == 0:
				ch_name = ch_names[name_idx]
			else:
				ch_name += " | " + ch_names[name_idx]
				
	doc_section = document.add_section()
	doc_section.start_type = WD_SECTION.NEW_PAGE
	
	print("ch_name: " + ch_name)
	section_title = 'Overview'
	generator.generate_section_header(document, doc_section, book_title, ch_name, section_title, ch_num)
	
	if len(ch_names) > 0:
		
		ch_title = ch_num + ". " + ch_name
		ch_title_heading = document.add_paragraph(ch_title, style='GB Heading 1')
	
		num_rows = len(ch_ov_section) 
		num_cols = len(ch_titles)
		ov_table = document.add_table(num_rows,num_cols)
	
		for ch_ov_section_idx in range(len(ch_ov_sections)): # for each chapter
			ch_ov_section = ch_ov_sections[ch_ov_section_idx]
			print("ch_ov_section: " + str(ch_ov_section))
			for ov_idx in range(len(ch_ov_section)):
				ov = ch_ov_section[ov_idx]
				print("overview point: " + str(ov))
				#overview_point = document.add_paragraph(ov, style="GB Overview") 
				ov_cell = ov_table.cell(ov_idx,ch_ov_section_idx)
				ov_cell.text = ov
				
				ov_cell_p = ov_cell.paragraphs[0]
				ov_cell_p.style = document.styles['GB Overview']
				ov_cell_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
				
				if ch_ov_section_idx != len(ch_ov_sections)-1: # all but last ch
					set_cell_border(
						ov_cell,
						top={},
						bottom={},
						start={},
						end={"sz": 9, "color": "#830303", "val": "dashed"},
					)
					
				if ch_ov_section_idx == 0: # first ch
					set_cell_margins(ov_cell, end=360) # 1/1440 of an inch
				elif ch_ov_section_idx == len(ch_ov_sections)-1: # last ch
					set_cell_margins(ov_cell, start=360) # 1/1440 of an inch
				else: # all but first and last ch
					set_cell_margins(ov_cell, start=360, end=360) # 1/1440 of an inch
	
	
def generate_comparison_ch_sections(document, book_title, ch_all_book_data, ch_all_book_paragraphs):

	print("\n=== Generate Comparison Chapter Sections ===\n")
	
	ch_titles = []
	ch_all_book_princ_data = []
	ch_names = []
	
	ch_num = '1'
	ch_name = ''
	ch_title = ''
	
	for cur_book_all_ch_data_idx in range(len(ch_all_book_data)):
		cur_book_all_ch_data = ch_all_book_data[cur_book_all_ch_data_idx]
		print("cur_book_all_ch_data: " + str(cur_book_all_ch_data))
	
		ch_title_idx = 0
		cur_ch_title = re.sub("\s+"," ",cur_book_all_ch_data[ch_title_idx])
		#print("cur_ch_title: " + cur_ch_title)
		ch_titles.append(cur_ch_title)
		
		if cur_book_all_ch_data_idx == 0:
			ch_num = isolator.isolate_ch_num(cur_ch_title)
	
		ch_sections_idx = 1
		ch_sections = cur_book_all_ch_data[ch_sections_idx]
		print("ch_sections: " + str(ch_sections))
	
		ch_princ_data_idx = 2
		ch_princ_data = cur_book_all_ch_data[ch_princ_data_idx]
		print("ch_princ_data: " + str(ch_princ_data))
		ch_all_book_princ_data.append(ch_princ_data)
	
		# to get ch name remove numbers, dots, and leading space from ch title
		cur_ch_name = re.sub("\d|\.","",cur_ch_title).strip()
		print("cur_ch_name: " + cur_ch_name)
		ch_names.append(cur_ch_name)
		
		all_ch_section_titles = reader.read_comparison_section_titles(book_title, ch_num) # read section titles

	# format chapter name
	if len(ch_names) > 0:
		print("ch_names: " + str(ch_names))
		for name_idx in range(len(ch_names)):
			if name_idx == 0:
				ch_name = ch_names[name_idx]
			else:
				ch_name += " | " + ch_names[name_idx]

	for section_idx in range(len(ch_sections)):
		section = ch_sections[section_idx]
		#print("section " + str(section_idx+1) + ": " + str(section))
		
		if section_idx != 0: # skip the first one bc it is the overview already displayed
		
			doc_section = document.add_section()
			doc_section.start_type = WD_SECTION.NEW_PAGE
			
			#section_title = '' # blank for comparison bc too long to fit on one line. ch name already too long for one line so maybe keep?
			section_title = ''
			if len(all_ch_section_titles) > 0:
				for ch_section_titles_idx in range(len(all_ch_section_titles)):
					ch_section_titles = all_ch_section_titles[ch_section_titles_idx]
					if len(ch_section_titles) > 0:
						if ch_section_titles_idx == 0:
							section_title = ch_section_titles[section_idx-1] # subtract one bc first section always overview so not written
						else:
							section_title += " | " + ch_section_titles[section_idx-1]
					else:
						print("Warning: No ch_section_titles so we cannot get section_title!")
			else:
				print("Warning: No all_ch_section_titles so we cannot get ch_section_titles!")
			
			
			generator.generate_section_header(document, doc_section, book_title, ch_name, section_title, ch_num)
	
			for group_idx in range(len(section)):
				group = section[group_idx]
				#print("group: " + str(group))
				
				writer.display_comparison_principle_tables(document, group, ch_all_book_princ_data, ch_all_book_paragraphs)					
					
				# if not last group in section, bc already section break				
				if group_idx != len(section) - 1:
					#print("Add paragraph between groups.")
					document.add_paragraph()

def generate_comparison_chapter(document, book_title, ch_all_book_data, ch_all_book_paragraphs):

	print("\n=== Generate Comparison Chapter ===\n")

	generate_comparison_ch_overview(document, book_title, ch_all_book_data)
	
	#generate_ch_sections(document, ch_all_book_data[0], ch_all_book_paragraphs[0], book_title.split(" ")[0])
	generate_comparison_ch_sections(document, book_title, ch_all_book_data, ch_all_book_paragraphs)

# format [b1c1, b1c2], [b2c1, b2c2]
def generate_comparison_chapters(document, book_title, all_books_all_ch_data, all_books_all_ch_paragraphs):

	print("\n=== Generate Comparison Chapters ===\n")

	book1_all_ch_data = all_books_all_ch_data[0] # all books have same no. chapters
	num_chapters = len(book1_all_ch_data)

	for ch_idx in range(num_chapters): # for each chapter
	
		cur_ch_all_book_paragraphs = []
		cur_ch_all_book_data = []
	
		# take chapter of each book 
		num_books = len(all_books_all_ch_paragraphs)
		for book_idx in range(num_books):
		
			cur_book_all_ch_paragraphs = all_books_all_ch_paragraphs[book_idx]
			cur_book_cur_ch_paragraphs = cur_book_all_ch_paragraphs[ch_idx]
			
			cur_book_all_ch_data = all_books_all_ch_data[book_idx]
			cur_book_cur_ch_data = cur_book_all_ch_data[ch_idx]
		
			cur_ch_all_book_paragraphs.append(cur_book_cur_ch_paragraphs)
			cur_ch_all_book_data.append(cur_book_cur_ch_data)
		
		generate_comparison_chapter(document, book_title, cur_ch_all_book_data, cur_ch_all_book_paragraphs)
		
		

def generate_document_styles(document, book_title='aw'):

	styles = document.styles
	#latent_styles = styles.latent_styles
	
	# EXISTING STYLES
	
	# Normal
	normal_style = styles['Normal']
	# font
	normal_style.font.name = 'Helvetica' 
	normal_style.font.size = Pt(14)
	normal_style.font.color.rgb = RGBColor(0,0,0) #from_string('000')
	normal_style.font.bold = False
	normal_style.font.italic = False
	# paragraph format
	normal_style.paragraph_format.space_before = Pt(0)
	normal_style.paragraph_format.space_after = Pt(12)
	normal_style.paragraph_format.line_spacing = 1.5
	normal_style.paragraph_format.keep_together = True
	if re.search(",",book_title):
		normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
	else:
		normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
	
	# CUSTOM STYLES
	
	# GB Heading 1
	ch_title_heading_style = styles.add_style('GB Heading 1', WD_STYLE_TYPE.PARAGRAPH)
	ch_title_heading_style.base_style = styles['Heading 1']
	#ch_title_heading_style = styles['Heading 1'] #'Heading 1' #document.styles['Heading 1']
	# font
	ch_title_heading_style.font.name = 'Helvetica' 
	ch_title_heading_style.font.size = Pt(20)
	ch_title_heading_style.font.bold = True
	ch_title_heading_style.font.color.rgb = RGBColor(0,0,0) #from_string('000')
	# paragraph format
	ch_title_heading_style.paragraph_format.space_before = Pt(0)
	ch_title_heading_style.paragraph_format.space_after = Pt(12)
	ch_title_heading_style.paragraph_format.line_spacing = 1.5
	
	# GB Heading 2
	h2_style = styles.add_style('GB Heading 2', WD_STYLE_TYPE.PARAGRAPH)
	h2_style.base_style = styles['Heading 2']
	# font
	h2_style.font.name = 'Helvetica' 
	h2_style.font.size = Pt(18)
	h2_style.font.bold = True
	h2_style.font.color.rgb = RGBColor(0,0,0) #from_string('000')
	# paragraph format
	h2_style.paragraph_format.space_before = Pt(0)
	h2_style.paragraph_format.space_after = Pt(12)
	h2_style.paragraph_format.line_spacing = 1.5
	
	# GB Heading 3
	h3_style = styles.add_style('GB Heading 3', WD_STYLE_TYPE.PARAGRAPH)
	h3_style.base_style = styles['Heading 3']
	# font
	h3_style.font.name = 'Helvetica' 
	h3_style.font.size = Pt(16)
	h3_style.font.bold = True
	h3_style.font.color.rgb = RGBColor(0,0,0) #from_string('000')
	# paragraph format
	h3_style.paragraph_format.space_before = Pt(0)
	h3_style.paragraph_format.space_after = Pt(12)
	h3_style.paragraph_format.line_spacing = 1.5
	
	# GB Heading 4
	h4_style = styles.add_style('GB Heading 4', WD_STYLE_TYPE.PARAGRAPH)
	h4_style.base_style = styles['Heading 4']
	# font
	h4_style.font.name = 'Helvetica' 
	h4_style.font.size = Pt(16)
	h4_style.font.bold = False
	h4_style.font.italic = True
	h4_style.font.color.rgb = RGBColor(0,0,0) #from_string('000')
	# paragraph format
	h4_style.paragraph_format.space_before = Pt(0)
	h4_style.paragraph_format.space_after = Pt(12)
	h4_style.paragraph_format.line_spacing = 1.5
	
	# GB Heading 5
	h5_style = styles.add_style('GB Heading 5', WD_STYLE_TYPE.PARAGRAPH)
	h5_style.base_style = styles['Heading 5']
	# font
	h5_style.font.name = 'Helvetica' 
	h5_style.font.size = Pt(14)
	h5_style.font.bold = True
	h5_style.font.color.rgb = RGBColor(0,0,0) #from_string('000')
	# paragraph format
	h5_style.paragraph_format.space_before = Pt(0)
	h5_style.paragraph_format.space_after = Pt(12)
	h5_style.paragraph_format.line_spacing = 1.5
	
	# GB Overview
	ov_style = styles.add_style('GB Overview', WD_STYLE_TYPE.PARAGRAPH)
	ov_style.base_style = styles['Normal']
	# font
	ov_style.font.name = 'Helvetica' 
	ov_style.font.size = Pt(14)
	ov_style.font.color.rgb = RGBColor(0,0,0) #from_string('000')
	ov_style.font.bold = False
	ov_style.font.italic = True
	# paragraph format
	ov_style.paragraph_format.space_before = Pt(0)
	ov_style.paragraph_format.space_after = Pt(12)
	ov_style.paragraph_format.line_spacing = 1.5
	ov_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
	ov_style.paragraph_format.keep_together = True
	
	# GB Header and Footer
	hf_style = styles.add_style('Header Footer', WD_STYLE_TYPE.PARAGRAPH)
	hf_style.base_style = styles['Normal']
	# font
	hf_style.font.name = 'Helvetica' 
	hf_style.font.size = Pt(12)
	hf_style.font.bold = False
	hf_style.font.color.rgb = RGBColor(0,0,0) #from_string('000')
	# paragraph format
	hf_style.paragraph_format.space_before = Pt(0)
	hf_style.paragraph_format.space_after = Pt(6)
	hf_style.paragraph_format.line_spacing = 1.15
	
	# GB Cover
	cover_style = styles.add_style('GB Cover', WD_STYLE_TYPE.PARAGRAPH)
	cover_style.base_style = styles['Normal']
	# font
	cover_style.font.name = 'Helvetica' 
	cover_style.font.size = Pt(100)
	cover_style.font.bold = True
	cover_style.font.color.rgb = RGBColor(0,0,0) #from_string('000')
	# paragraph format
	cover_style.paragraph_format.space_before = Pt(0)
	cover_style.paragraph_format.space_after = Pt(18)
	cover_style.paragraph_format.line_spacing = 1.15
	cover_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	
	# TOC Heading
	toc_heading_style = styles.add_style('GB TOC Heading', WD_STYLE_TYPE.PARAGRAPH)
	toc_heading_style.base_style = styles['Normal']
	# font
	toc_heading_style.font.name = 'Helvetica' 
	toc_heading_style.font.size = Pt(18)
	toc_heading_style.font.bold = True
	toc_heading_style.font.color.rgb = RGBColor(0,0,0) #from_string('000')
	# paragraph format
	toc_heading_style.paragraph_format.space_before = Pt(0)
	toc_heading_style.paragraph_format.space_after = Pt(12)
	
	# List of Tables Overview
	tables_overview_style = styles.add_style('Tables Overview', WD_STYLE_TYPE.PARAGRAPH)
	tables_overview_style.base_style = styles['Normal']
	# font
	tables_overview_style.font.name = 'Helvetica' 
	tables_overview_style.font.size = Pt(14)
	tables_overview_style.font.bold = False
	tables_overview_style.font.color.rgb = RGBColor(0,0,0) #from_string('000')
	# paragraph format
	tables_overview_style.paragraph_format.space_before = Pt(0)
	tables_overview_style.paragraph_format.space_after = Pt(6)
	tables_overview_style.paragraph_format.line_spacing = 1.15
	
	# Appendix table data
	ap_table_style = styles.add_style('Table Data', WD_STYLE_TYPE.PARAGRAPH)
	ap_table_style.base_style = styles['Normal']
	# font
	ap_table_style.font.name = 'Helvetica' 
	ap_table_style.font.size = Pt(13)
	ap_table_style.font.bold = False
	ap_table_style.font.color.rgb = RGBColor(0,0,0) #from_string('000')
	# paragraph format
	ap_table_style.paragraph_format.space_before = Pt(0)
	ap_table_style.paragraph_format.space_after = Pt(6)
	ap_table_style.paragraph_format.line_spacing = 1
	ap_table_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
	ap_table_style.paragraph_format.keep_together = True
	
	# Appendix table field titles
	ap_field_title_style = styles.add_style('Field Title', WD_STYLE_TYPE.PARAGRAPH)
	ap_field_title_style.base_style = styles['Normal']
	# font
	ap_field_title_style.font.name = 'Helvetica' 
	ap_field_title_style.font.size = Pt(13)
	ap_field_title_style.font.bold = True
	ap_field_title_style.font.color.rgb = RGBColor(0,0,0) #from_string('000')
	# paragraph format
	ap_field_title_style.paragraph_format.space_before = Pt(0)
	ap_field_title_style.paragraph_format.space_after = Pt(6)
	ap_field_title_style.paragraph_format.line_spacing = 1
	ap_field_title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
	ap_field_title_style.paragraph_format.keep_together = True
	
	# Table caption
	ap_table_caption_style = styles.add_style('GB Table Caption', WD_STYLE_TYPE.PARAGRAPH)
	ap_table_caption_style.base_style = styles['Normal']
	# font
	ap_table_caption_style.font.name = 'Helvetica' 
	ap_table_caption_style.font.size = Pt(13)
	ap_table_caption_style.font.bold = False
	ap_table_caption_style.font.color.rgb = RGBColor(0,0,0) #from_string('000')
	# paragraph format
	ap_table_caption_style.paragraph_format.space_before = Pt(18)
	ap_table_caption_style.paragraph_format.space_after = Pt(6)
	ap_table_caption_style.paragraph_format.line_spacing = 1
	ap_table_caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
	ap_table_caption_style.paragraph_format.keep_together = True
	ap_table_caption_style.paragraph_format.keep_with_next = True
	
	# reference entry
	ref_style = styles.add_style('GB Reference', WD_STYLE_TYPE.PARAGRAPH)
	ref_style.base_style = styles['Normal']
	# font
	ref_style.font.name = 'Helvetica' 
	ref_style.font.size = Pt(14)
	ref_style.font.bold = False
	ref_style.font.color.rgb = RGBColor(0,0,0) #from_string('000')
	# paragraph format
	ref_style.paragraph_format.space_before = Pt(0)
	ref_style.paragraph_format.space_after = Pt(12)
	ref_style.paragraph_format.line_spacing = 1.15
	ref_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
	ref_style.paragraph_format.keep_together = True
	ref_style.paragraph_format.keep_with_next = False
	# hanging indent
	ref_style.paragraph_format.left_indent = Inches(0.5)
	ref_style.paragraph_format.first_line_indent = Inches(-0.5)
	
	# acknowledgement entry
	ack_style = styles.add_style('GB Acknowledgement', WD_STYLE_TYPE.PARAGRAPH)
	ack_style.base_style = styles['Normal']
	# font
	ack_style.font.name = 'Helvetica' 
	ack_style.font.size = Pt(14)
	ack_style.font.bold = False
	ack_style.font.color.rgb = RGBColor(0,0,0) #from_string('000')
	# paragraph format
	ack_style.paragraph_format.space_before = Pt(0)
	ack_style.paragraph_format.space_after = Pt(6)
	ack_style.paragraph_format.line_spacing = 1.5
	ack_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
	ack_style.paragraph_format.keep_together = True
	ack_style.paragraph_format.keep_with_next = False
	
	# logo for back page
	logo_style = styles.add_style('GB Logo', WD_STYLE_TYPE.PARAGRAPH)
	logo_style.base_style = styles['Normal']
	# font
	logo_style.font.name = 'Helvetica' 
	logo_style.font.size = Pt(72)
	logo_style.font.bold = True
	logo_style.font.color.rgb = RGBColor(131,3,3) #from_string('000')
	# paragraph format
	logo_style.paragraph_format.space_before = Pt(0)
	logo_style.paragraph_format.space_after = Pt(6)
	logo_style.paragraph_format.line_spacing = 1
	logo_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
	logo_style.paragraph_format.keep_together = True
	logo_style.paragraph_format.keep_with_next = False
	
	# back page text
	back_page_style = styles.add_style('GB Back Page', WD_STYLE_TYPE.PARAGRAPH)
	back_page_style.base_style = styles['Normal']
	# font
	back_page_style.font.name = 'Helvetica' 
	back_page_style.font.size = Pt(14)
	back_page_style.font.bold = False
	back_page_style.font.color.rgb = RGBColor(0,0,0) #from_string('000')
	# paragraph format
	back_page_style.paragraph_format.space_before = Pt(0)
	back_page_style.paragraph_format.space_after = Pt(6)
	back_page_style.paragraph_format.line_spacing = 1
	back_page_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
	back_page_style.paragraph_format.keep_together = True
	back_page_style.paragraph_format.keep_with_next = False
	
	# intro
	intro_style = styles.add_style('GB Intro', WD_STYLE_TYPE.PARAGRAPH)
	intro_style.base_style = styles['Normal']
	# font
	intro_style.font.name = 'Helvetica' 
	intro_style.font.size = Pt(12)
	intro_style.font.bold = False
	intro_style.font.color.rgb = RGBColor(0,0,0) #from_string('000')
	# paragraph format
	intro_style.paragraph_format.space_before = Pt(0)
	intro_style.paragraph_format.space_after = Pt(0)
	intro_style.paragraph_format.line_spacing = 1.5
	intro_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
	intro_style.paragraph_format.keep_together = True
	intro_style.paragraph_format.keep_with_next = True
	
	# feedback
	fdbk_style = styles.add_style('GB Feedback', WD_STYLE_TYPE.PARAGRAPH)
	fdbk_style.base_style = styles['Normal']
	# font
	fdbk_style.font.name = 'Helvetica' 
	fdbk_style.font.size = Pt(18)
	fdbk_style.font.bold = True
	fdbk_style.font.color.rgb = RGBColor(0,0,0) #from_string('000')
	# paragraph format
	fdbk_style.paragraph_format.space_before = Pt(0)
	fdbk_style.paragraph_format.space_after = Pt(12)
	fdbk_style.paragraph_format.line_spacing = 1.5
	fdbk_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
	fdbk_style.paragraph_format.keep_together = True
	fdbk_style.paragraph_format.keep_with_next = True
	
	
def prevent_document_break(document):
    """https://github.com/python-openxml/python-docx/issues/245#event-621236139
       Globally prevent table cells from splitting across pages.
    """
    tags = document.element.xpath('//w:tr')
    rows = len(tags)
    for row in range(0, rows):
        tag = tags[row]  # Specify which <w:r> tag you want
        child = OxmlElement('w:cantSplit')  # Create arbitrary tag
        tag.append(child)  # Append in the new tag

# set the page margins
def set_page_margins(sections):
	page_margin = 1
	for section in sections:
		section.top_margin = Inches(page_margin)
		section.bottom_margin = Inches(page_margin)
		section.left_margin = Inches(page_margin)
		section.right_margin = Inches(page_margin)

# generate an entire book (eg art of war), one page at a time
def generate_book_pages(book_title):

	print("\n=== Generate Book Pages: " + book_title + " ===\n")

	document = Document()
	
	generate_document_styles(document, book_title) # generate styles, set style properties
	
	page_sections = document.sections
	set_page_margins(page_sections)
	
	# generate book pages by section
	
	# cover image
	#if book_title == 'gb':
	if re.search("gb",book_title):
	
		pic_p = document.add_paragraph()
		cover_pic = pic_p.add_run().add_picture("../data/images/GB-cover-large.gif",width=Inches(8.5),height=Inches(11)) # alt img GB-cover.tif
	
		# change cover image margins
		cover_pic_sect = document.add_section(start_type=WD_SECTION.CONTINUOUS) 
	
	generator.generate_book_cover_page(document, book_title)
	
	# dedication page
	if re.search("gb",book_title):
		ded_string = "This book is dedicated to everyone who bought it."
		ded_paragraph = document.add_paragraph(ded_string)
		ded_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		document.add_section(start_type=WD_SECTION.NEW_PAGE) 
	
	generator.generate_book_toc(document, book_title)
	
	# intro page
	if re.search('gb',book_title): # 'gb' and 'aw gb'
		generator.generate_book_intro(document, book_title)
	
	# prepare to read raw principles where each column is a different book
	data_type = 'gb principles'
	input = 'Raw Data'
	extension = 'tsv'

	# use if we need data from spreadsheet, rather than from word docx with index entries pre-added
	raw_principle_data = reader.extract_data(data_type, input, extension)
	#print('raw_principle_data: ' + str(raw_principle_data))
	#list_title = "raw_principle_data"
	#writer.display_list_with_title(raw_principle_data, list_title)
	
	# isolate data for a single book, determined by book_title
	# if the book_title string contains multiple book titles, then separate the book titles and gather data for each given book title
	table_title = data_type
	if re.search(",",book_title): # if you find the comma split by the comma
		print("User input multiple book titles.")
		all_input_book_titles = book_title.split(",")
		
		all_input_book_data = []
		for input_book_title in all_input_book_titles:
			
			field_title = input_book_title # eg 'aw'
			input_book_data = isolator.isolate_data_field(raw_principle_data, table_title, field_title)
			all_input_book_data.append(input_book_data)
			
		print("all_input_book_data: " + str(all_input_book_data)) # format [ ['b1p1','b1p2'],['b2p1','b2p2'] ]
			
		all_book_chapters = []
		for input_book_data in all_input_book_data:
			book_chapters = isolator.isolate_whole_chapters(input_book_data) # whereas isolate_chapters starts new chapt when encounters "1.", isolate_whole_chapters starts new chapt when encounters chapter title line with format ^\d\.\s+\w+$
			all_book_chapters.append(book_chapters)
			
		print("all_book_chapters: " + str(all_book_chapters)) # format [ [ ['b1c1p1','b1c1p2'],['b1c2p1','b1c2p2'] ], [ ['b2c1p1','b2c1p2'],['b2c2p1','b2c2p2'] ] ]
			
		all_books_all_ch_data = []
		all_books_all_ch_paragraphs = []
		
		for book_idx in range(len(all_book_chapters)):
			book_chapters = all_book_chapters[book_idx]
			print("book_chapters: " + str(book_chapters))
			# isolate chapter data. all chapter raw data. 
			all_ch_overviews = isolator.isolate_all_ch_overviews(book_chapters)
			print("all_ch_overviews: " + str(all_ch_overviews))
			all_ch_principles = isolator.isolate_all_ch_principles(book_chapters) # isolate principle data in chapter
			
			all_ch_titles = isolator.isolate_all_ch_titles(book_chapters)
			all_ch_title_paragraphs = isolator.isolate_all_ch_title_paragraphs(book_title, book_idx)
			if len(all_ch_titles) > 0:
				writer.display_all_ch_titles(all_ch_titles, all_ch_title_paragraphs) # we do not need to assign ch titles from docx bc we have them already from spreadsheet
			else:
				all_ch_titles = writer.display_all_ch_titles(all_ch_titles, all_ch_title_paragraphs)
			print("all_ch_titles: " + str(all_ch_titles))

			all_ch_paragraphs = isolator.isolate_all_ch_paragraphs(book_title, book_idx) # take from input word docx (eg AW-input.docx)
			#writer.display_all_ch_paragraphs(all_ch_paragraphs)

			all_ch_data = generator.generate_all_ch_valid_data(all_ch_titles, all_ch_overviews, all_ch_principles) #[ch1_data]
			
			all_books_all_ch_data.append(all_ch_data)
			all_books_all_ch_paragraphs.append(all_ch_paragraphs)
			
		generate_comparison_chapters(document, book_title, all_books_all_ch_data, all_books_all_ch_paragraphs)
		
	else:
		field_title = book_title # eg 'aw'
		book_data = isolator.isolate_data_field(raw_principle_data, table_title, field_title) #eg aw_data; all_principle_data[aw_principle_idx]
		#writer.display_list_with_title(aw_data, field_title)
	
		#aw_data = ['1. Planning','O1','1. Sun Tzu said: The art of war is of vital importance to the State.','2. Waging War']
	
		book_chapters = isolator.isolate_whole_chapters(book_data) # whereas isolate_chapters starts new chapt when encounters "1.", isolate_whole_chapters starts new chapt when encounters chapter title line with format ^\d\.\s+\w+$
	
		# isolate chapter data. all chapter raw data. 
		all_ch_overviews = isolator.isolate_all_ch_overviews(book_chapters)
		all_ch_principles = isolator.isolate_all_ch_principles(book_chapters) # isolate principle data in chapter
	
		all_ch_titles = isolator.isolate_all_ch_titles(book_chapters)
		all_ch_title_paragraphs = isolator.isolate_all_ch_title_paragraphs(book_title)
		if len(all_ch_titles) > 0:
			writer.display_all_ch_titles(all_ch_titles, all_ch_title_paragraphs) # we do not need to assign ch titles from docx bc we have them already from spreadsheet
		else:
			all_ch_titles = writer.display_all_ch_titles(all_ch_titles, all_ch_title_paragraphs)


		all_ch_paragraphs = isolator.isolate_all_ch_paragraphs(book_title) # take from input word docx (eg AW-input.docx)
		writer.display_all_ch_paragraphs(all_ch_paragraphs)

		all_ch_data = generator.generate_all_ch_valid_data(all_ch_titles, all_ch_overviews, all_ch_principles) #[ch1_data]

		generate_book_chapters(document, book_title, all_ch_paragraphs, all_ch_data)
		
		
	# conclusion pages
	if re.search('gb',book_title): # 'gb' and 'aw gb'
		generator.generate_book_conclusion(document, book_title)
		
	# generate appendix
	generator.generate_principles_appendix(document, book_title)
	
	generator.generate_principles_references(document, book_title)
	
	generator.generate_principles_index(document, book_title)
	
	generator.generate_principles_acknowledgements(document, book_title)
	
	generator.generate_principles_back_page(document, book_title)
	
	prevent_document_break(document) # keep table row content together on the same page
	
	writer.display_page_numbers(document, book_title) # add page numbers
	
	# change cover image margins
	if re.search('gb',book_title):
		cover_pic_sect = document.sections[0] 
		cover_pic_sect.top_margin = 0
		cover_pic_sect.left_margin = 0
	
	#document.save('../test.docx')
	book_title = re.sub(",","-",book_title)
	out_filename = '../../../Troob/Game-of-Business/' + book_title.upper() + '-drafts/' + book_title.upper() + '-test.docx'
	document.save(out_filename)

# user input
book_title = 'aw,gb' # 'aw' or 'gb' or 'aw,gb'
	
generate_book_pages(book_title)
