# writer.py
# functions for a writer

import re
import generator, reader, isolator, determiner

# display word docx
from docx import Document
#set cell border
from docx.table import _Cell
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH


# order of detail fields
sku_idx = 0
handle_idx = 1 # old
collection_idx = 1
title_idx = 2
intro_idx = 3
color_idx = 4
mat_idx = 5
finish_idx = 6
width_idx = 7
depth_idx = 8
height_idx = 9
weight_idx = 10
features_idx = 11
cost_idx = 12
img_src_idx = 13
barcode_idx = 14
gen_handle_idx = 15

# === Helper Functions ===

def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
 
    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
 
    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
 
            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
 
            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))
                    
def set_cell_margins(cell: _Cell, **kwargs):
    """
    cell:  actual cell instance you want to modify

    usage:

        set_cell_margins(cell, top=50, start=50, bottom=50, end=50) if we want x=1", then y=1440, so to get x=1/4", set y=1440/4=360

    provided values are in twentieths of a point (1/1440 of an inch).
    read more here: http://officeopenxml.com/WPtableCellMargins.php
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')

    for m in [
        "top",
        "start",
        "bottom",
        "end",
    ]:
        if m in kwargs:
            node = OxmlElement("w:{}".format(m))
            node.set(qn('w:w'), str(kwargs.get(m)))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)

    tcPr.append(tcMar)
    
def remove_unwanted_paragraphs(outer_table):

	# remove unwanted empty paragraph lines
	outer_table_element = outer_table._element # containing whole sub-principle, eg 2.4
	#print("outer_table_element: " + str(outer_table_element))
	outer_table_children = outer_table_element.getchildren()
	#print("outer_table_children: " + str(outer_table_children))
	num_table_children = len(outer_table_children)
	outer_table_row = outer_table_children[num_table_children-1]
	#print("outer_table_row: " + str(outer_table_row))
	outer_table_row_children = outer_table_row.getchildren()
	#print("outer_table_row_children: " + str(outer_table_row_children))
	outer_table_content = outer_table_row_children[0]
	outer_table_content_children = outer_table_content.getchildren()
	#print("outer_table_content_children: " + str(outer_table_content_children))
	# loop thru odd numbered indexes to remove unwanted paragraphs, except <p> at last index (last child)
	num_content_children = len(outer_table_content_children)
	#print("num_content_children: " + str(num_content_children))
	for child_idx in range(num_content_children - 2): # valid range bc last <p> is needed
		if child_idx % 2 == 1:
			#print("child_idx: " + str(child_idx))
			outer_table_unwanted_pg = outer_table_content_children[child_idx]
			#print("outer_table_unwanted_pg: " + str(outer_table_unwanted_pg))
			outer_table_unwanted_pg.getparent().remove(outer_table_unwanted_pg)
			outer_table_unwanted_pg._p = outer_table_unwanted_pg._element = None
    
    
    
    

def display_field_values(values):
	print("=== Display Field Values ===")
	for value in values:
		print(value)
	print()
	
def display_valid_data_headers(vendor):
	data_type = "valid data"
	value_type = "field"
	valid_data_fields = reader.read_json(data_type, value_type)
	product_variable_names = valid_data_fields[vendor]
	product_headers = ''
	for name_idx in range(len(product_variable_names)):
		name = product_variable_names[name_idx]
		if name_idx == 0:
			product_headers += name
		else:
			product_headers += ";" + name
	print(product_headers)
	
def display_catalog_headers():
	catalog_variable_names = ["SKU","Collection","Title","Intro","Colors","Materials","Finish","Width","Depth","Height","Weight","Features","Cost","Image Link","Barcode"]
	catalog_headers = ''
	for name_idx in range(len(catalog_variable_names)):
		name = catalog_variable_names[name_idx]
		if name_idx == 0:
			catalog_headers += name
		else:
			catalog_headers += ";" + name
	print(catalog_headers)

def display_shopify_variant_headers():
	product_variable_names = ["Ref Num","Variant SKU","Handle","Variant Weight","Variant Barcode","Body HTML","Option1 Name","Option1 Value","Option2 Name","Option2 Value","Option3 Name","Option3 Value","Tags","Image Src","Type","Title","Published","Published Scope","Variant Inventory Tracker","Variant Inventory Policy","Variant Weight Unit","Command","Vendor","Variant Cost","Variant Price","Variant Compare At Price","Template Suffix"]
	product_headers = ''
	for name_idx in range(len(product_variable_names)):
		name = product_variable_names[name_idx]
		if name_idx == 0:
			product_headers += name
		else:
			product_headers += ";" + name
	print(product_headers)

def display_zoho_item_headers():
	product_variable_names = ["SKU","Item Name","Package Width","Package Depth","Package Height","Package Weight","Type","Reference Number","Purchase Account","Reason","Vendor","Preferred Vendor","Date","Warehouse Name","Description","Quantity Adjusted","New Quantity On Hand","Unit","Commission","Cost Price","Selling Price","Sales Account","Inventory Account","Item Type"]
	product_headers = ''
	for name_idx in range(len(product_variable_names)):
		name = product_variable_names[name_idx]
		if name_idx == 0:
			product_headers += name
		else:
			product_headers += ";" + name
	print(product_headers)
	
def display_zoho_composite_headers():
	product_variable_names = ["SKU","Composite Item Name","Package Width","Package Depth","Package Height","Package Weight","Type","Reference Number","Purchase Account","Reason","Vendor","Preferred Vendor","Date","Warehouse Name","Description","Quantity Adjusted","New Quantity On Hand","Unit","Commission","Cost Price","Selling Price","Sales Account","Inventory Account","Item Type","Mapped Item Name","Mapped Quantity","Mapped Item SKU"]
	product_headers = ''
	for name_idx in range(len(product_variable_names)):
		name = product_variable_names[name_idx]
		if name_idx == 0:
			product_headers += name
		else:
			product_headers += ";" + name
	print(product_headers)

# replace commas with semicolons for shopify import and add commas if needed to allow space for 3 options
def format_option_string(init_option_string):
	option_string = ''

	option_string = re.sub(',',';',init_option_string)
	product_option_data = option_string.split(";")
	# make sure always 3 option spaces, even if less than 3 options
	max_options = 3 # limited by shopify
	num_option_fields = 2 * max_options # name and value for each option
	# if we already have maximum number of options, do nothing
	num_defined_option_fields = len(product_option_data) # number of fields filled in with either name or value
	if num_defined_option_fields != num_option_fields:
		# add spaces for options, so count defined options so we know difference
		num_blank_option_fields = num_option_fields - num_defined_option_fields
		for blank_option in range(num_blank_option_fields):
			option_string += ';' # original maybe "size;king" so final would be "size;king;;;;"

	return option_string
	
# given dims in format W33" x D37" x H35.00"
# convert to 33;37;35 for catalog gen
def format_dim_string(init_dim_string):
	#print("\n=== Format Dim String ===\n")
	
	# remove ", W, D, H 
	dim_string = re.sub("\"|W|D|H","",init_dim_string)
	
	# split data by ' x '
	dim_data = dim_string.split(' x ')
	
	final_dim_string = ''
	# add semicolons
	for dim_idx in range(len(dim_data)):
		dim = dim_data[dim_idx]
		
		if dim_idx == 0:
			final_dim_string += dim
		else:
			final_dim_string += ";" + dim
		
	return final_dim_string
			

def display_all_item_details(all_dtls):
	print("\n=== Item Details ===\n")
	for item_details in all_dtls:
		#handle = item_details[handle_idx].strip().lower()
		handle = generator.generate_handle(item_details)

		print(handle + ": " + str(item_details))
	print()
	
def display_list(list):
	print("\n=== List ===\n")
	for element in list:
		print(element)
		
def display_list_with_title(list, title):
	print("\n=== " + title.title() + " ===\n")
	if len(list) > 0:
		for element in list:
			if element != '':
				print(element)
	else:
		print("\nWarning: Empty List!\n")
		
def format_title_for_dim_fmla(init_title, coll_name):
	title = re.sub(coll_name.title(),"",init_title)
	
	# remove parentheses bc in context of dims
	all_parentheses = ["(1 Stool Only)","(1 Chair Only)","(2 Pack)","(5 Piece)","(7 Piece)"]
	for parenthesized in all_parentheses:
		title = re.sub(parenthesized,"",title)
		title = re.sub("\(|\)","",title)
		
	title = title.strip()
	print("title: \"" + title + "\"")
	
	return title
	
# ====== Principle Generator ======

principle_num_idx = 0
principle_content_idx = 1
	
# principle elements are the elemental html code blocks used to show the principles on the webpage, including all html tags used for functionality such as dropdowns revealing more info
def display_principle_elements(desired_principles, principle_ids=[], demo_principles=[]):

	all_principle_elements = ''
	
	# init content bt <table> tags, which includes all table rows
	table_content = ''

	for principle_idx in range(len(desired_principles)):

		principle = desired_principles[principle_idx]
	
		print("principle: " + principle)

		principle_data = isolator.isolate_principle_num_and_content(principle) #re.split("(?<=\d\.)\s",principle) # separate/isolate princ num and princ content
		print("principle_data: " + str(principle_data))

		if len(principle_data) > 0: 
	
			principle_num = principle_data[principle_num_idx] #'1.'
		
			principle_id = principle_ids[principle_idx]
			#print("principle_id: " + principle_id)

			principle_class = determiner.determine_principle_class(principle_num) #'gob-preview-principle'

			principle_content = generator.generate_principle_content(principle_data, principle_content_idx, demo_principles, principle_id);
			
		#<table>
		#<tr class="gob-preview-principle">
		#<td>1. </td>
		#<td>Sun Zi said: The Game of Business is of vital importance to the Company. </td>
		#</tr>
		#</table>

		table_content += '\n\t<tr id=\"' + principle_id + '\" class=\"' + principle_class + '\">\n\t\t<td>' + principle_num + ' </td>\n\t\t<td>' + principle_content + ' </td>\n\t</tr>'

	preview = True
	table_class = ''
	if preview:
		table_class = "gob-preview"

	all_principle_elements = '<table class=\"' + table_class + '\">' + table_content + '\n</table>'

	#print('all_principle_elements:\n' + all_principle_elements)
	
	return all_principle_elements
	
def display_principle_in_table(document, table, princ_num, ch_princ_data, ch_paragraphs=[]):
	
	#print("\n=== Display Principle " + princ_num + " in Table ===\n")
	
	# princ_num_cell = table.cell(0,0)
# 	princ_content_cell = table.cell(0,1)
# 	princ_num_cell.text = princ_num
# 	princ_content_cell.text = 'Principle content. '
# 	princ_num_cell.width = Inches(0.75) # arbitrary, based on font size and maximum content length. must align with other table columns
# 	princ_content_cell.width = Inches(5.75) # 6.5-princ_num_cell.width

	# Format Table Cells
	princ_num_cell = table.cell(0,0)
	princ_content_cell = table.cell(0,1)
	
	princ_num_cell.width = Inches(0.95)
	princ_content_cell.width = Inches(5.55)
	
	set_cell_border(
		princ_num_cell,
		top={},
		bottom={},
		start={},
		end={"sz": 18, "color": "#830303", "val": "single"},
	)

	if princ_num.count('.') == 3: # if sub-sub-principle
		#print("Principle " + princ_num + " is a sub-sub-principle!")
		set_cell_margins(princ_content_cell, start=1080) # 1/1440 of an inch
	elif princ_num.count('.') == 2:# if sub-principle
		#print("Principle " + princ_num + " is a sub-principle!")
		set_cell_margins(princ_content_cell, start=720) # 1/1440 of an inch
	else: # if main principle
		#print("Principle " + princ_num + " is a main principle!")
		set_cell_margins(princ_content_cell, start=360) # 1/1440 of an inch

	# write cell text
	# ensure princ num ends with dot
	princ_num = reader.add_rear_dot(princ_num)
	#print("final_princ_num: " + princ_num)
	princ_num_cell.text = princ_num # write princ num
	
	if len(ch_paragraphs) == 0:
		print("Length of ch_paragraphs = 0 so populate principle content from Raw Data.")
		
		princ_content = isolator.isolate_principle_content(princ_num, ch_princ_data)
		princ_content_cell.text = princ_content
	else:
		# get princ content paragraph from existing doc with index entries
		princ_content_p = isolator.isolate_principle_paragraph(princ_num, ch_princ_data, ch_paragraphs)
		
		#princ_content_p = generator.generate_runs_from_paragraph(princ_content_p, keywords) # gen runs to italicize keywords. modify existing paragraph that may have index entry runs mixed in with text runs
		
		inserted_p = document._body._body._insert_p(princ_content_p._p) # directly insert a paragraph from a src doc into a dest doc
		#print("inserted_p: " + str(inserted_p))
		
		princ_content_cell._element.addprevious(inserted_p)
		#print("princ_content_cell paragraphs: " + str(princ_content_cell.paragraphs))
		cell_unwanted_pg = princ_content_cell.paragraphs[0]._element
		cell_unwanted_pg.getparent().remove(cell_unwanted_pg)
		
		
		
# need to create separate table for each sub-principle set bc rows with sets of sub-principles start with one column and then get nested table
def display_principle_tables(document, group, ch_princ_data, ch_paragraphs=[]):

	#print("\n=== Diplay Principle Tables ===\n")
	#print("ch_paragraphs: " + str(ch_paragraphs))

	num_rows = 1

	all_sub_principles = isolator.isolate_sub_principles(group, ch_princ_data) # iso sub-principles to put in same row, so they stay on same page
	#print("all_sub_principles: " + str(all_sub_principles))
	
	for sub_princ_set in all_sub_principles:
		if len(sub_princ_set) == 1: # only 1 principle in the set
			num_cols = 2
			princ_table = document.add_table(rows=num_rows,cols=num_cols)
			princ_table.autofit = False
			princ_table.allow_autofit = False
			princ_table.alignment = WD_TABLE_ALIGNMENT.CENTER
			princ_num = sub_princ_set[0]
			display_principle_in_table(document, princ_table, princ_num, ch_princ_data, ch_paragraphs)
		
		# create outer table that will contain the nested table
		# if the current principle is a sub-principle that contains sub-sub-principles, use a nested table to keep the whole sub-principle, including sub-sub-principles, on the same page
		elif len(sub_princ_set) > 1:
			num_cols = 1
			outer_table = document.add_table(rows=num_rows,cols=num_cols) # always 1 row with 1 cell containing whole sub-principle, eg 2.4
			outer_table.autofit = False
			outer_table.allow_autofit = False
			outer_table.alignment = WD_TABLE_ALIGNMENT.CENTER

			outer_table_cell = outer_table.cell(0,0)
			outer_table_cell.width = Inches(6.5)

			num_cols = 2

			for princ_num in sub_princ_set:
				nested_table = outer_table_cell.add_table(rows=num_rows,cols=num_cols)
				nested_table.allow_autofit = False
				nested_table.alignment = WD_TABLE_ALIGNMENT.CENTER
				display_principle_in_table(document, nested_table, princ_num, ch_princ_data, ch_paragraphs)

			remove_unwanted_paragraphs(outer_table)
			
def display_comparison_principles_in_table(document, table, princ_num, ch_all_book_princ_data, ch_all_book_paragraphs=[]):
	
	print("\n=== Display Comparison Principles " + princ_num + " in Table ===\n")
	
	# princ_num_cell = table.cell(0,0)
# 	princ_content_cell = table.cell(0,1)
# 	princ_num_cell.text = princ_num
# 	princ_content_cell.text = 'Principle content. '
# 	princ_num_cell.width = Inches(0.75) # arbitrary, based on font size and maximum content length. must align with other table columns
# 	princ_content_cell.width = Inches(5.75) # 6.5-princ_num_cell.width

	# Format Table Cells
	princ_num_cell = table.cell(0,0)
	
	princ_content_cells = []
	for book_idx in range(len(ch_all_book_princ_data)):
		#print("book_idx: " + str(book_idx))
		col_idx=book_idx+1 #offset 1 for princ num col
		princ_content_cell = table.cell(0,col_idx)
		princ_content_cells.append(princ_content_cell)
	
	# set cell width
	princ_num_cell_width = 0.95
	princ_num_cell.width = Inches(princ_num_cell_width)
	all_princ_content_width = 6.5 - princ_num_cell_width
	remaining_width = all_princ_content_width
	princ_content_width = all_princ_content_width / len(princ_content_cells)
	for cell_idx in range(len(princ_content_cells)):
		#print("remaining_width: " + str(remaining_width))
		#print("princ_content_width: " + str(princ_content_width))
	
		cell = princ_content_cells[cell_idx]
		
		if cell_idx == len(princ_content_cells)-1:
			cell.width = Inches(remaining_width)
		else:
			cell.width = Inches(princ_content_width)
			
			remaining_width -= princ_content_width
	
	set_cell_border(
		princ_num_cell,
		top={},
		bottom={},
		start={},
		end={"sz": 18, "color": "#830303", "val": "single"},
	)
	
	set_cell_border(
		princ_content_cells[0],
		top={},
		bottom={},
		start={},
		end={"sz": 9, "color": "#830303", "val": "dashed"},
	)

	for cell_idx in range(len(princ_content_cells)):
		cell = princ_content_cells[cell_idx]
		
		start_margin = 360
		end_margin = 180
		
		if princ_num.count('.') == 3: # if sub-sub-principle
			#print("Principle " + princ_num + " is a sub-sub-principle!")
			set_cell_margins(cell, start=1080, end=end_margin) # 1/1440 of an inch
		elif princ_num.count('.') == 2:# if sub-principle
			#print("Principle " + princ_num + " is a sub-principle!")
			set_cell_margins(cell, start=720, end=end_margin) # 1/1440 of an inch
		else: # if main principle
			#print("Principle " + princ_num + " is a main principle!")
			set_cell_margins(cell, start=start_margin, end=end_margin) # 1/1440 of an inch
				

	# write cell text
	# ensure princ num ends with dot
	princ_num = reader.add_rear_dot(princ_num)
	#print("final_princ_num: " + princ_num)
	princ_num_cell.text = princ_num # write princ num
	
	if len(ch_all_book_paragraphs) == 0:
		print("Length of ch_all_book_paragraphs = 0 so populate principle content from Raw Data.")
		
		for cell_idx in range(len(princ_content_cells)):
			cell = princ_content_cells[cell_idx]
			
			princ_content = isolator.isolate_principle_content(princ_num, ch_all_book_princ_data[cell_idx])
			cell.text = princ_content
	else:
		for cell_idx in range(len(princ_content_cells)):
			cell = princ_content_cells[cell_idx]
			#print("cell_idx: " + str(cell_idx))
			
			# get princ content paragraph from existing doc with index entries
			princ_content_p = isolator.isolate_principle_paragraph(princ_num, ch_all_book_princ_data[cell_idx], ch_all_book_paragraphs[cell_idx])
		
			#princ_content_p = generator.generate_runs_from_paragraph(princ_content_p, keywords) # gen runs to italicize keywords. modify existing paragraph that may have index entry runs mixed in with text runs
		
			inserted_p = document._body._body._insert_p(princ_content_p._p) # directly insert a paragraph from a src doc into a dest doc
			#print("inserted_p: " + str(inserted_p))
		
			cell._element.addprevious(inserted_p)
			#print("princ_content_cell paragraphs: " + str(princ_content_cell.paragraphs))
			cell_unwanted_pg = cell.paragraphs[0]._element
			cell_unwanted_pg.getparent().remove(cell_unwanted_pg)
			
# need to create separate table for each sub-principle set bc rows with sets of sub-principles start with one column and then get nested table
def display_comparison_principle_tables(document, group, ch_all_book_princ_data, ch_all_book_paragraphs=[]):

	#print("\n=== Diplay Comparison Principle Tables ===\n")
	#print("ch_paragraphs: " + str(ch_paragraphs))

	num_rows = 1 # each row is new table for formatting purposes

	all_sub_principles = isolator.isolate_sub_principles(group, ch_all_book_princ_data[0]) # iso sub-principles to put in same row, so they stay on same page. we can take the first ch_princ_data bc used to check if contains "the following"
	#print("all_sub_principles: " + str(all_sub_principles))
	
	# do not use nested tables for comparisons bc the nested table cannot fit on one page bc the proportions of text and space do not allow it
	for sub_princ_set in all_sub_principles:
	
		for princ_num in sub_princ_set:
			num_cols = len(ch_all_book_princ_data) + 1 # no. books + 1 for princ num col
			princ_table = document.add_table(rows=num_rows,cols=num_cols)
			princ_table.autofit = False
			princ_table.allow_autofit = False
			princ_table.alignment = WD_TABLE_ALIGNMENT.CENTER
			#princ_num = sub_princ_set[0] # sub_princ_set format ["1."]
			display_comparison_principles_in_table(document, princ_table, princ_num, ch_all_book_princ_data, ch_all_book_paragraphs)
	
		# if len(sub_princ_set) == 1: # only 1 principle in the set
# 			num_cols = len(ch_all_book_princ_data) + 1 # no. books + 1 for princ num col
# 			princ_table = document.add_table(rows=num_rows,cols=num_cols)
# 			princ_table.autofit = False
# 			princ_table.allow_autofit = False
# 			princ_table.alignment = WD_TABLE_ALIGNMENT.CENTER
# 			princ_num = sub_princ_set[0] # sub_princ_set format ["1."]
# 			display_comparison_principles_in_table(document, princ_table, princ_num, ch_all_book_princ_data, ch_all_book_paragraphs)
# 		
# 		# create outer table that will contain the nested table
# 		# if the current principle is a sub-principle that contains sub-sub-principles, use a nested table to keep the whole sub-principle, including sub-sub-principles, on the same page
# 		elif len(sub_princ_set) > 1:
# 			num_cols = 1
# 			outer_table = document.add_table(rows=num_rows,cols=num_cols) # always 1 row with 1 cell containing whole sub-principle, eg 2.4
# 			outer_table.autofit = False
# 			outer_table.allow_autofit = False
# 			outer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
# 
# 			outer_table_cell = outer_table.cell(0,0)
# 			outer_table_cell.width = Inches(6.5)
# 
# 			num_cols = len(ch_all_book_princ_data) + 1 # no. books + 1 for princ num col
# 
# 			for princ_num in sub_princ_set:
# 				nested_table = outer_table_cell.add_table(rows=num_rows,cols=num_cols)
# 				nested_table.allow_autofit = False
# 				nested_table.alignment = WD_TABLE_ALIGNMENT.CENTER
# 				display_comparison_principles_in_table(document, nested_table, princ_num, ch_all_book_princ_data, ch_all_book_paragraphs)
# 
# 			remove_unwanted_paragraphs(outer_table)
			
def display_nested_table_row(table, row, row_idx, num_total_rows, document, table_type='appendix'):

	#print("\n=== Display Nested Table Row " + str(row_idx) + " ===\n")
	
	#print("row: " + str(row))

	# Format Table Cells
	cells = []
	for col_idx in range(len(row)):
		cell = table.cell(0,col_idx)
		cells.append(cell)

	# write cell text
	for col_idx in range(len(row)):
		cell_content = row[col_idx]
		cell = cells[col_idx]
		cell.text = cell_content
		
	# what we do to the paragraph within the cell must be done after setting cell text
	if table_type != 'ordered list':
		for cell in cells:
			paragraph = cell.paragraphs[0]
			#print("paragraph: " + paragraph.text)
		
			# set paragraph styles
			if row_idx == 0:
				paragraph.style = document.styles['Field Title']
			else: 
				paragraph.style = document.styles['Table Data']
				#print('set appendix table style: ' + str(paragraph.style))
			
			# set cell borders
			if row_idx == 0:
				set_cell_border(
					cell,
					top={"sz": 18, "color": "#830303", "val": "single"},
					bottom={"sz": 9, "color": "#830303", "val": "single"},
					start={},
					end={},
				)
			elif row_idx == num_total_rows-1:
				#print("last row")
				set_cell_border(
					cell,
					top={"sz": 9, "color": "#830303", "val": "dashed"},
					bottom={"sz": 18, "color": "#830303", "val": "single"},
					start={},
					end={},
				)
			else: # middle row 
				set_cell_border(
					cell,
					top={},
					bottom={"sz": 9, "color": "#830303", "val": "dashed"},
					start={},
					end={},
				)
			
			# set cell margins, top and bottom bc too crowded next to border
			set_cell_margins(cell, top=115, bottom=115) # 1/1440 of an inch
		
	# if it is a table in the conclusion, the first col is the principle number which takes up little space
	if table_type == 'conclusion':
		for col_idx in range(len(cells)):
			cell = cells[col_idx]
			# assuming 3 cols
			# must add up to 6.5
			if col_idx == 0:
				cell.width = Inches(0.5)
			elif col_idx == 1:
				cell.width = Inches(2.5)
			elif col_idx == 2:
				cell.width = Inches(3.5)
				
	if table_type == 'ordered list':
		
		id_cell = cells[0]
		part_cell = cells[1]
	
		id_cell.width = Inches(0.5)
		part_cell.width = Inches(6)
		
		id = row[0]
		# if id has two dots then change margin for sub-bullet
		if id.count(".") == 2:
			set_cell_margins(part_cell, start=720)
			
		
def display_nested_table(document, nested_table_data, caption='', table_type='appendix'):

	#print("\n=== Diplay Nested Table ===\n")
	
	if table_type == 'ordered list':
		final_table_data = []
		ids = nested_table_data[0]
		parts = nested_table_data[1]
		for idx in range(len(ids)):
			id = ids[idx]
			part = parts[idx]
			row = [id, part]
			final_table_data.append(row)
			
		nested_table_data = final_table_data
	
	#print("nested_table_data: " + str(nested_table_data))
	
	# display caption
	if caption != '':
		caption_paragraph = document.add_paragraph(caption,style='GB Table Caption')
	
	num_total_rows = len(nested_table_data)
	
	# create outer table that will contain the nested table
	# use nested tables for tables to keep them on the same page
	num_rows = 1
	num_cols = 1
	outer_table = document.add_table(rows=num_rows,cols=num_cols) # always 1 row with 1 cell containing whole sub-principle, eg 2.4
	outer_table.autofit = False
	outer_table.allow_autofit = False
	outer_table.alignment = WD_TABLE_ALIGNMENT.CENTER

	outer_table_cell = outer_table.cell(0,0)
	outer_table_cell.width = Inches(6.5)

	for row_idx in range(len(nested_table_data)):
		row = nested_table_data[row_idx]
		num_cols = len(row)
		
		nested_table = outer_table_cell.add_table(rows=num_rows,cols=num_cols)
		nested_table.allow_autofit = False
		nested_table.alignment = WD_TABLE_ALIGNMENT.CENTER
		display_nested_table_row(nested_table, row, row_idx, num_total_rows, document, table_type)

	# remove unwanted empty paragraph lines
	remove_unwanted_paragraphs(outer_table) # nested table
	
def display_appendix_table_row(table, row, row_idx, num_total_rows, document):

	#print("\n=== Display Appendix Table Row " + str(row_idx) + " ===\n")
	
	#print("row: " + str(row))

	# Format Table Cells
	cells = []
	for col_idx in range(len(row)):
		cell = table.cell(0,col_idx)
		cells.append(cell)

	# write cell text
	for col_idx in range(len(row)):
		cell_content = row[col_idx]
		cell = cells[col_idx]
		cell.text = cell_content
		
	# what we do to the paragraph within the cell must be done after setting cell text
	for cell in cells:
		paragraph = cell.paragraphs[0]
		#print("paragraph: " + paragraph.text)
		
		# set paragraph styles
		if row_idx == 0:
			paragraph.style = document.styles['Field Title']
		else: 
			paragraph.style = document.styles['Table Data']
			#print('set appendix table style: ' + str(paragraph.style))
			
		# set cell borders
		if row_idx == 0:
			set_cell_border(
				cell,
				top={"sz": 18, "color": "#830303", "val": "single"},
				bottom={"sz": 9, "color": "#830303", "val": "single"},
				start={},
				end={},
			)
		elif row_idx == num_total_rows-1:
			#print("last row")
			set_cell_border(
				cell,
				top={"sz": 9, "color": "#830303", "val": "dashed"},
				bottom={"sz": 18, "color": "#830303", "val": "single"},
				start={},
				end={},
			)
		else: # middle row 
			set_cell_border(
				cell,
				top={},
				bottom={"sz": 9, "color": "#830303", "val": "dashed"},
				start={},
				end={},
			)
			
		# set cell margins, top and bottom bc too crowded next to border
		set_cell_margins(cell, top=115, bottom=115) # 1/1440 of an inch
		
def display_appendix_table(document, appendix_table, caption=''):

	#print("\n=== Diplay Appendix Table ===\n")
	
	#print("appendix_table: " + str(appendix_table))
	
	# display caption
	if caption != '':
		caption_paragraph = document.add_paragraph(caption,style='GB Table Caption')
	
	num_total_rows = len(appendix_table)
	
	if num_total_rows < 3: # only 1 row is invalid for appendix tables bc they at least have title row and 2 content rows (3 rows)
		print("Warning: invalid appendix table!")
	
	# create outer table that will contain the nested table
	# use nested tables for appendix tables to keep them on the same page
	elif num_total_rows >= 3:
		num_rows = 1
		num_cols = 1
		outer_table = document.add_table(rows=num_rows,cols=num_cols) # always 1 row with 1 cell containing whole sub-principle, eg 2.4
		outer_table.autofit = False
		outer_table.allow_autofit = False
		outer_table.alignment = WD_TABLE_ALIGNMENT.CENTER

		outer_table_cell = outer_table.cell(0,0)
		outer_table_cell.width = Inches(6.5)

		for row_idx in range(len(appendix_table)):
			row = appendix_table[row_idx]
			num_cols = len(row)
			
			nested_table = outer_table_cell.add_table(rows=num_rows,cols=num_cols)
			nested_table.allow_autofit = False
			nested_table.alignment = WD_TABLE_ALIGNMENT.CENTER
			display_appendix_table_row(nested_table, row, row_idx, num_total_rows, document)

		# remove unwanted empty paragraph lines
		remove_unwanted_paragraphs(outer_table) # nested table
		
		#temp_table = document.add_table(0,0)
		#document.tables.remove(temp_table)
		#space_paragraph = document.add_paragraph(style='Appendix Table Data')
		#space_paragraph.style = document.styles['Appendix Table Data']
					
def display_all_appendix_tables(document, appendix_tables, ap_table_captions=[]):

	#print("\n=== Diplay Appendix Tables ===\n")
	
	for ap_table_idx in range(len(appendix_tables)):
		appendix_table = appendix_tables[ap_table_idx]
		ap_table_caption = ap_table_captions[ap_table_idx]
	
		display_appendix_table(document, appendix_table, ap_table_caption)
		
def display_all_comparison_appendix_tables(document, all_books_appendix_tables, all_books_ap_table_captions=[]):

	print("\n=== Diplay All Comparison Appendix Tables ===\n")
	
	for ap_table_idx in range(len(all_books_appendix_tables[0])): # display all ap tables with the same num together
		
		for book_idx in range(len(all_books_appendix_tables)):
		
			appendix_tables = all_books_appendix_tables[book_idx]
			appendix_table = appendix_tables[ap_table_idx]
			
			ap_table_captions = all_books_ap_table_captions[book_idx]
			print("ap_table_captions: " + str(ap_table_captions))
			ap_table_caption = ap_table_captions[ap_table_idx]
			print("ap_table_caption: " + str(ap_table_caption))
	
			display_appendix_table(document, appendix_table, ap_table_caption)
			
		document.add_page_break()

def create_element(name):
	return OxmlElement(name)
	
def create_attribute(element, name, value):
	element.set(ns.qn(name), value)
	
def format_book_title(title):

	book_title = ''
	
	if title.lower() == 'aw':
		book_title = 'The Art of War'
	elif re.search("gb",title.lower()):
		book_title = 'Game of Business'
	
	return book_title
	
def split_docx_keywords(content, keywords):

	#print("\n=== Split DOCX Keywords ===\n")

	split_content = [content]
	
	#print("content: \"" + content + "\"")

	# even if there is one key, perform split
	# we can loop thru keys to find matches and form search string
	key_search_string = ''
	found_key = False
	for key in keywords:
		# if we find key in content, add it to key string for delim later
		#print("key: " + key)
		if re.search(key, content):
			if found_key == False:
				key_string = "(" + key + ")" # first key to search
			else:
				key_string = "|(" + key + ")"
			
			key_search_string += key_string
			found_key = True
	
	if found_key == True:
	
		#print("found key!")
	
		split_content = []
	
		key_string = "(" + key + ")"
		# make string with all keywords as delimiters so we can split by all delims at once!
		content_data = re.split(key_search_string,content)
		#print("content_data: " + str(content_data))
		
		for data in content_data:
			if data != '' and data != None:
				split_content.append(data)
					
	#print("split_content: " + str(split_content))
	
	return split_content
	
def display_ch_paragraphs(ch):

	#print("\n=== Display Chapter Paragraphs ===\n")
	
	for ch_paragraph in ch:
		print(ch_paragraph.text)
	
def display_all_ch_paragraphs(all_ch_paragraphs):

	print("\n=== Display All Chapter Paragraphs ===\n")
	
	for ch in all_ch_paragraphs:
		print("New Chapter")
		for ch_paragraph in ch:
			print(ch_paragraph.text)
			
def display_all_ch_titles(all_ch_titles, all_ch_title_paragraphs=[]):

	print("\n=== Display All Chapter Titles ===\n")
	
	all_final_ch_titles = []
	
	# if given ch titles use that, else use the paragraphs from the word doc
	if len(all_ch_titles) > 0:
		print("Found chapter titles in spreadsheet!")
		for ch_title in all_ch_titles:
			print(ch_title)
			
		all_final_ch_titles = all_ch_titles
	elif len(all_ch_title_paragraphs) > 0:
		print("Found chapter titles in docx!")
		for ch_title_paragraph in all_ch_title_paragraphs:
			print("New Chapter")
			print(ch_title_paragraph.text)
			all_final_ch_titles.append(ch_title_paragraph.text)
				
	else:
		print("Warning: No chapter titles found!")
		
	return all_final_ch_titles
			
# def add_page_number(paragraph):
# 	paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
# 
# 	page_run = paragraph.add_run()
# 	t1 = create_element('w:t')
# 	create_attribute(t1, 'xml:space', 'preserve')
# 	t1.text = 'Page '
# 	page_run._r.append(t1)
# 
# 	page_num_run = paragraph.add_run()
# 
# 	fldChar1 = create_element('w:fldChar')
# 	create_attribute(fldChar1, 'w:fldCharType', 'begin')
# 
# 	instrText = create_element('w:instrText')
# 	create_attribute(instrText, 'xml:space', 'preserve')
# 	instrText.text = "PAGE"
# 
# 	fldChar2 = create_element('w:fldChar')
# 	create_attribute(fldChar2, 'w:fldCharType', 'end')
# 
# 	page_num_run._r.append(fldChar1)
# 	page_num_run._r.append(instrText)
# 	page_num_run._r.append(fldChar2)
# 
# 	of_run = paragraph.add_run()
# 	t2 = create_element('w:t')
# 	create_attribute(t2, 'xml:space', 'preserve')
# 	t2.text = ' of '
# 	of_run._r.append(t2)
# 
# 	fldChar3 = create_element('w:fldChar')
# 	create_attribute(fldChar3, 'w:fldCharType', 'begin')
# 
# 	instrText2 = create_element('w:instrText')
# 	create_attribute(instrText2, 'xml:space', 'preserve')
# 	instrText2.text = "NUMPAGES"
# 
# 	fldChar4 = create_element('w:fldChar')
# 	create_attribute(fldChar4, 'w:fldCharType', 'end')
# 
# 	num_pages_run = paragraph.add_run()
# 	num_pages_run._r.append(fldChar3)
# 	num_pages_run._r.append(instrText2)
# 	num_pages_run._r.append(fldChar4)
	
# def add_page_number(run):
#     fldStart = create_element('w:fldChar')
#     create_attribute(fldStart, 'w:fldCharType', 'begin')
# 
#     instrText = create_element('w:instrText')
#     create_attribute(instrText, 'xml:space', 'preserve')
#     instrText.text = "PAGE"
# 
#     fldChar1 = create_element('w:fldChar')
#     create_attribute(fldChar1, 'w:fldCharType', 'separate')
# 
#     fldChar2 = create_element('w:t')
#     fldChar2.text = "2"
# 
#     fldEnd = create_element('w:fldChar')
#     create_attribute(fldEnd, 'w:fldCharType', 'end')
# 
#     run._r.append(fldStart)
# 
#     run._r.append(instrText)
#     run._r.append(fldChar1)
#     run._r.append(fldChar2)
# 
#     run._r.append(fldEnd)

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
			
def display_page_numbers(document, book_title='aw'):

	doc_sections = document.sections
	
	first_main_section = doc_sections[0]
	
	if re.search("gb",book_title):
		intro_sect_num = 4
		first_main_section = doc_sections[intro_sect_num] # intro ideally needs page number in different scheme than main content bc the first page should be the first page of main content but it already is the overview of the first chapter which is technically not the main content of the first chapter so for simplicity start page numbering at intro page
	else: # aw or aw,gb
		first_main_section = doc_sections[2]	
	
	first_main_footer = first_main_section.footer
	first_main_footer.is_linked_to_previous = False
	
	doc_sect_footer_block = first_main_footer.paragraphs[0]
	doc_sect_footer_block.style = document.styles['Header Footer']
	doc_sect_footer_block.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	
	add_page_number(doc_sect_footer_block.add_run())
	
	sectPr = first_main_section._sectPr
	
	pgNumType = OxmlElement('w:pgNumType')
	pgNumType.set(ns.qn('w:start'), "1")
	sectPr.append(pgNumType)
	
def display_table(document, ids, parts):
		
	num_rows = len(ids)
	num_cols = 2
	table = document.add_table(num_rows, num_cols)
	
	if len(ids) == len(parts):
	
		for row_idx in range(len(ids)):
			id = ids[row_idx]
			part = parts[row_idx]
			
			id_cell = table.cell(row_idx,0)
			part_cell = table.cell(row_idx,1)
			
			id_cell.text = id
			part_cell.text = part
			
			id_cell.width = Inches(0.5)
			part_cell.width = Inches(6)
			
			# if id has two dots then change margin for sub-bullet
			if id.count(".") == 2:
				set_cell_margins(part_cell, start=720)
			
	else:
		print("Warning: No. IDs different than no. parts!")
		
def display_ordered_list_table(document, ids, parts):
		
	num_rows = len(ids)
	num_cols = 2
	table = document.add_table(num_rows, num_cols)
	
	if len(ids) == len(parts):
	
		for row_idx in range(len(ids)):
			id = ids[row_idx]
			part = parts[row_idx]
			
			id_cell = table.cell(row_idx,0)
			part_cell = table.cell(row_idx,1)
			
			id_cell.text = id
			part_cell.text = part
			
			id_cell.width = Inches(0.5)
			part_cell.width = Inches(6)
			
			# if id has two dots then change margin for sub-bullet
			if id.count(".") == 2:
				set_cell_margins(part_cell, start=720)
		
def add_formatted_paragraph(document, p_string, p_style="Normal"):

	raw_keywords = reader.extract_data("gb-keywords","Supplements","tsv")
	keywords = []
	for key in raw_keywords:
		keywords.append(key[0])
	#keywords = ["Sun Zi\'s Art of War"] # needed to gen runs. depends on context

	# if paragraph ends with number, take it as reference and make it superscript
	p_string = p_string.strip()
	# sub special characters
	p_string = reader.sub_special_characters(p_string)
	ref_nums = ''
	p_without_ref_nums = p_string
	if re.search("\d$",p_string):
		p_parts = p_string.split(".")
		ref_nums = p_parts[len(p_parts)-1]
		#print("ref_nums: " + ref_nums)
		p_without_ref_nums = re.sub("\d+(,\d+)*$","",p_string)
		#print("p_without_ref_nums: " + p_without_ref_nums)
		
	paragraph = document.add_paragraph(style=p_style)
	
	generator.generate_runs(p_without_ref_nums, paragraph, keywords) # gen runs to italicize keywords
	
	if len(ref_nums) > 0:
		ref_run = paragraph.add_run(ref_nums)
		ref_run.font.superscript = True
		
def format_cell_paragraph(cell_p, p_string):

	print("\n=== Format Cell Paragraph: " + p_string + " ===\n")

	raw_keywords = reader.extract_data("gb-keywords","Supplements","tsv")
	keywords = []
	for key in raw_keywords:
		keywords.append(key[0])
	#keywords = ["Sun Zi\'s Art of War"] # needed to gen runs. depends on context

	# if paragraph ends with number, take it as reference and make it superscript
	p_string = p_string.strip()
	# sub special characters
	p_string = reader.sub_special_characters(p_string)
	ref_nums = ''
	p_without_ref_nums = p_string
	if re.search("\d$",p_string):
		p_parts = p_string.split(".")
		ref_nums = p_parts[len(p_parts)-1]
		#print("ref_nums: " + ref_nums)
		p_without_ref_nums = re.sub("\d+(,\d+)*$","",p_string)
		#print("p_without_ref_nums: " + p_without_ref_nums)
	
	generator.generate_runs(p_without_ref_nums, cell_p, keywords) # gen runs to italicize keywords
	
	if len(ref_nums) > 0:
		ref_run = cell_p.add_run(ref_nums)
		ref_run.font.superscript = True